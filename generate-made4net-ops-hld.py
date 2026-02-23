#!/usr/bin/env python3
"""
Made4Net Operational Excellence Architecture - HLD Generator
Focus: How hosting team manages 800+ warehouse endpoints
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_ops_hld():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========================================
    # COVER PAGE
    # ========================================
    title = doc.add_heading('Made4Net Systems', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Operational Excellence Architecture')
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(30, 60, 114)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('Managing 800+ Warehouse Endpoints at Scale')
    tagline_run.font.size = Pt(16)
    tagline_run.font.italic = True
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n')

    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Document Type:', 'High-Level Architecture Design'),
        ('Focus Area:', 'Operational Excellence & Remote Management'),
        ('Target Role:', 'Global Hosting Team Manager'),
        ('Prepared For:', 'Sagi Van - Made4Net Leadership'),
        ('Date:', datetime.now().strftime('%B %d, %Y'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================
    # EXECUTIVE SUMMARY
    # ========================================
    doc.add_heading('Executive Summary', 1)
    
    exec_summary = [
        "This architecture addresses the core operational challenge: How does a hosting team "
        "efficiently manage, monitor, troubleshoot, and maintain 800+ warehouse endpoints "
        "distributed globally, while ensuring security, performance, and minimal downtime?",
        
        "The solution follows AWS best practices with a multi-account architecture and focuses on six operational pillars:",
        
        "1. CONNECTIVITY: Secure, high-performance connections from warehouses to AWS cloud",
        "2. REMOTE ACCESS: Zero-exposure remote management without SSH/RDP ports",
        "3. CENTRALIZED MONITORING: Single pane of glass for all endpoints",
        "4. TROUBLESHOOTING: Rapid diagnosis and resolution workflows",
        "5. DEPLOYMENT & PATCHING: Automated, zero-downtime updates",
        "6. MULTI-ACCOUNT ARCHITECTURE: Enterprise-grade security and operational isolation",
        
        "Additionally, this document outlines the vision for establishing a Global Hosting Team "
        "with clear roles, responsibilities, and operational processes to manage this infrastructure at scale.",
        
        "This architecture eliminates manual operations, reduces MTTR from hours to minutes, "
        "and provides the hosting team with enterprise-grade tools to manage infrastructure at scale."
    ]
    
    for para in exec_summary:
        if para.startswith(('1.', '2.', '3.', '4.', '5.')):
            p = doc.add_paragraph(para)
            p.style = 'List Number'
        else:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

    
    # ========================================
    # PILLAR 1: CONNECTIVITY
    # ========================================
    doc.add_heading('1. Connectivity Architecture', 1)
    doc.add_heading('How 800+ Warehouses Connect to AWS Cloud', 2)
    
    doc.add_paragraph(
        "The connectivity layer ensures secure, high-performance, and resilient connections "
        "between warehouse endpoints and AWS-hosted infrastructure."
    )
    
    doc.add_heading('1.1 Connection Options', 2)
    
    connectivity_table = doc.add_table(rows=4, cols=4)
    connectivity_table.style = 'Medium Grid 3 Accent 1'
    
    connectivity_data = [
        ('Method', 'Use Case', 'Performance', 'Security'),
        ('AWS Site-to-Site VPN', 'Standard warehouses (most common)', 
         '1.25 Gbps per tunnel, <100ms latency', 'IPsec encryption, MFA'),
        ('AWS Direct Connect', 'High-volume warehouses (critical sites)', 
         '1-100 Gbps dedicated, <10ms latency', 'Private connection, no internet'),
        ('AWS Client VPN', 'Individual admin access', 
         'Up to 2 Gbps shared', 'TLS 1.2+, certificate-based')
    ]
    
    for i, row_data in enumerate(connectivity_data):
        for j, cell_text in enumerate(row_data):
            cell = connectivity_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('1.2 Transit Gateway - The Hub', 2)
    
    tgw_points = [
        "AWS Transit Gateway acts as a central hub connecting all 800+ warehouses to a single Production VPC",
        
        "Multi-Tenancy Architecture:",
        "• Single Production VPC (not 800 VPCs - operationally efficient)",
        "• Transit Gateway with 800+ VPN attachments",
        "• Route table per customer for network isolation",
        "• Customer A's warehouse can only route to Customer A's application subnet",
        "• Application-level multi-tenancy with schema-per-tenant database",
        
        "Key Features:",
        "• Network Segmentation: Route tables enforce customer isolation",
        "• Scalability: Supports 5,000 attachments (VPNs, VPCs, Direct Connects)",
        "• Performance: 50 Gbps per VPC attachment",
        "• Monitoring: CloudWatch metrics for bandwidth, packet loss, connection status",
        "• Cost Efficient: Single VPC saves $360K/year vs VPC-per-tenant"
    ]
    
    for point in tgw_points:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('1.3 Performance Optimization', 2)
    
    perf_table = doc.add_table(rows=5, cols=2)
    perf_table.style = 'Light List Accent 1'
    
    perf_data = [
        ('Optimization', 'Implementation'),
        ('Accelerated VPN', 'AWS Global Accelerator for VPN endpoints (up to 60% latency reduction)'),
        ('Multi-Tunnel VPN', 'ECMP (Equal-Cost Multi-Path) for 2.5 Gbps aggregate bandwidth'),
        ('Regional Endpoints', 'Deploy VPN endpoints in closest AWS region to warehouse'),
        ('QoS Policies', 'Prioritize real-time warehouse traffic over batch jobs')
    ]
    
    for i, (opt, impl) in enumerate(perf_data):
        perf_table.rows[i].cells[0].text = opt
        perf_table.rows[i].cells[1].text = impl
        if i == 0:
            perf_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            perf_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('1.4 Connectivity Use Cases', 2)
    
    doc.add_paragraph(
        "Below are the 3 options for the connectivity:"
    )
    
    doc.add_paragraph(
        "The architecture supports three primary connectivity patterns for warehouse operations:"
    )
    
    doc.add_paragraph()
    
    # Case 1: End User Connectivity
    doc.add_heading('Case 1: End User (Tenant Login)', 3)
    
    end_user_flow = [
        "Scenario: Warehouse manager logs into Made4Net WMS application",
        
        "Connection Flow:",
        "1. User opens web browser and navigates to https://customer-a.made4net.com",
        "2. Request routes through Cloudflare proxy for optimization and firewall protection",
        "3. Cloudflare forwards to Application Load Balancer (ALB) in Production VPC",
        "4. ALB performs SSL termination and routes to appropriate tenant",
        "5. Request reaches EC2 Auto Scaling Group (tenant-specific instances)",
        "6. Application authenticates user via Amazon Cognito (multi-tenant SSO)",
        "7. Session established with tenant-isolated database (RDS)",
        
        "Security Features:",
        "• Cloudflare WAF (Web Application Firewall) for edge protection",
        "• DDoS protection via Cloudflare",
        "• SSL/TLS encryption end-to-end",
        "• Cognito user pools for tenant isolation",
        "• IAM roles for fine-grained access control",
        
        "Performance:",
        "• Cloudflare caching reduces latency (static assets served from edge)",
        "• ALB distributes load across multiple availability zones",
        "• Auto Scaling adjusts capacity based on demand",
        "• Typical response time: <200ms for UI, <50ms for API calls"
    ]
    
    for point in end_user_flow:
        if point.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
            doc.add_paragraph(point, style='List Number')
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Case 2: IoT Device Connectivity
    doc.add_heading('Case 2: IoT Devices (Robots, Sensors, Smart Shelves)', 3)
    
    iot_flow = [
        "Scenario: Warehouse robot sends inventory data to cloud",
        
        "Connection Flow:",
        "1. IoT device (robot/sensor/smart shelf) connects to local warehouse network",
        "2. Device authenticates to AWS IoT Core using X.509 certificates",
        "3. Establishes MQTT connection over TLS 1.2+ (port 8883)",
        "4. Publishes telemetry data to IoT topic: warehouse/robot-123/telemetry",
        "5. IoT Core routes message via IoT Rules Engine",
        "6. Data flows to multiple destinations:",
        "   • Real-time: Kinesis Data Streams → Lambda → DynamoDB (hot path)",
        "   • Analytics: Kinesis Firehose → S3 → Athena (cold path)",
        "   • Alerts: IoT Events → SNS → Operations team (anomaly detection)",
        "7. Application queries DynamoDB for real-time inventory status",
        
        "Note: Legacy systems do not have Outposts or external endpoints",
        
        "Security Features:",
        "• Device certificates managed by AWS IoT Device Management",
        "• Certificate rotation every 90 days (automated)",
        "• IoT Device Defender monitors device behavior for anomalies",
        "• Network isolation via VPC and security groups",
        "• Encrypted at rest (S3, DynamoDB) and in transit (TLS)",
        
        "Device Types & Protocols:",
        "• Autonomous Robots: MQTT over TLS (position, battery, task status)",
        "• RFID Sensors: HTTPS REST API (tag reads, inventory counts)",
        "• Smart Shelves: MQTT (weight sensors, stock levels)",
        "• Barcode Scanners: WebSocket over TLS (real-time scan events)",
        "• Temperature Sensors: MQTT (cold storage monitoring)",
        
        "Performance & Scalability:",
        "• IoT Core supports 1 million concurrent connections per account",
        "• Message throughput: 20,000 messages/second per connection",
        "• Typical latency: <50ms (device → IoT Core → Lambda → DynamoDB)",
        "• Auto-scaling for Lambda and DynamoDB based on load"
    ]
    
    for point in iot_flow:
        if point.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
            if '\n' in point:
                p = doc.add_paragraph(point.split('\n')[0])
                for line in point.split('\n')[1:]:
                    doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(point, style='List Number')
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Case 3: Hosting Engineer Connectivity
    doc.add_heading('Case 3: Hosting Engineer (Troubleshooting & Support)', 3)
    
    engineer_flow = [
        "Scenario: Hosting engineer needs to troubleshoot a slow warehouse application",
        
        "Connection Flow:",
        "1. Engineer opens AWS Management Console (https://console.aws.amazon.com)",
        "2. Authenticates with IAM credentials + MFA",
        "3. Navigates to AWS Systems Manager",
        "4. Opens Fleet Manager to view all managed instances",
        "5. Selects target instance (e.g., warehouse-app-server-247)",
        "6. Initiates Session Manager connection (no SSH/RDP ports required)",
        "7. Secure shell session established via SSM Agent",
        "8. Engineer runs diagnostic commands (top, ps, logs)",
        "9. Alternatively, uses CloudWatch for metrics or X-Ray for traces",
        "10. Session logged to S3 for audit compliance",
        
        "Troubleshooting Tools:",
        "• Fleet Manager: View instance inventory, health status, patch compliance",
        "• Session Manager: Secure shell access without SSH/RDP ports",
        "• CloudWatch: View metrics (CPU, memory, disk), query logs",
        "• X-Ray: Distributed tracing for application performance",
        "• Run Command: Execute commands on multiple instances simultaneously",
        
        "Security Features:",
        "• No SSH/RDP ports exposed to internet (zero attack surface)",
        "• IAM-based access control (who can access which instances)",
        "• MFA enforcement for production access",
        "• Session logging to S3 (complete audit trail)",
        "• Encrypted connections via AWS API (TLS 1.2+)",
        
        "Performance:",
        "• Session establishment: <5 seconds",
        "• Command execution latency: <100ms",
        "• CloudWatch query response: <2 seconds",
        "• X-Ray trace retrieval: <3 seconds",
        
        "Access Control Example:",
        "• Junior Engineer: Read-only access to dev/test instances",
        "• Senior Engineer: Full access to dev/test, read-only to production",
        "• Team Lead: Full access to all environments",
        "• Auditor: View session logs only, no instance access"
    ]
    
    for point in engineer_flow:
        if point.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.')):
            doc.add_paragraph(point, style='List Number')
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Comparison Table
    doc.add_heading('Connectivity Comparison', 3)
    
    comparison_table = doc.add_table(rows=7, cols=4)
    comparison_table.style = 'Medium Grid 3 Accent 1'
    
    comparison_data = [
        ('Aspect', 'End User (Tenant Login)', 'IoT Device (Robot/Sensor)', 'Hosting Engineer (Troubleshooting)'),
        ('Protocol', 'HTTPS (443)', 'MQTT over TLS (8883)', 'HTTPS (443)'),
        ('Authentication', 'Cognito (username/password + MFA)', 'X.509 certificates', 'IAM credentials + MFA'),
        ('Entry Point', 'Cloudflare → ALB', 'AWS IoT Core', 'AWS Console → Systems Manager'),
        ('Latency Requirement', '<200ms (interactive)', '<50ms (real-time telemetry)', '<100ms (interactive)'),
        ('Data Volume', 'Low (KB per request)', 'High (MB per second aggregate)', 'Low (KB per request)'),
        ('Connection Pattern', 'Short-lived (request/response)', 'Long-lived (persistent MQTT)', 'Interactive (session-based)')
    ]
    
    for i, row_data in enumerate(comparison_data):
        for j, cell_text in enumerate(row_data):
            cell = comparison_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # PILLAR 2: REMOTE ACCESS
    # ========================================
    doc.add_heading('2. Remote Access Architecture', 1)
    doc.add_heading('How Hosting Team Connects to Instances (No SSH/RDP Exposure)', 2)
    
    doc.add_paragraph(
        "Traditional remote access (SSH port 22, RDP port 3389) exposes infrastructure to attacks. "
        "This architecture uses AWS Systems Manager Session Manager for zero-exposure access."
    )
    
    doc.add_heading('2.1 AWS Systems Manager Session Manager', 2)
    
    ssm_benefits = [
        "How It Works:",
        "1. Hosting team opens AWS Console or uses AWS CLI",
        "2. Selects target instance from Fleet Manager",
        "3. Initiates session through SSM Agent (no inbound ports required)",
        "4. Encrypted connection established via AWS API",
        "5. Full shell access (bash for Linux, PowerShell for Windows)",
        
        "Security Benefits:",
        "• No SSH/RDP ports exposed to internet",
        "• No bastion hosts to manage",
        "• IAM-based access control (who can access which instances)",
        "• Session logging to S3 for audit compliance",
        "• MFA enforcement for production access"
    ]
    
    for point in ssm_benefits:
        if point.startswith(('1.', '2.', '3.', '4.', '5.')):
            doc.add_paragraph(point, style='List Number')
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 Access Methods Comparison', 2)
    
    access_table = doc.add_table(rows=4, cols=4)
    access_table.style = 'Medium Grid 3 Accent 1'
    
    access_data = [
        ('Method', 'Security Risk', 'Audit Trail', 'Hosting Team Experience'),
        ('Traditional SSH/RDP', 'HIGH (exposed ports)', 'Limited', 'Requires VPN + credentials'),
        ('Bastion Host', 'MEDIUM (single point)', 'Partial', 'Extra hop, complex'),
        ('SSM Session Manager', 'LOW (no exposure)', 'Complete', 'One-click from console')
    ]
    
    for i, row_data in enumerate(access_data):
        for j, cell_text in enumerate(row_data):
            cell = access_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('2.3 Role-Based Access Control', 2)
    
    rbac_example = [
        "IAM policies define who can access which instances:",
        
        "• Junior Admin: Read-only access to dev/test instances",
        "• Senior Admin: Full access to dev/test, read-only to production",
        "• Team Lead: Full access to all environments",
        "• Auditor: View session logs only, no instance access",
        
        "Example: A junior admin attempting to access production will be denied by IAM policy, "
        "and the attempt is logged to CloudTrail for security review."
    ]
    
    for point in rbac_example:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_page_break()

    
    # ========================================
    # PILLAR 3: CENTRALIZED MONITORING
    # ========================================
    doc.add_heading('3. Centralized Monitoring Architecture', 1)
    doc.add_heading('Single Pane of Glass for 800+ Endpoints', 2)
    
    doc.add_paragraph(
        "The hosting team needs visibility across all warehouse endpoints and hosted instances "
        "from a single dashboard. This architecture provides real-time monitoring, alerting, "
        "and historical analysis."
    )
    
    doc.add_heading('3.1 AWS Systems Manager Fleet Manager', 2)
    
    fleet_features = [
        "Fleet Manager provides a unified view of all managed instances:",
        
        "• Instance Inventory: OS version, installed software, patch status",
        "• Health Status: CPU, memory, disk usage across all instances",
        "• Compliance Dashboard: Which instances are out of compliance",
        "• Grouping: View by customer, region, environment (dev/test/prod)",
        "• Quick Actions: Start session, run command, view logs - all from one screen"
    ]
    
    for point in fleet_features:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('3.2 CloudWatch Unified Dashboard', 2)
    
    cw_metrics = [
        "Custom CloudWatch dashboard showing:",
        
        "Warehouse Connectivity:",
        "• VPN tunnel status (up/down) per warehouse",
        "• Bandwidth utilization and packet loss",
        "• Connection latency trends",
        
        "Instance Health:",
        "• CPU, memory, disk usage per instance",
        "• Application-specific metrics (database connections, queue depth)",
        "• Error rates and response times",
        
        "Business Metrics:",
        "• Number of active warehouses",
        "• Transactions per second across all sites",
        "• Failed transactions requiring investigation"
    ]
    
    for point in cw_metrics:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('3.3 Alerting Strategy', 2)
    
    alert_table = doc.add_table(rows=4, cols=4)
    alert_table.style = 'Medium Grid 3 Accent 1'
    
    alert_data = [
        ('Alert Type', 'Trigger', 'Notification', 'Response Time'),
        ('P1 - Critical', 'Warehouse offline, database down', 
         'PagerDuty → On-call team', 'Immediate (24/7)'),
        ('P2 - High', 'High CPU (>80%), disk full warning', 
         'Slack + Email → Team channel', '< 30 minutes'),
        ('P3 - Medium', 'Patch compliance drift, backup failure', 
         'Email → Team lead', 'Next business day')
    ]
    
    for i, row_data in enumerate(alert_data):
        for j, cell_text in enumerate(row_data):
            cell = alert_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # PILLAR 4: TROUBLESHOOTING
    # ========================================
    doc.add_heading('4. Troubleshooting Architecture', 1)
    doc.add_heading('How Hosting Team Diagnoses and Resolves Issues', 2)
    
    doc.add_paragraph(
        "When a warehouse reports 'the system is slow' or an alert fires, the hosting team "
        "needs tools to quickly identify the root cause. This architecture provides a "
        "systematic troubleshooting workflow."
    )
    
    doc.add_heading('4.1 Troubleshooting Workflow', 2)
    
    workflow_steps = [
        "Step 1: Identify Scope",
        "• Is it one warehouse or multiple?",
        "• Is it one customer or all customers?",
        "• Check CloudWatch dashboard for patterns",
        
        "Step 2: Check Connectivity",
        "• VPN tunnel status in Transit Gateway console",
        "• Packet loss and latency metrics",
        "• Network ACLs and Security Group rules",
        
        "Step 3: Check Application Layer",
        "• AWS X-Ray traces for slow API calls",
        "• Application logs in CloudWatch Logs Insights",
        "• Database query performance (RDS Performance Insights)",
        
        "Step 4: Check Infrastructure",
        "• EC2 instance CPU/memory/disk via CloudWatch",
        "• Auto Scaling Group health checks",
        "• Load balancer target health",
        
        "Step 5: Execute Remediation",
        "• SSM Run Command to restart services",
        "• Auto Scaling to replace unhealthy instances",
        "• Route 53 failover to DR region if needed"
    ]
    
    for step in workflow_steps:
        if step.startswith('Step'):
            p = doc.add_paragraph(step)
            p.runs[0].font.bold = True
        elif step.startswith('•'):
            doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('4.2 Key Troubleshooting Tools', 2)
    
    tools_table = doc.add_table(rows=6, cols=3)
    tools_table.style = 'Medium Grid 3 Accent 1'
    
    tools_data = [
        ('Tool', 'Purpose', 'Example Use Case'),
        ('CloudWatch Logs Insights', 'Query logs across all instances', 
         'Find all 500 errors in last hour'),
        ('AWS X-Ray', 'Distributed tracing', 
         'Identify slow database query causing latency'),
        ('SSM Run Command', 'Execute commands on multiple instances', 
         'Restart application service on 50 instances'),
        ('RDS Performance Insights', 'Database query analysis', 
         'Find top 10 slowest SQL queries'),
        ('VPC Flow Logs', 'Network traffic analysis', 
         'Identify blocked connections from warehouse')
    ]
    
    for i, row_data in enumerate(tools_data):
        for j, cell_text in enumerate(row_data):
            cell = tools_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('4.3 Real-World Scenario: "Warehouse 247 is Slow"', 2)
    
    scenario = [
        "Problem: Warehouse 247 reports slow system response.",
        
        "Troubleshooting Steps:",
        "1. Check CloudWatch dashboard → CPU normal, network normal",
        "2. Check X-Ray traces → API calls taking 5 seconds (normally 200ms)",
        "3. Drill into X-Ray → Database query taking 4.8 seconds",
        "4. Check RDS Performance Insights → Missing index on orders table",
        "5. Resolution: Create index via SSM Session Manager to database",
        "6. Verify: X-Ray shows API calls back to 200ms",
        "7. Total time: 8 minutes from alert to resolution",
        
        "Root Cause: Database index missing after recent schema change.",
        "Prevention: Add index creation to deployment checklist."
    ]
    
    for point in scenario:
        if point.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.')):
            doc.add_paragraph(point, style='List Number')
        else:
            p = doc.add_paragraph(point)
            if point.startswith(('Problem:', 'Root Cause:', 'Prevention:')):
                p.runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # PILLAR 5: DEPLOYMENT & PATCHING
    # ========================================
    doc.add_heading('5. Deployment & Patching Architecture', 1)
    doc.add_heading('Automated, Zero-Downtime Updates at Scale', 2)
    
    doc.add_paragraph(
        "Deploying new application versions and OS patches to 800+ endpoints manually is "
        "impossible. This architecture automates deployments with rollback capabilities "
        "and zero downtime."
    )
    
    doc.add_heading('5.1 Application Deployment Pipeline', 2)
    
    deployment_flow = [
        "Deployment Flow (using AWS CodeDeploy):",
        
        "1. Developer commits code to Git repository",
        "2. CodePipeline triggers automatically",
        "3. CodeBuild compiles and runs tests",
        "4. CodeDeploy stages deployment:",
        "   • Dev environment: Deploy to all instances",
        "   • Test environment: Deploy after dev passes",
        "   • Production: Deploy in waves (10% → 50% → 100%)",
        "5. Health checks validate each wave",
        "6. Automatic rollback if health checks fail",
        
        "Deployment Strategies:",
        "• Blue/Green: New version deployed to separate instances, traffic switched atomically",
        "• Rolling: Update instances in batches (e.g., 10 at a time)",
        "• Canary: Deploy to 5% of fleet, monitor, then proceed"
    ]
    
    for point in deployment_flow:
        if point.startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
            doc.add_paragraph(point, style='List Number')
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        elif point.startswith('   •'):
            doc.add_paragraph(point.strip())
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 OS Patch Management (SSM Patch Manager)', 2)
    
    patch_strategy = [
        "Automated Patching Strategy:",
        
        "Patch Baselines:",
        "• Critical Security Patches: Auto-approve, deploy within 24 hours",
        "• Important Updates: Auto-approve, deploy within 7 days",
        "• Optional Updates: Manual approval required",
        
        "Maintenance Windows:",
        "• Dev/Test: Patch Tuesday, 2:00 AM local time",
        "• Production: Following Saturday, 2:00 AM local time (after 4 days of testing)",
        "• Emergency Patches: On-demand with change approval",
        
        "Patching Process:",
        "1. SSM scans all instances for missing patches",
        "2. Patches staged during maintenance window",
        "3. Instances patched in rolling fashion (10% at a time)",
        "4. Health checks validate each batch",
        "5. Automatic reboot if required (during maintenance window)",
        "6. Compliance report generated and reviewed",
        
        "Zero-Downtime Approach:",
        "• Auto Scaling Group maintains minimum capacity",
        "• Load balancer drains connections before patching",
        "• Instance returns to service after health check passes",
        "• No user-facing downtime"
    ]
    
    for point in patch_strategy:
        if point.startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
            doc.add_paragraph(point, style='List Number')
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('5.3 Patch Compliance Tracking', 2)
    
    compliance_table = doc.add_table(rows=4, cols=3)
    compliance_table.style = 'Light List Accent 1'
    
    compliance_data = [
        ('Compliance Level', 'Criteria', 'Action'),
        ('Compliant (Green)', '95%+ instances patched', 'No action required'),
        ('At Risk (Yellow)', '85-94% instances patched', 'Review non-compliant instances'),
        ('Non-Compliant (Red)', '<85% instances patched', 'Immediate remediation required')
    ]
    
    for i, (level, criteria, action) in enumerate(compliance_data):
        compliance_table.rows[i].cells[0].text = level
        compliance_table.rows[i].cells[1].text = criteria
        compliance_table.rows[i].cells[2].text = action
        if i == 0:
            for j in range(3):
                compliance_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # PILLAR 6: AWS OUTPOSTS - HYBRID ON-PREMISES
    # ========================================
    doc.add_heading('6. AWS Outposts - Hybrid On-Premises Architecture', 1)
    doc.add_heading('Managing Warehouses with Local Compute Requirements', 2)
    
    doc.add_paragraph(
        "Some warehouses require on-premises compute due to low-latency requirements, "
        "data residency regulations, or local system dependencies. AWS Outposts extends "
        "AWS infrastructure to warehouse locations while maintaining the same operational "
        "model as cloud-based resources."
    )
    
    doc.add_heading('6.1 When to Use AWS Outposts', 2)
    
    outposts_use_cases = [
        "Use Cases for Outposts Deployment:",
        
        "• Low Latency Requirements: Warehouse automation systems requiring <10ms response times",
        "• Data Residency: Regulatory requirements to keep data on-premises",
        "• Local Data Processing: Real-time processing of IoT sensor data from warehouse devices",
        "• Migration Strategy: Gradual cloud migration while maintaining local dependencies",
        "• High-Volume Sites: Critical distribution centers with 24/7 operations"
    ]
    
    for point in outposts_use_cases:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 Outposts Monitoring Architecture', 2)
    
    outposts_monitoring = [
        "Critical Monitoring Areas:",
        
        "Capacity Monitoring:",
        "• EC2 Instance Capacity: Track available instance types and utilization",
        "• EBS Volume Capacity: Monitor storage consumption and available space",
        "• CapacityExceptions Metric: Alert when capacity limits are reached",
        "• Recommendation: Maintain N+M availability model (N required + M spare servers)",
        
        "Network Connectivity:",
        "• Service Link Status: Monitor ConnectedStatus metric (Outpost to AWS Region)",
        "• Local Gateway BGP: Monitor BGP peering status with on-premises network",
        "• Traffic Metrics: Track IfTrafficIn/IfTrafficOut for bandwidth utilization",
        "• Local Network Interface (LNI): Monitor connectivity to warehouse network",
        
        "Health Events:",
        "• AWS_EC2_INSTANCE_RETIREMENT_SCHEDULED: Hardware failure requiring replacement",
        "• AWS_OUTPOSTS_SERVICE_LINK_DOWN: Loss of connectivity to AWS Region",
        "• Impact: Service link down prevents new resource creation but existing workloads continue"
    ]
    
    for point in outposts_monitoring:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('6.3 Outposts Operational Best Practices', 2)
    
    outposts_table = doc.add_table(rows=6, cols=3)
    outposts_table.style = 'Medium Grid 3 Accent 1'
    
    outposts_data = [
        ('Practice Area', 'Implementation', 'Benefit'),
        ('Capacity Planning', 'Order N+M servers (spare capacity for failures)', 
         'Ensure failover capacity during hardware issues'),
        ('Service Link Monitoring', 'CloudWatch alarm on ConnectedStatus metric', 
         'Immediate alert when Region connectivity lost'),
        ('Cross-Account Sharing', 'Use AWS RAM to share Outposts across accounts', 
         'Multiple teams use same Outpost securely'),
        ('Centralized Observability', 'CloudWatch cross-account observability', 
         'Single dashboard for all Outpost metrics'),
        ('Event Automation', 'AWS Health Aware for custom notifications', 
         'Route alerts to correct teams automatically')
    ]
    
    for i, row_data in enumerate(outposts_data):
        for j, cell_text in enumerate(row_data):
            cell = outposts_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('6.4 Outposts Troubleshooting Workflow', 2)
    
    outposts_troubleshooting = [
        "Scenario: Outpost Service Link Down",
        
        "1. Alert received: AWS_OUTPOSTS_SERVICE_LINK_DOWN Health event",
        "2. Verify: Check ConnectedStatus CloudWatch metric (confirms link down)",
        "3. Impact assessment:",
        "   • Existing EC2 instances continue running",
        "   • New resource creation (ec2 run-instances) will fail",
        "   • CloudWatch metrics may not update in real-time",
        "4. Escalate to network team: Use Outposts network troubleshooting checklist",
        "5. Check physical connectivity:",
        "   • Verify network cables and switches",
        "   • Check BGP peering status",
        "   • Review firewall rules for AWS API endpoints",
        "6. Monitor restoration: ConnectedStatus returns to 'Connected'",
        "7. Verify operations: Test new resource creation",
        "8. Post-incident: Document root cause and prevention steps"
    ]
    
    for step in outposts_troubleshooting:
        if step.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.')):
            if '\n' in step:
                p = doc.add_paragraph(step.split('\n')[0])
                for line in step.split('\n')[1:]:
                    doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(step, style='List Number')
        else:
            p = doc.add_paragraph(step)
            if step.startswith('Scenario:'):
                p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('6.5 Outposts Monitoring Tools', 2)
    
    outposts_tools = [
        "AWS Tools for Outposts Observability:",
        
        "• CloudWatch Metrics: Capacity, network traffic, service link status",
        "• CloudWatch Dashboards: Centralized view across multiple Outposts",
        "• AWS Health Events: Critical alerts for hardware failures and connectivity",
        "• CloudTrail Logs: API call auditing for Outposts resources",
        "• VPC Flow Logs: Network traffic analysis within Outpost",
        "• Traffic Mirroring: Copy traffic to security appliances for inspection",
        "• AWS X-Ray: Distributed tracing for applications on Outposts",
        "• CloudWatch Synthetics: Canaries to test connectivity (Region ↔ Outpost)",
        "• AWS Health Aware: Custom notifications across accounts and teams"
    ]
    
    for point in outposts_tools:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            if point.endswith(':'):
                p.runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # PILLAR 7: MULTI-ACCOUNT ARCHITECTURE
    # ========================================
    doc.add_heading('7. Multi-Account Architecture', 1)
    doc.add_heading('Enterprise-Grade Security and Operational Isolation', 2)
    
    doc.add_paragraph(
        "Following AWS best practices, the Made4Net architecture uses a multi-account strategy "
        "for security isolation, compliance, and operational excellence. This approach separates "
        "security services, operational tools, and workloads into dedicated AWS accounts."
    )
    
    doc.add_heading('7.1 Account Structure Overview', 2)
    
    account_structure = [
        "The architecture consists of 6 AWS accounts organized under AWS Organizations:",
        
        "1. Management Account: AWS Organizations root, billing, Control Tower",
        "2. Security Account: GuardDuty, Inspector, Config, Security Hub, CloudTrail",
        "3. Operations Account: Systems Manager, CloudWatch, X-Ray, Backup",
        "4. Production Account: VPC, Transit Gateway, application workloads",
        "5. Outposts Account #1: Warehouse Group A (NY, Boston, Philadelphia)",
        "6. Outposts Account #2: Warehouse Group B (Chicago, Detroit, Milwaukee)",
        
        "Additional accounts:",
        "• DR Account: Disaster recovery resources in us-west-2"
    ]
    
    for point in account_structure:
        if point.startswith(('1.', '2.', '3.', '4.', '5.', '6.')):
            doc.add_paragraph(point, style='List Number')
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('7.2 Security Account - Centralized Security Monitoring', 2)
    
    security_account = [
        "The Security Account monitors all other accounts for threats, vulnerabilities, "
        "and compliance violations:",
        
        "Amazon GuardDuty:",
        "• Threat detection across all accounts",
        "• Monitors VPC Flow Logs from Production and Outposts accounts",
        "• Detects SSH brute force, port scanning, malicious IPs",
        "• Delegated administrator for organization-wide monitoring",
        
        "Amazon Inspector:",
        "• Vulnerability scanning across all accounts (including Outposts)",
        "• Uses SSM Agent to collect software inventory",
        "• Continuous CVE scanning and network exposure detection",
        "• Automated rescanning when new vulnerabilities published",
        
        "AWS Config:",
        "• Compliance monitoring across all accounts",
        "• Tracks configuration changes",
        "• Enforces compliance rules (encryption, security groups)",
        "• Aggregated view of all resources",
        
        "AWS Security Hub:",
        "• Centralized security findings from GuardDuty, Inspector, Config",
        "• Security standards compliance (CIS, PCI-DSS)",
        "• Automated remediation workflows",
        
        "AWS CloudTrail:",
        "• Organization trail logs all API calls across all accounts",
        "• Immutable audit log stored in dedicated S3 bucket",
        "• Encrypted with KMS for compliance"
    ]
    
    for point in security_account:
        if point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('7.3 Operations Account - Centralized Management', 2)
    
    operations_account = [
        "The Operations Account provides centralized management and monitoring for all workloads:",
        
        "AWS Systems Manager:",
        "• Fleet Manager: Unified view of all EC2 instances across accounts",
        "• Session Manager: Secure remote access to instances",
        "• Patch Manager: Automated patching across all accounts",
        "• Run Command: Bulk operations on multiple instances",
        "• Cross-account access to Production and Outposts accounts",
        
        "Amazon CloudWatch:",
        "• Centralized monitoring dashboard for all accounts",
        "• Cross-account observability",
        "• Unified alarms and notifications",
        "• Log aggregation via CloudWatch Logs Insights",
        
        "AWS Backup:",
        "• Centralized backup management",
        "• Cross-account backup policies",
        "• Cross-region backup replication",
        "• Compliance reporting"
    ]
    
    for point in operations_account:
        if point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('7.4 Outposts Accounts - Hybrid Deployment Isolation', 2)
    
    outposts_accounts = [
        "Outposts are deployed in separate accounts for security and billing isolation:",
        
        "Why Separate Accounts:",
        "• Billing Isolation: Track Outposts costs per warehouse group",
        "• Security Boundary: Isolate on-premises resources",
        "• Compliance: Meet data residency requirements",
        "• Blast Radius: Limit impact of security incidents",
        
        "VPC Extension Model:",
        "• VPC extends from AWS Region to Outpost (same VPC, different subnets)",
        "• Outpost subnet: 10.0.10.0/24 (part of Production VPC 10.0.0.0/16)",
        "• Seamless communication between cloud and on-premises resources",
        
        "VPC Endpoints (PrivateLink):",
        "• ssm: Systems Manager service",
        "• ec2messages: SSM Agent communication",
        "• ssmmessages: Session Manager",
        "• ec2: EC2 API calls",
        "• s3: S3 access for patches/logs",
        "• inspector2: Inspector scanning",
        "• Benefit: No internet gateway required, enhanced security",
        
        "Service Link:",
        "• Connectivity from Outpost to AWS Region",
        "• Minimum: 500 Mbps, Recommended: 1 Gbps+",
        "• Monitored via ConnectedStatus CloudWatch metric",
        
        "Cross-Account Monitoring:",
        "• Security Account monitors Outposts via GuardDuty and Inspector",
        "• Operations Account manages Outposts via Systems Manager",
        "• Unified dashboard shows all accounts"
    ]
    
    for point in outposts_accounts:
        if point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('7.5 Multi-Account Benefits', 2)
    
    benefits_table = doc.add_table(rows=5, cols=3)
    benefits_table.style = 'Medium Grid 3 Accent 1'
    
    benefits_data = [
        ('Benefit Category', 'Implementation', 'Business Value'),
        ('Security Isolation', 'Breach in one account does not affect others', 
         'Reduced blast radius, improved security posture'),
        ('Cost Allocation', 'Per-account cost tracking and budgets', 
         'Accurate chargeback to business units'),
        ('Compliance', 'Separate accounts for data residency requirements', 
         'Meet regulatory requirements (GDPR, HIPAA)'),
        ('Operational Excellence', 'Centralized monitoring and management', 
         'Single pane of glass for 800+ endpoints')
    ]
    
    for i, row_data in enumerate(benefits_data):
        for j, cell_text in enumerate(row_data):
            cell = benefits_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('7.6 Unified Security Dashboard Example', 2)
    
    dashboard_example = [
        "The Security Account provides a unified view of security posture across all accounts:",
        
        "GuardDuty Findings (Last 24 Hours):",
        "• Production Account: 0 high, 1 medium",
        "• Outposts Account #1: 1 high (SSH brute force), 2 medium",
        "• Outposts Account #2: 0 high, 0 medium",
        
        "Inspector Vulnerabilities:",
        "• Production Account: 5 high, 12 medium",
        "• Outposts Account #1: 12 high, 19 medium",
        "• Outposts Account #2: 8 high, 15 medium",
        
        "Config Compliance:",
        "• Production Account: 98% compliant",
        "• Outposts Account #1: 95% compliant",
        "• Outposts Account #2: 97% compliant",
        
        "Overall Security Hub Score: 95/100"
    ]
    
    for point in dashboard_example:
        if point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        elif point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            doc.add_paragraph(point)
    
    doc.add_page_break()

    
    # ========================================
    # OPERATIONAL WORKFLOWS
    # ========================================
    doc.add_heading('8. Day-to-Day Operational Workflows', 1)
    
    doc.add_paragraph(
        "This section describes typical daily activities for the hosting team using "
        "the architecture described above."
    )
    
    doc.add_heading('8.1 Morning Health Check (15 minutes)', 2)
    
    morning_check = [
        "1. Open CloudWatch Unified Dashboard",
        "   • Review overnight alerts (any P1/P2 incidents?)",
        "   • Check VPN tunnel status (all 800+ warehouses connected?)",
        "   • Verify Auto Scaling Group health (any instances replaced?)",
        
        "2. Open Systems Manager Fleet Manager",
        "   • Review patch compliance score (target: 95%+)",
        "   • Check for instances with high CPU/memory/disk",
        "   • Verify SSM Agent status (all instances reporting?)",
        
        "3. Review CloudWatch Logs Insights",
        "   • Query for application errors in last 24 hours",
        "   • Identify any patterns or spikes",
        
        "4. Check AWS Backup Dashboard",
        "   • Verify all backups completed successfully",
        "   • Review cross-region replication status"
    ]
    
    for step in morning_check:
        if step.startswith(('1.', '2.', '3.', '4.')):
            p = doc.add_paragraph(step.split('\n')[0])
            p.runs[0].font.bold = True
            if '\n' in step:
                for line in step.split('\n')[1:]:
                    doc.add_paragraph(line.strip(), style='List Bullet')
    
    doc.add_paragraph()
    
    doc.add_heading('8.2 Responding to Alert: "High CPU on Instance i-abc123"', 2)
    
    alert_response = [
        "1. Receive alert via PagerDuty/Slack",
        "2. Open CloudWatch dashboard → Identify affected instance",
        "3. Check X-Ray traces → Is application slow?",
        "4. Open SSM Session Manager → Connect to instance",
        "5. Run diagnostic commands:",
        "   • Linux: top, ps aux, df -h, netstat",
        "   • Windows: Task Manager equivalent via PowerShell",
        "6. Identify root cause (e.g., runaway process)",
        "7. Remediate:",
        "   • Option A: Kill process via SSM",
        "   • Option B: Restart instance via EC2 console",
        "   • Option C: Terminate instance (Auto Scaling replaces it)",
        "8. Verify resolution in CloudWatch",
        "9. Document incident in ticketing system",
        "10. Total time: 10-15 minutes"
    ]
    
    for step in alert_response:
        if step.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.')):
            if '\n' in step:
                p = doc.add_paragraph(step.split('\n')[0])
                p.runs[0].font.bold = True if step.startswith(('5.', '7.')) else False
                for line in step.split('\n')[1:]:
                    doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(step, style='List Number')
    
    doc.add_paragraph()
    
    doc.add_heading('8.3 Deploying New Application Version', 2)
    
    deploy_workflow = [
        "1. Developer merges code to main branch",
        "2. CodePipeline automatically triggers",
        "3. Hosting team receives Slack notification: 'Deployment started'",
        "4. Monitor CodeDeploy dashboard:",
        "   • Dev environment: Deployed in 5 minutes",
        "   • Test environment: Deployed in 10 minutes",
        "   • Production: Canary deployment (5% of fleet)",
        "5. Monitor CloudWatch for errors during canary",
        "6. If canary successful:",
        "   • CodeDeploy proceeds to 50% of fleet",
        "   • Then 100% of fleet",
        "7. If canary fails:",
        "   • Automatic rollback to previous version",
        "   • Alert sent to development team",
        "8. Total deployment time: 30-45 minutes",
        "9. Zero downtime for end users"
    ]
    
    for step in deploy_workflow:
        if step.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            if '\n' in step:
                p = doc.add_paragraph(step.split('\n')[0])
                for line in step.split('\n')[1:]:
                    doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(step, style='List Number')
    
    doc.add_page_break()

    
    # ========================================
    # AWS SERVICES SUMMARY
    # ========================================
    doc.add_heading('9. AWS Services Summary', 1)
    
    doc.add_paragraph(
        "This architecture leverages the following AWS services to achieve operational excellence:"
    )
    
    doc.add_paragraph()
    
    services_table = doc.add_table(rows=19, cols=3)
    services_table.style = 'Medium List 1 Accent 1'
    
    services_data = [
        ('AWS Service', 'Operational Purpose', 'Key Benefit'),
        ('Transit Gateway', 'Hub for 800+ warehouse VPN connections', 'Network segmentation & scalability'),
        ('Site-to-Site VPN', 'Secure warehouse-to-cloud connectivity', 'IPsec encryption, low latency'),
        ('Direct Connect', 'Dedicated connection for critical sites', 'Predictable performance'),
        ('AWS Outposts', 'On-premises AWS infrastructure at warehouses', 'Low latency, data residency'),
        ('Systems Manager Session Manager', 'Remote access without SSH/RDP', 'Zero-exposure security'),
        ('Systems Manager Fleet Manager', 'Unified view of all instances', 'Single pane of glass'),
        ('Systems Manager Patch Manager', 'Automated OS patching', 'Zero-downtime updates'),
        ('Systems Manager Run Command', 'Execute commands at scale', 'Bulk operations'),
        ('CloudWatch Dashboards', 'Centralized monitoring', 'Real-time visibility'),
        ('CloudWatch Logs Insights', 'Log analysis and querying', 'Fast troubleshooting'),
        ('CloudWatch Alarms', 'Proactive alerting', 'Reduce MTTR'),
        ('AWS X-Ray', 'Distributed tracing', 'Identify performance bottlenecks'),
        ('AWS Health', 'Service and resource health events', 'Proactive issue detection'),
        ('AWS Health Aware', 'Custom health notifications', 'Multi-account alerting'),
        ('CodePipeline', 'CI/CD automation', 'Consistent deployments'),
        ('CodeDeploy', 'Application deployment', 'Blue/green, canary strategies'),
        ('RDS Performance Insights', 'Database query analysis', 'Optimize slow queries'),
        ('VPC Flow Logs', 'Network traffic analysis', 'Troubleshoot connectivity')
    ]
    
    for i, row_data in enumerate(services_data):
        for j, cell_text in enumerate(row_data):
            cell = services_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================
    # BEST PRACTICES SUMMARY
    # ========================================
    doc.add_heading('10. Operational Best Practices Summary', 1)
    
    best_practices = [
        "Connectivity Best Practices:",
        "• Use Transit Gateway for centralized network management",
        "• Deploy VPN endpoints in region closest to warehouse",
        "• Implement multi-tunnel VPN for bandwidth aggregation",
        "• Monitor VPN health with CloudWatch metrics",
        
        "Remote Access Best Practices:",
        "• Never expose SSH/RDP ports to internet",
        "• Use SSM Session Manager for all remote access",
        "• Implement IAM role-based access control",
        "• Log all sessions to S3 for audit compliance",
        "• Enforce MFA for production access",
        
        "Monitoring Best Practices:",
        "• Create unified CloudWatch dashboard for team",
        "• Set up tiered alerting (P1/P2/P3)",
        "• Use CloudWatch Logs Insights for log analysis",
        "• Implement synthetic monitoring with CloudWatch Canaries",
        "• Review metrics daily during morning health check",
        
        "Troubleshooting Best Practices:",
        "• Follow systematic workflow (connectivity → app → infrastructure)",
        "• Use X-Ray for distributed tracing",
        "• Leverage SSM Run Command for bulk diagnostics",
        "• Document all incidents and root causes",
        "• Conduct post-incident reviews",
        
        "Deployment & Patching Best Practices:",
        "• Automate deployments with CodePipeline/CodeDeploy",
        "• Use canary or blue/green deployment strategies",
        "• Patch dev/test before production (4-day soak period)",
        "• Maintain 95%+ patch compliance",
        "• Schedule maintenance windows during off-peak hours",
        "• Always have rollback plan",
        
        "AWS Outposts Best Practices:",
        "• Monitor ConnectedStatus metric with CloudWatch alarms",
        "• Maintain N+M capacity model for hardware failures",
        "• Use AWS Health Aware for multi-account event notifications",
        "• Implement cross-account observability for shared Outposts",
        "• Monitor capacity metrics to prevent resource exhaustion",
        "• Test service link failover procedures regularly",
        "• Document network troubleshooting procedures with on-site teams"
    ]
    
    for practice in best_practices:
        if practice.endswith(':'):
            p = doc.add_paragraph(practice)
            p.runs[0].font.bold = True
        else:
            doc.add_paragraph(practice, style='List Bullet')
    
    doc.add_page_break()

    
    # ========================================
    # METRICS & KPIs
    # ========================================
    doc.add_heading('11. Operational Metrics & KPIs', 1)
    
    doc.add_paragraph(
        "These metrics measure the effectiveness of the operational excellence architecture:"
    )
    
    doc.add_paragraph()
    
    metrics_table = doc.add_table(rows=13, cols=3)
    metrics_table.style = 'Light Grid Accent 1'
    
    metrics_data = [
        ('Metric', 'Target', 'Measurement Method'),
        ('Warehouse Connectivity', '99.9% uptime', 'VPN tunnel status in CloudWatch'),
        ('Outposts Service Link', '99.9% uptime', 'ConnectedStatus CloudWatch metric'),
        ('Outposts Capacity Utilization', '<80%', 'EC2/EBS capacity metrics'),
        ('Mean Time to Detect (MTTD)', '<5 minutes', 'Time from issue to alert'),
        ('Mean Time to Resolve (MTTR)', '8-15 minutes', 'Time from alert to resolution'),
        ('Patch Compliance', '95%+', 'SSM Patch Manager compliance report'),
        ('Deployment Success Rate', '98%+', 'CodeDeploy success/failure ratio'),
        ('Deployment Frequency', 'Daily (dev/test), Weekly (prod)', 'CodePipeline execution count'),
        ('Rollback Rate', '<5%', 'CodeDeploy rollback count'),
        ('Security Incidents', '0 per month', 'GuardDuty findings + manual review'),
        ('Audit Compliance', '100%', 'AWS Config compliance score'),
        ('Team Productivity', '80%+ automation', 'Manual tasks vs automated tasks')
    ]
    
    for i, row_data in enumerate(metrics_data):
        for j, cell_text in enumerate(row_data):
            cell = metrics_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================
    # GLOBAL HOSTING TEAM STRUCTURE
    # ========================================
    doc.add_heading('12. Global Hosting Team Structure & Vision', 1)
    
    doc.add_paragraph(
        "This section outlines the organizational structure and vision for the Global Hosting Team "
        "responsible for managing Made4Net's infrastructure across 800+ warehouse endpoints."
    )
    
    doc.add_heading('12.1 Team Leadership Role', 2)
    
    doc.add_paragraph(
        "Global Hosting Team Manager - Key Responsibilities:"
    )
    
    leadership_responsibilities = [
        "People Management:",
        "• Lead and manage a global hosting team across multiple time zones",
        "• Mentor team members and drive professional development",
        "• Build a culture of operational excellence and continuous improvement",
        "• Conduct performance reviews and career development planning",
        
        "Technical Leadership:",
        "• Oversee day-to-day operations of hosting infrastructure across multiple environments",
        "• Own and manage AWS environments with focus on availability, scalability, and cost efficiency",
        "• Lead resolution of complex incidents and manage critical infrastructure events",
        "• Define and implement operational processes, automation, and best practices",
        
        "Cross-Functional Collaboration:",
        "• Collaborate closely with Development, DevOps, Security teams, and external vendors",
        "• Lead infrastructure upgrades, migrations, and continuous improvement initiatives",
        "• Ensure SLA compliance and define service performance metrics",
        "• Act as technical escalation point for production issues"
    ]
    
    for point in leadership_responsibilities:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('12.2 Required Skills & Experience', 2)
    
    skills_table = doc.add_table(rows=9, cols=2)
    skills_table.style = 'Medium Grid 3 Accent 1'
    
    skills_data = [
        ('Skill Area', 'Requirements'),
        ('Team Management', 'Proven experience managing hosting or infrastructure teams'),
        ('AWS Expertise', 'Strong hands-on experience with AWS in complex environments'),
        ('Operating Systems', 'Solid knowledge of Linux and Windows operating systems'),
        ('Networking', 'Strong understanding of DNS, load balancing, and firewalls'),
        ('Virtualization', 'Experience with virtualization and storage technologies'),
        ('Automation', 'Experience with automation tools and scripting (Python, Bash, PowerShell)'),
        ('Communication', 'Excellent English communication skills (written and verbal)'),
        ('Leadership', 'Strong leadership and problem-solving skills in global environments')
    ]
    
    for i, (skill, req) in enumerate(skills_data):
        skills_table.rows[i].cells[0].text = skill
        skills_table.rows[i].cells[1].text = req
        if i == 0:
            skills_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            skills_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('12.3 Team Structure Vision', 2)
    
    team_structure = [
        "Recommended Team Organization:",
        
        "Tier 1 - Operations Engineers (3-4 members):",
        "• Monitor dashboards and respond to alerts",
        "• Handle routine maintenance and patching",
        "• Execute runbooks for common issues",
        "• Escalate complex issues to Tier 2",
        "• Coverage: 24/7 follow-the-sun model",
        
        "Tier 2 - Senior Infrastructure Engineers (2-3 members):",
        "• Deep troubleshooting and root cause analysis",
        "• Infrastructure design and implementation",
        "• Automation development and optimization",
        "• Mentor Tier 1 engineers",
        "• On-call rotation for critical escalations",
        
        "Tier 3 - Principal/Architect (1 member):",
        "• Architecture design and strategic planning",
        "• Complex migrations and infrastructure upgrades",
        "• Vendor management and technology evaluation",
        "• Disaster recovery planning and testing",
        "• Technical escalation point for organization",
        
        "DevOps/Automation Specialist (1-2 members):",
        "• CI/CD pipeline management",
        "• Infrastructure as Code (Terraform, CloudFormation)",
        "• Automation framework development",
        "• Integration with development workflows"
    ]
    
    for point in team_structure:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        elif point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('12.4 Operational Model', 2)
    
    operational_model = [
        "Follow-the-Sun Coverage:",
        "• Americas team: 8am-5pm EST (New York)",
        "• EMEA team: 8am-5pm GMT (London)",
        "• APAC team: 8am-5pm JST (Tokyo)",
        "• Ensures 24/7 coverage with minimal on-call burden",
        
        "Incident Management:",
        "• P1 (Critical): Immediate response, all hands on deck",
        "• P2 (High): Response within 30 minutes during business hours",
        "• P3 (Medium): Response within 4 hours",
        "• P4 (Low): Scheduled maintenance window",
        
        "Change Management:",
        "• All production changes require CAB approval",
        "• Automated deployments during maintenance windows",
        "• Rollback procedures tested and documented",
        "• Post-implementation review for all major changes",
        
        "Knowledge Management:",
        "• Centralized runbook repository (Confluence/SharePoint)",
        "• Post-mortem documentation for all P1/P2 incidents",
        "• Weekly knowledge sharing sessions",
        "• Quarterly architecture review meetings"
    ]
    
    for point in operational_model:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        elif point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        else:
            doc.add_paragraph(point)
    
    doc.add_paragraph()
    
    doc.add_heading('12.5 Success Metrics', 2)
    
    success_metrics_table = doc.add_table(rows=7, cols=3)
    success_metrics_table.style = 'Light List Accent 1'
    
    success_metrics_data = [
        ('Metric', 'Target', 'Measurement'),
        ('System Availability', '99.9%+', 'Uptime monitoring across all environments'),
        ('Mean Time To Repair (MTTR)', '<15 minutes', 'Average time from alert to resolution'),
        ('Incident Response Time', 'P1: <5 min, P2: <30 min', 'Time from alert to engineer engagement'),
        ('Automation Coverage', '80%+', 'Percentage of tasks automated vs manual'),
        ('Team Satisfaction', '4.0+/5.0', 'Quarterly team engagement survey'),
        ('Customer SLA Compliance', '100%', 'Percentage of SLAs met per customer contract')
    ]
    
    for i, (metric, target, measurement) in enumerate(success_metrics_data):
        success_metrics_table.rows[i].cells[0].text = metric
        success_metrics_table.rows[i].cells[1].text = target
        success_metrics_table.rows[i].cells[2].text = measurement
        if i == 0:
            success_metrics_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            success_metrics_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
            success_metrics_table.rows[i].cells[2].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('12.6 Growth & Development Path', 2)
    
    growth_path = [
        "Career Progression:",
        "• Operations Engineer → Senior Engineer → Principal Engineer → Architect",
        "• Specialization tracks: Networking, Security, Automation, Cloud Architecture",
        "• Leadership track: Team Lead → Manager → Director",
        
        "Training & Certification:",
        "• AWS certifications: Solutions Architect, SysOps Administrator, DevOps Engineer",
        "• Vendor training: Cloudflare, Terraform, Kubernetes",
        "• Soft skills: Leadership, communication, incident management",
        "• Annual training budget per team member",
        
        "Innovation Time:",
        "• 10% time for automation projects and process improvements",
        "• Quarterly hackathons for infrastructure innovation",
        "• Recognition program for operational excellence contributions"
    ]
    
    for point in growth_path:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        elif point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        else:
            doc.add_paragraph(point)
    
    doc.add_page_break()
    
    # ========================================
    # POC/DEMO SYSTEM ARCHITECTURE
    # ========================================
    doc.add_heading('13. POC/Demo System Architecture', 1)
    
    doc.add_paragraph(
        "This section describes a proof-of-concept (POC) system designed for demonstration and testing purposes, "
        "using AWS free-tier resources with a simplified architecture."
    )
    
    doc.add_heading('13.1 System Overview', 2)
    
    poc_overview = [
        "The POC system consists of two Windows-based EC2 instances running in AWS free tier:",
        
        "1. Frontend Instance: Web server hosting the user interface",
        "2. Backend Instance: Application server with integrated database",
        
        "This architecture demonstrates the core Made4Net WMS functionality in a cost-effective manner "
        "suitable for demos, training, and proof-of-concept deployments."
    ]
    
    for para in poc_overview:
        if para.startswith(('1.', '2.')):
            doc.add_paragraph(para, style='List Number')
        else:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    doc.add_heading('13.2 Architecture Components', 2)
    
    components_table = doc.add_table(rows=3, cols=3)
    components_table.style = 'Medium Grid 3 Accent 1'
    
    components_data = [
        ('Component', 'Specification', 'Purpose'),
        ('Frontend EC2', 'Windows Server 2022, t2.micro (1 vCPU, 1GB RAM), IIS Web Server',
         'Hosts web UI for inventory management, handles user authentication'),
        ('Backend EC2', 'Windows Server 2022, t2.micro (1 vCPU, 1GB RAM), SQL Server Express',
         'Application logic, REST API, database (SQL Server Express 10GB limit)')
    ]
    
    for i, (component, spec, purpose) in enumerate(components_data):
        components_table.rows[i].cells[0].text = component
        components_table.rows[i].cells[1].text = spec
        components_table.rows[i].cells[2].text = purpose
        if i == 0:
            components_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            components_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
            components_table.rows[i].cells[2].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('13.3 Frontend Instance', 2)
    
    frontend_details = [
        "Web Server Configuration:",
        "• Operating System: Windows Server 2022 (Free Tier eligible)",
        "• Web Server: IIS (Internet Information Services) 10.0",
        "• Instance Type: t2.micro (1 vCPU, 1GB RAM)",
        "• Storage: 30GB EBS (General Purpose SSD)",
        "• Branding: Made4Net logo, color scheme, and branding throughout UI",
        
        "Application Features (Made4Net WMS):",
        "• Inventory Management UI: View, add, edit, delete inventory items",
        "• Internal Authentication: Application-level username/password for warehouse users",
        "• User Roles: Admin, Manager, Operator (stored in InventoryDB.Users table)",
        "• Responsive Design: HTML5, CSS3, JavaScript (Bootstrap framework)",
        "• API Integration: RESTful API calls to backend instance",
        "• Made4Net Branding: Custom logo, color scheme (#0066CC primary, #FF6600 accent)",
        
        "Internal Authentication (Application Users):",
        "• Purpose: Warehouse staff accessing Made4Net WMS application",
        "• Method: Username/password stored in SQL Server (InventoryDB.Users table)",
        "• Password Policy: Minimum 8 characters, complexity requirements",
        "• Session Management: 30-minute timeout, secure cookies",
        "• User Roles: Admin (full access), Manager (read/write), Operator (read-only)",
        "• Authentication Flow: Login page → Backend API validation → JWT token → Session",
        
        "Security:",
        "• HTTPS enabled with self-signed certificate (demo) or Let's Encrypt (production)",
        "• Security Group: Allow inbound HTTPS (443) from internet, RDP (3389) from admin IP only",
        "• Windows Firewall: Enabled with IIS exceptions",
        "• Application-level authentication separate from AWS/Windows authentication"
    ]
    
    for point in frontend_details:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('13.4 Backend Instance', 2)
    
    backend_details = [
        "Application Server Configuration:",
        "• Operating System: Windows Server 2022 (Free Tier eligible)",
        "• Runtime: .NET Framework 4.8 or .NET 6",
        "• Instance Type: t2.micro (1 vCPU, 1GB RAM)",
        "• Storage: 30GB EBS (General Purpose SSD)",
        
        "Database:",
        "• SQL Server Express 2019 (Free, 10GB database size limit)",
        "• Database Name: InventoryDB",
        "• Tables: Users, Inventory, AuditLog",
        "• Backup: Daily automated backup to S3 (free tier: 5GB)",
        
        "API Endpoints:",
        "• POST /api/auth/login - User authentication",
        "• GET /api/inventory - List all inventory items",
        "• POST /api/inventory - Add new inventory item",
        "• PUT /api/inventory/{id} - Update inventory item",
        "• DELETE /api/inventory/{id} - Delete inventory item",
        "• GET /api/health - Health check endpoint",
        
        "Security:",
        "• Security Group: Allow inbound HTTPS (443) from frontend instance only",
        "• SQL Server: Windows Authentication mode",
        "• API: JWT token-based authentication",
        "• Encryption: TLS 1.2+ for all API communication"
    ]
    
    for point in backend_details:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('13.5 Authentication & Access Control', 2)
    
    auth_details = [
        "This POC system implements two distinct authentication layers:",
        
        "Layer 1: Internal Application Authentication (Made4Net Users)",
        "• Purpose: Warehouse staff accessing the Made4Net WMS application",
        "• User Types: Warehouse Managers, Operators, Supervisors",
        "• Authentication Method: Username/password stored in SQL Server",
        "• Access Level: Application features (inventory management, reporting)",
        "• Login URL: https://[elastic-ip]/login",
        "• User Management: Admin users can create/modify application users",
        "• Password Storage: Hashed with bcrypt/PBKDF2 in InventoryDB.Users table",
        
        "Layer 2: External System Administration (AWS/Infrastructure)",
        "• Purpose: Hosting team managing AWS infrastructure and EC2 instances",
        "• User Types: System Administrators, DevOps Engineers, Hosting Team",
        "• Authentication Method: AWS IAM credentials + MFA for AWS Console",
        "• Windows Authentication: RDP access to EC2 instances with Windows credentials",
        "• Access Level: AWS resources (EC2, VPC, S3, CloudWatch, Security Groups)",
        "• Management Tools: AWS Console, Systems Manager, RDP, PowerShell",
        "• Separation: System admins do NOT have Made4Net application user accounts",
        
        "Access Control Matrix:"
    ]
    
    for point in auth_details:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        elif point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        else:
            p = doc.add_paragraph(point)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    access_matrix_table = doc.add_table(rows=5, cols=4)
    access_matrix_table.style = 'Medium Grid 3 Accent 1'
    
    access_matrix_data = [
        ('User Type', 'Authentication', 'Access Scope', 'Tools/Interface'),
        ('Warehouse Manager', 'Made4Net App Login', 'Inventory Management UI', 'Web Browser (HTTPS)'),
        ('Warehouse Operator', 'Made4Net App Login', 'Read-only Inventory', 'Web Browser (HTTPS)'),
        ('System Administrator', 'AWS IAM + MFA', 'EC2, VPC, S3, CloudWatch', 'AWS Console, RDP'),
        ('Hosting Engineer', 'AWS IAM + Windows Auth', 'EC2 Instances, IIS, SQL Server', 'RDP, PowerShell, SSMS')
    ]
    
    for i, (user_type, auth, scope, tools) in enumerate(access_matrix_data):
        access_matrix_table.rows[i].cells[0].text = user_type
        access_matrix_table.rows[i].cells[1].text = auth
        access_matrix_table.rows[i].cells[2].text = scope
        access_matrix_table.rows[i].cells[3].text = tools
        if i == 0:
            access_matrix_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            access_matrix_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
            access_matrix_table.rows[i].cells[2].paragraphs[0].runs[0].font.bold = True
            access_matrix_table.rows[i].cells[3].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('13.6 Failover & High Availability', 2)
    
    failover_details = [
        "The POC system implements basic failover capabilities to demonstrate resilience:",
        
        "Health Monitoring:",
        "• Frontend health check: HTTP GET to /health endpoint every 60 seconds",
        "• Backend health check: SQL Server connection test every 60 seconds",
        "• CloudWatch alarms: Alert on instance status check failures",
        
        "Failover Strategy:",
        "• Automated Instance Recovery: AWS automatically recovers failed instances (same AZ)",
        "• Manual Failover: Documented procedure to launch replacement instance from AMI",
        "• Database Backup: Daily automated backups to S3 for disaster recovery",
        "• AMI Snapshots: Weekly automated AMI creation for both instances",
        
        "Recovery Time Objectives:",
        "• Automated Recovery: 5-10 minutes (AWS instance recovery)",
        "• Manual Failover: 15-30 minutes (launch from AMI + restore database)",
        "• Data Loss: Maximum 24 hours (daily backup schedule)",
        
        "Limitations (Free Tier):",
        "• No Multi-AZ deployment (requires additional instances)",
        "• No load balancer (requires ALB, not free tier)",
        "• No auto-scaling (requires multiple instances)",
        "• Single point of failure for each tier"
    ]
    
    for point in failover_details:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        elif point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        else:
            p = doc.add_paragraph(point)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_paragraph()
    
    doc.add_heading('13.7 Network Architecture', 2)
    
    network_details = [
        "VPC Configuration:",
        "• VPC: 10.0.0.0/16 (65,536 IP addresses)",
        "• Public Subnet: 10.0.1.0/24 (Frontend instance)",
        "• Private Subnet: 10.0.2.0/24 (Backend instance)",
        "• Internet Gateway: Attached to VPC for public internet access",
        "• NAT Gateway: Not used (cost optimization - backend uses direct internet via IGW for updates)",
        
        "Security Groups:",
        "• Frontend-SG: Inbound HTTPS (443) from 0.0.0.0/0, RDP (3389) from admin IP",
        "• Backend-SG: Inbound HTTPS (443) from Frontend-SG only, RDP (3389) from admin IP",
        "• Outbound: All traffic allowed (for Windows updates and package downloads)",
        
        "DNS & Routing:",
        "• Frontend: Elastic IP attached for static public IP",
        "• Backend: Private IP only (10.0.2.x)",
        "• Route 53: Optional custom domain (e.g., demo.made4net.com)",
        "• Internal DNS: Backend accessible via private IP from frontend"
    ]
    
    for point in network_details:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        else:
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('13.8 Cost Analysis', 2)
    
    cost_table = doc.add_table(rows=7, cols=3)
    cost_table.style = 'Light List Accent 1'
    
    cost_data = [
        ('Resource', 'Free Tier Allowance', 'Monthly Cost (After Free Tier)'),
        ('Frontend EC2 (t2.micro)', '750 hours/month (1 instance)', '$0 (within free tier)'),
        ('Backend EC2 (t2.micro)', '750 hours/month (1 instance)', '$0 (within free tier)'),
        ('EBS Storage (30GB x 2)', '30GB', '$0 (within free tier)'),
        ('Data Transfer Out', '15GB/month', '$0.09/GB after 15GB'),
        ('Elastic IP', 'Free when attached', '$0'),
        ('Total Monthly Cost', 'Free Tier (12 months)', '~$0-5 (minimal overage)')
    ]
    
    for i, (resource, free_tier, cost) in enumerate(cost_data):
        cost_table.rows[i].cells[0].text = resource
        cost_table.rows[i].cells[1].text = free_tier
        cost_table.rows[i].cells[2].text = cost
        if i == 0:
            cost_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            cost_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
            cost_table.rows[i].cells[2].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('13.8 Deployment Process', 2)
    
    deployment_steps = [
        "Step 1: VPC Setup",
        "1. Create VPC (10.0.0.0/16)",
        "2. Create public subnet (10.0.1.0/24)",
        "3. Create private subnet (10.0.2.0/24)",
        "4. Create and attach Internet Gateway",
        "5. Configure route tables",
        
        "Step 2: Security Groups",
        "1. Create Frontend-SG with HTTPS and RDP rules",
        "2. Create Backend-SG with restricted access",
        
        "Step 3: Launch Instances",
        "1. Launch Frontend EC2 (Windows Server 2022, t2.micro, public subnet)",
        "2. Allocate and attach Elastic IP to frontend",
        "3. Launch Backend EC2 (Windows Server 2022, t2.micro, private subnet)",
        
        "Step 4: Configure Frontend",
        "1. RDP to frontend instance",
        "2. Install IIS via Server Manager",
        "3. Deploy web application files",
        "4. Configure IIS bindings and SSL",
        
        "Step 5: Configure Backend",
        "1. RDP to backend instance (via frontend as jump box)",
        "2. Install SQL Server Express 2019",
        "3. Deploy application code and API",
        "4. Create database schema and seed data",
        
        "Step 6: Testing & Validation",
        "1. Test frontend UI access via browser",
        "2. Test user authentication",
        "3. Test inventory CRUD operations",
        "4. Verify backend API connectivity",
        "5. Test failover procedures"
    ]
    
    for point in deployment_steps:
        if point.startswith(('1.', '2.', '3.', '4.', '5.')):
            doc.add_paragraph(point, style='List Number')
        else:
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('13.10 Use Cases', 2)
    
    use_cases = [
        "This POC system is suitable for:",
        
        "• Sales Demonstrations: Show Made4Net WMS capabilities to prospects",
        "• Training Environment: Hands-on training for new users and administrators",
        "• Development Testing: Test new features before production deployment",
        "• Proof of Concept: Validate architecture decisions with minimal cost",
        "• Customer Trials: Provide temporary access for evaluation purposes",
        
        "Limitations:",
        "• Not suitable for production workloads (single instance, no HA)",
        "• Limited to 10GB database size (SQL Server Express)",
        "• Performance constraints (1 vCPU, 1GB RAM per instance)",
        "• No enterprise features (clustering, replication, advanced monitoring)"
    ]
    
    for point in use_cases:
        if point.startswith('•'):
            doc.add_paragraph(point, style='List Bullet')
        elif point.endswith(':'):
            p = doc.add_paragraph(point)
            p.runs[0].font.bold = True
        else:
            p = doc.add_paragraph(point)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()
    
    # ========================================
    # CONCLUSION
    # ========================================
    doc.add_heading('14. Conclusion', 1)
    
    conclusion = [
        "This operational excellence architecture transforms how the hosting team manages "
        "800+ warehouse endpoints and hosted infrastructure, including hybrid deployments "
        "with AWS Outposts. By leveraging AWS-native tools, Cloudflare optimization, and automation, the team achieves:",
        
        "Key Outcomes:",
        "• Secure Connectivity: Cloudflare proxy for optimization and firewall protection",
        "• Unified Visibility: Single pane of glass for all endpoints via Fleet Manager",
        "• Hybrid Flexibility: AWS Outposts for low-latency and data residency requirements",
        "• Rapid Troubleshooting: 8-15 minute MTTR using systematic workflows",
        "• Automated Operations: Zero-downtime deployments and patching",
        "• Audit Compliance: Complete session logging and change tracking",
        "• Team Excellence: Clear organizational structure with defined roles and growth paths",
        
        "The architecture eliminates manual, error-prone operations and provides the hosting "
        "team with enterprise-grade tools to manage infrastructure at scale—whether in the "
        "cloud or on-premises. Combined with a well-structured Global Hosting Team, this directly "
        "addresses the core responsibilities of the Global Hosting Team Manager role: maintaining "
        "secure, performant, and compliant infrastructure for Made4Net's global customer base.",
        
        "By implementing these best practices and building a strong team foundation, the hosting "
        "organization can focus on strategic initiatives rather than firefighting, while maintaining "
        "99.9%+ uptime for mission-critical warehouse operations across cloud and on-premises deployments."
    ]
    
    for para in conclusion:
        if para.startswith('•'):
            doc.add_paragraph(para, style='List Bullet')
        elif para == 'Key Outcomes:':
            p = doc.add_paragraph(para)
            p.runs[0].font.bold = True
        else:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Save document
    doc.save('Made4Net-Operational-Excellence-HLD.docx')
    print("✓ Operational Excellence HLD generated: Made4Net-Operational-Excellence-HLD.docx")
    print(f"  Document size: {len(doc.element.xml)} bytes")
    print("  Sections: 14 focused sections on operational excellence")
    print("  Focus: Connectivity, Remote Access, Monitoring, Troubleshooting, Deployment, Outposts, Multi-Account, Team Structure, POC System")

if __name__ == '__main__':
    create_ops_hld()
