#!/usr/bin/env python3
"""
Convert docs/AI-Chat-Logic.md into a Microsoft Word document (.docx) in the root folder.

Uses python-docx (same library as the other _gen_*.py document generators).
Supports: headings (#/##/###), bullet lists, numbered + checkbox lists, GitHub
pipe tables, fenced code blocks, blockquotes, horizontal rules, and inline
**bold** / `code` formatting.

Run:  python _gen_ai_chat_logic_doc.py
"""

import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE_MD = os.path.join(SCRIPT_DIR, 'docs', 'AI-Chat-Logic.md')
OUTPUT_DOCX = os.path.join(SCRIPT_DIR, 'AI-Chat-Logic.docx')

ACCENT = RGBColor(0xE8, 0x71, 0x4A)   # SlashMyBill orange
DARK = RGBColor(0x1F, 0x29, 0x37)
CODE_BG = 'F3F4F6'


def _shade(cell_or_para_element, fill):
    """Apply a background shade to a table cell or paragraph."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    cell_or_para_element.append(shd)


def _add_inline(paragraph, text):
    """Render inline **bold** and `code` segments into runs."""
    # Split on bold and inline-code tokens, keeping the delimiters
    tokens = re.split(r'(\*\*.+?\*\*|`[^`]+`)', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xB1, 0x1A, 0x6B)
        else:
            paragraph.add_run(tok)


def _add_code_block(doc, lines):
    """Render a fenced code block as a single shaded, monospaced paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    _shade(p._p.get_or_add_pPr(), CODE_BG)
    run = p.add_run('\n'.join(lines))
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = DARK


def _add_table(doc, rows):
    """rows: list of lists of cell strings; first row is the header."""
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    try:
        table.style = 'Light Grid Accent 1'
    except KeyError:
        table.style = 'Table Grid'

    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for c in range(cols):
            text = row[c] if c < len(row) else ''
            cell = cells[c]
            cell.text = ''
            para = cell.paragraphs[0]
            _add_inline(para, text)
            if i == 0:
                _shade(cell._tc.get_or_add_tcPr(), 'E0E7FF')
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()


def _is_table_sep(line):
    return bool(re.match(r'^\s*\|?[\s:|-]+\|?\s*$', line)) and '-' in line


def _split_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def build():
    if not os.path.exists(SOURCE_MD):
        print(f"Source markdown not found: {SOURCE_MD}")
        sys.exit(1)

    with open(SOURCE_MD, 'r', encoding='utf-8') as f:
        md_lines = f.read().split('\n')

    doc = Document()
    # Base font
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    i = 0
    n = len(md_lines)
    while i < n:
        line = md_lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < n and not md_lines[i].strip().startswith('```'):
                code_lines.append(md_lines[i])
                i += 1
            _add_code_block(doc, code_lines)
            i += 1
            continue

        # Table (header line followed by separator)
        if '|' in line and i + 1 < n and _is_table_sep(md_lines[i + 1]):
            rows = [_split_table_row(line)]
            i += 2  # skip header + separator
            while i < n and '|' in md_lines[i] and md_lines[i].strip():
                rows.append(_split_table_row(md_lines[i]))
                i += 1
            _add_table(doc, rows)
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', stripped):
            i += 1
            continue

        # Headings
        if stripped.startswith('### '):
            h = doc.add_heading(level=3)
            _add_inline(h, stripped[4:])
            i += 1
            continue
        if stripped.startswith('## '):
            h = doc.add_heading(level=2)
            _add_inline(h, stripped[3:])
            i += 1
            continue
        if stripped.startswith('# '):
            h = doc.add_heading(level=1)
            _add_inline(h, stripped[2:])
            for run in h.runs:
                run.font.color.rgb = ACCENT
            i += 1
            continue

        # Blockquote
        if stripped.startswith('> '):
            p = doc.add_paragraph(style='Intense Quote' if 'Intense Quote' in [s.name for s in doc.styles] else 'Normal')
            _add_inline(p, stripped[2:])
            i += 1
            continue

        # Checkbox list item
        m_chk = re.match(r'^[-*]\s+\[[ xX]\]\s+(.*)$', stripped)
        if m_chk:
            p = doc.add_paragraph(style='List Bullet')
            _add_inline(p, '\u2610 ' + m_chk.group(1))
            i += 1
            continue

        # Bullet list item
        if re.match(r'^[-*]\s+', stripped):
            p = doc.add_paragraph(style='List Bullet')
            _add_inline(p, re.sub(r'^[-*]\s+', '', stripped))
            i += 1
            continue

        # Numbered list item
        if re.match(r'^\d+\.\s+', stripped):
            p = doc.add_paragraph(style='List Number')
            _add_inline(p, re.sub(r'^\d+\.\s+', '', stripped))
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _add_inline(p, stripped)
        i += 1

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")


if __name__ == '__main__':
    build()
