from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helper Functions (same pattern as _gen_ug.py) ─────────────────────────────

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
        if ri % 2 == 0:
            for ci in range(len(row)):
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E8EDF5')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'), 'clear')
                tr.cells[ci]._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def tip_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('\U0001f4a1 Tip: ' + text)
    run.font.color.rgb = RGBColor(22, 101, 52)
    run.font.italic = True

def note_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('\u26a0 Note: ' + text)
    run.font.color.rgb = RGBColor(146, 64, 14)
    run.font.italic = True

# ── Document Setup ────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SlashMyBill')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('High-Level Design Document')
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(99, 102, 241)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Version 13  |  {datetime.date.today().strftime("%B %Y")}  |  AWS FinOps Platform').font.size = Pt(10)

doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# 1. Executive Summary
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Executive Summary', 1)
doc.add_paragraph(
    'SlashMyBill is an AI-powered AWS FinOps platform that enables organizations to analyze, '
    'optimize, and reduce cloud spending across multiple AWS accounts. Members connect their '
    'AWS accounts via a secure cross-account IAM role, then use a conversational AI agent to '
    'ask natural language questions about costs, detect anomalies, receive actionable savings '
    'recommendations, and execute automated cleanup actions — all backed by live AWS API data.\n\n'
    'The platform combines three analysis engines:\n'
    '  1. Waste Scan Engine — a tips-driven scanner that evaluates 92 optimization rules against '
    'live AWS resource state to identify idle, unattached, and misconfigured resources.\n'
    '  2. AI Chat Engine — a data-gathering pipeline that collects cost, resource, and metric data '
    'from connected accounts and feeds it to Amazon Bedrock (Nova Lite) for natural language analysis.\n'
    '  3. Bedrock Agent — an autonomous agent with 12 action tools that decides which data to fetch '
    'based on the user\'s question and iterates until it has enough context to answer.\n\n'
    'Platform URL: https://www.slashmycloudbill.com/members/\n'
    'AWS Platform Account: 991105135552 (us-east-1)\n'
    'AI Engine: Amazon Bedrock Nova Lite v1 + Bedrock Agent (IDG5VJGUOZ5W)'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. System Architecture
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '2. System Architecture', 1)

add_heading(doc, '2.1 Platform Topology', 2)
doc.add_paragraph(
    'All platform infrastructure runs in a single AWS account (991105135552, us-east-1). '
    'Customer AWS accounts are accessed remotely via cross-account IAM roles — no agents or '
    'software is deployed into customer environments.\n\n'
    'Architecture layers:\n'
    '  Frontend:  Static site (HTML/CSS/JS + ECharts v5) hosted on S3, served via CloudFront\n'
    '             Domain: slashmycloudbill.com with Route 53 DNS\n\n'
    '  API:       API Gateway HTTP API v2 (25+ routes) → Lambda function handlers\n\n'
    '  Compute:   8 Lambda functions (Python 3.12) handling member operations, AI queries,\n'
    '             scan engine, agent actions, bill analysis, scheduling, admin, and OTP\n\n'
    '  AI:        Amazon Bedrock Nova Lite model for chat analysis\n'
    '             Bedrock Agent (IDG5VJGUOZ5W) with 12 action tools for autonomous queries\n\n'
    '  Storage:   DynamoDB (9 tables, single-table design for member data)\n'
    '             S3 for static assets and bill uploads\n\n'
    '  Auth:      Cognito User Pool for member authentication (JWT tokens)\n\n'
    '  Scheduler: Amazon EventBridge Scheduler for automated stop/start actions'
)

add_heading(doc, '2.2 Core Components', 2)
add_table(doc,
    ['Component', 'Technology', 'Purpose'],
    [
        ['Frontend', 'HTML/CSS/JS + ECharts v5', 'Member Portal — Observe, Chat, Act, Plan, Configure tabs'],
        ['API Gateway', 'HTTP API v2 (25+ routes)', 'Request routing to Lambda handlers'],
        ['Member Handler', 'Python 3.12 Lambda (256 MB, 120s)', 'Auth, accounts, AI queries, scan engine, data gathering'],
        ['Agent Action', 'Python 3.12 Lambda (256 MB, 120s)', 'Bedrock Agent action group — cross-account data fetching'],
        ['Bill Analyzer', 'Python 3.12 Lambda (1024 MB, 900s)', 'PDF bill parsing + Bedrock analysis'],
        ['Scheduler Executor', 'Python 3.12 Lambda (512 MB, 300s)', 'EventBridge-triggered cross-account stop/start/scale'],
        ['AI Engine', 'Amazon Bedrock Nova Lite v1', 'Natural language analysis + recommendations'],
        ['Bedrock Agent', 'Agent IDG5VJGUOZ5W (12 tools)', 'Autonomous multi-step data retrieval + analysis'],
        ['Storage', 'DynamoDB (9 tables) + S3', 'Members, accounts, tips, OTP, feedback, metrics, scans, invoices'],
        ['CDN', 'CloudFront + Route 53', 'HTTPS delivery, DNS, caching'],
        ['CI/CD', 'GitHub Actions + OIDC', 'Automated deployment on push to main'],
    ],
    [2.0, 2.5, 3.5]
)

add_heading(doc, '2.3 Cross-Account Access Model', 2)
doc.add_paragraph(
    'Customer accounts are accessed via STS AssumeRole with a SHA-256 ExternalId for security.\n\n'
    'Role setup:\n'
    '  • Members deploy a CloudFormation template in their AWS account\n'
    '  • Template creates IAM Role: SlashMyBill-{AccountID}\n'
    '  • Trust policy: Platform account (991105135552) with ExternalId = SHA-256(member_email)\n'
    '  • Permissions: ReadOnlyAccess managed policy + inline policy for Cost Explorer, Budgets,\n'
    '    Pricing, Trusted Advisor, S3 write (lifecycle/delete), EC2 write (stop/snapshot delete),\n'
    '    RDS write (delete with snapshot), and stack self-management\n\n'
    'Access flow:\n'
    '  1. Member makes API request (e.g., scan, AI query, dashboard load)\n'
    '  2. Backend retrieves member email and account list from DynamoDB\n'
    '  3. For each account, calls sts:AssumeRole with:\n'
    '       RoleArn: arn:aws:iam::{AccountID}:role/SlashMyBill-{AccountID}\n'
    '       ExternalId: SHA-256(member_email)\n'
    '  4. Receives temporary credentials (1-hour expiry)\n'
    '  5. Uses credentials to call AWS APIs in the customer account\n\n'
    'The ExternalId prevents confused deputy attacks — even if an attacker knows the role ARN, '
    'they cannot assume it without the correct SHA-256 hash of the member\'s email.'
)
note_box(doc, 'No long-lived credentials are stored. Every cross-account operation uses fresh STS temporary tokens.')


# ══════════════════════════════════════════════════════════════════════════════
# 3. Analysis Engine — Decision Flow (MAIN SECTION)
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '3. Analysis Engine — Decision Flow', 1)
doc.add_paragraph(
    'This section describes the core analysis logic that powers SlashMyBill. It covers how the '
    'system scans for waste, answers natural language questions, runs autonomous agent queries, '
    'audits FinOps settings, assembles dashboard data, and analyzes tag coverage. Each subsection '
    'traces the complete decision flow from trigger to output.'
)

# ── 3.1 Waste Scan Engine ────────────────────────────────────────────────────

add_heading(doc, '3.1 Waste Scan Engine (Act \u2192 Waste Cleanup)', 2)
doc.add_paragraph(
    'The waste scan engine is a tips-driven scanner that evaluates 92 optimization rules from '
    'the DynamoDB knowledge base against live AWS resource state. It runs on demand when a member '
    'clicks "Scan for Waste" in the Act tab or "Scan for Savings Opportunities" in the Chat tab.'
)

add_heading(doc, '3.1.1 Scan Execution Flow', 3)
doc.add_paragraph(
    'Step 1 — Load Tips\n'
    '  Load all 92 optimization tips from the ViewMyBill-CostOptimizationTips DynamoDB table.\n'
    '  Each tip contains: tipId, service, category, automatedCheck specification, checkImplemented\n'
    '  flag, actionType (delete/modify/advisory/deep-link), and estimatedSavings guidance.\n\n'
    'Step 2 — Assume Cross-Account Roles\n'
    '  For each connected account (up to 5), call sts:AssumeRole with the member\'s SHA-256\n'
    '  ExternalId. Obtain temporary credentials scoped to that account.\n\n'
    'Step 3 — Gather Resource Data\n'
    '  Using the temporary credentials, collect data from AWS APIs across multiple services:\n'
    '    \u2022 EC2: DescribeInstances, DescribeVolumes, DescribeAddresses, DescribeSnapshots\n'
    '    \u2022 RDS: DescribeDBInstances\n'
    '    \u2022 EBS: DescribeVolumes (state, type, size, attachments)\n'
    '    \u2022 S3: ListBuckets, GetBucketLifecycleConfiguration, ListObjectsV2\n'
    '    \u2022 ELB: DescribeLoadBalancers, DescribeTargetHealth\n'
    '    \u2022 CloudWatch: GetMetricStatistics (CPUUtilization, DatabaseConnections, Invocations)\n'
    '    \u2022 VPC: DescribeNatGateways, DescribeVpcEndpoints\n'
    '    \u2022 KMS: ListKeys, DescribeKey\n'
    '    \u2022 Cost Explorer: GetCostAndUsage (to determine active services)\n\n'
    'Step 4 — Evaluate Each Tip\n'
    '  For each of the 92 tips, the engine runs the corresponding automated check function\n'
    '  from the _SCAN_REGISTRY. Each check function receives the gathered data and returns:\n'
    '    \u2022 status: "found" (waste detected) or "clean" (no issue)\n'
    '    \u2022 savings: Estimated monthly dollar savings\n'
    '    \u2022 resources: List of affected resource IDs with account labels\n'
    '    \u2022 cardData: Structured data for rendering the Act tab card\n\n'
    'Step 5 — Deduplicate Findings\n'
    '  When scanning multiple accounts, findings are deduplicated by tipId. Resources from\n'
    '  different accounts are merged into a single finding card with per-resource account labels.\n\n'
    'Step 6 — Sort and Return\n'
    '  Findings are sorted by savings amount (highest first). The top findings are returned\n'
    '  to the frontend and cached in MemberPortal-Members.lastScan for the Chat widget.\n\n'
    'Step 7 — Calculate Efficiency Score\n'
    '  Efficiency Score = (total_spend - potential_savings) / total_spend \u00d7 100\n'
    '  This score is displayed on the dashboard KPI bar and in AI responses.'
)

add_heading(doc, '3.1.2 Automated Check Functions — Detail', 3)
doc.add_paragraph(
    'Each check function maps to a specific AWS API call pattern and evaluation logic:'
)

add_table(doc,
    ['Check', 'API Call', 'Detection Logic', 'Savings Calculation'],
    [
        ['EBS Unattached Volumes',
         'ec2:DescribeVolumes',
         'Filter state=available (not attached to any instance)',
         '$0.10/GB/month \u00d7 volume size'],
        ['Stale EBS Snapshots',
         'ec2:DescribeSnapshots',
         'Filter snapshots older than 180 days (owner=self)',
         'Storage cost based on snapshot size'],
        ['Unassociated Elastic IPs',
         'ec2:DescribeAddresses',
         'Filter addresses with no AssociationId',
         '$3.65/month per unassociated EIP'],
        ['S3 Missing Lifecycle',
         's3:GetBucketLifecycleConfiguration',
         'API returns NoSuchLifecycleConfiguration error = missing',
         'Advisory — savings depend on data tiering potential'],
        ['Idle Load Balancers',
         'elbv2:DescribeTargetHealth',
         '0 healthy targets across all target groups',
         '$16.43/month per idle ALB/NLB'],
        ['Idle EC2 Instances',
         'CloudWatch CPUUtilization',
         'Average CPU < 5% over 14-day period',
         'Full instance hourly rate \u00d7 730 hours/month'],
        ['Idle RDS Instances',
         'CloudWatch CPU + Connections',
         'Average CPU < 5% AND DatabaseConnections < 2 over 14 days',
         'Full RDS instance hourly rate \u00d7 730 hours/month'],
        ['Customer-Managed KMS Keys',
         'kms:ListKeys + kms:DescribeKey',
         'Filter customer-managed keys (not AWS-managed)',
         '$1.00/month per customer-managed key'],
        ['Spot Pricing Candidates',
         'EC2 instances + CloudWatch',
         'Non-production instances with low CPU utilization',
         'Up to 90% savings vs On-Demand pricing'],
        ['gp2 \u2192 gp3 Migration',
         'ec2:DescribeVolumes',
         'EBS volumes with type=gp2',
         '~20% savings by migrating to gp3'],
        ['Graviton Candidates',
         'ec2:DescribeInstances',
         'x86 instance families (m5, c5, r5, t3) \u2192 ARM equivalents (m7g, c7g, r7g, t4g)',
         '~20% savings from Graviton migration'],
    ],
    [1.8, 1.8, 2.2, 2.2]
)

tip_box(doc, 'The scan engine gates checks by active services — if Cost Explorer shows $0 spend on RDS, '
    'all RDS-related checks are skipped to reduce API calls and scan time.')


# ── 3.2 AI Chat Analysis Flow ────────────────────────────────────────────────

add_heading(doc, '3.2 AI Chat Analysis Flow', 2)
doc.add_paragraph(
    'The AI Chat engine answers natural language questions about AWS costs by gathering live data '
    'from connected accounts, assembling it into a structured prompt, and sending it to Amazon '
    'Bedrock (Nova Lite) for analysis. This is the primary analysis path used in the Chat tab.'
)

add_heading(doc, '3.2.1 Data Gathering Pipeline', 3)
doc.add_paragraph(
    'When a user submits a question, the system gathers ALL relevant data from the selected '
    'accounts before invoking the AI model. The gathering pipeline collects:\n\n'
    '  Cost Data:\n'
    '    \u2022 Cost by service (30 days) via ce:GetCostAndUsage grouped by SERVICE\n'
    '    \u2022 Daily cost trend (7 days) via ce:GetCostAndUsage with DAILY granularity\n'
    '    \u2022 Monthly trend (6 months) for comparison questions — triggered by keywords like\n'
    '      "compare", "trend", "month over month", "growing"\n'
    '    \u2022 Usage type breakdown for top-cost services\n\n'
    '  Compute Resources:\n'
    '    \u2022 EC2 instances + 14-day CPU utilization metrics from CloudWatch\n'
    '    \u2022 RDS instances + CPU utilization + DatabaseConnections (14 days)\n'
    '    \u2022 Lambda functions + invocation counts + error rates (7 days)\n\n'
    '  Storage Resources:\n'
    '    \u2022 S3 buckets + lifecycle policy status (present/missing)\n'
    '    \u2022 EBS volumes + attachment state + volume type (gp2/gp3)\n\n'
    '  Network Resources:\n'
    '    \u2022 NAT Gateways + processed bytes\n'
    '    \u2022 VPC Endpoints\n'
    '    \u2022 Elastic IPs + association status\n\n'
    '  Optimization Data:\n'
    '    \u2022 AWS Compute Optimizer recommendations (rightsizing)\n'
    '    \u2022 AWS Pricing API for Savings Plans / Reserved Instance comparisons\n'
    '    \u2022 Knowledge base tips matching the question context\n'
    '    \u2022 FinOps Settings healthcheck results (cached from last scan)\n\n'
    'Data collection is gated by:\n'
    '  \u2022 Service presence: Only fetch data for services with actual CE spend > $0.01\n'
    '  \u2022 Question keywords: Specific questions (KMS, S3 lifecycle, budgets) skip irrelevant\n'
    '    EC2/VPC/NAT data collection\n'
    '  \u2022 Top cost services: NAT/VPC/EBS data fetched only when EC2-Other or VPC are in\n'
    '    the top 6 cost services'
)

add_heading(doc, '3.2.2 Prompt Assembly', 3)
doc.add_paragraph(
    'All gathered data is assembled into a structured prompt with the following sections:\n\n'
    '  1. System Instructions — 20+ rules governing response quality:\n'
    '     \u2022 Answer the specific question first — no generic cost summaries\n'
    '     \u2022 Use exact dollar amounts from actual data (never invent numbers)\n'
    '     \u2022 Include resource IDs and account IDs in recommendations\n'
    '     \u2022 Exclude Tax, Amazon Registrar, AWS Cost Explorer from savings analysis\n'
    '     \u2022 Group services < $0.50 as "Minor costs"\n'
    '     \u2022 Reference SlashMyBill features (Act \u2192 Waste Cleanup, Plan \u2192 Budget) instead\n'
    '       of directing users to the AWS Console\n'
    '     \u2022 Rightsize before recommending Savings Plans commitments\n'
    '     \u2022 Identify Graviton migration candidates from instance type family\n'
    '     \u2022 Show Cost Efficiency Score prominently for general questions\n'
    '     \u2022 Explain deleted-mid-month resources (not flagged as waste)\n'
    '     \u2022 Hebrew language support for comparison questions\n\n'
    '  2. Account Data — Per-account structured blocks containing:\n'
    '     \u2022 cost_by_service, daily_trend, monthly_trend\n'
    '     \u2022 ec2_instances, rds_instances, lambda_metrics\n'
    '     \u2022 s3_buckets, ebs_volumes, network_resources\n'
    '     \u2022 compute_optimizer_recommendations, budgets\n'
    '     \u2022 kms_customer_keys, pricing_data\n\n'
    '  3. Knowledge Base Context — Relevant tips from the 92-tip knowledge base,\n'
    '     filtered by question keywords and active services\n\n'
    '  4. User Question — The original natural language question'
)

add_heading(doc, '3.2.3 AI Response Generation', 3)
doc.add_paragraph(
    'The assembled prompt is sent to Amazon Bedrock (Nova Lite model). The AI generates a\n'
    'response that includes:\n\n'
    '  \u2022 Specific dollar amounts derived from actual Cost Explorer data\n'
    '  \u2022 Resource IDs (instance IDs, volume IDs, bucket names) and account IDs\n'
    '  \u2022 Actionable recommendations linking to SlashMyBill features:\n'
    '    - "Go to Act \u2192 Waste Cleanup" for idle resource cleanup\n'
    '    - "Go to Act \u2192 Scheduler" for automated stop/start\n'
    '    - "Go to Plan \u2192 Budget" for cost alerts\n'
    '    - "Go to Configure \u2192 FinOps Settings" for account health\n'
    '  \u2022 Follow-up drill-down questions for deeper analysis\n'
    '  \u2022 Chart suggestions when data is suitable for visualization'
)

add_heading(doc, '3.2.4 Post-Processing', 3)
doc.add_paragraph(
    'After the AI generates its response, the backend applies post-processing:\n\n'
    '  1. Navigation Link Conversion — Text patterns like "Go to Act \u2192 Waste Cleanup" are\n'
    '     converted to clickable navigation links that switch the user to the correct tab\n'
    '  2. Commands Log — The list of AWS API calls made during data gathering is attached\n'
    '     to the response for transparency\n'
    '  3. Drill-Down Buttons — Follow-up questions suggested by the AI are rendered as\n'
    '     clickable buttons below the answer\n'
    '  4. Table Detection — Responses containing tabular data get a "Show as table" button\n'
    '  5. Feedback Hooks — Each answer gets \U0001f44d/\U0001f44e feedback buttons that store\n'
    '     ratings in MemberPortal-AgentFeedback'
)


# ── 3.3 Bedrock Agent Flow ───────────────────────────────────────────────────

add_heading(doc, '3.3 Bedrock Agent Flow (Autonomous)', 2)
doc.add_paragraph(
    'The Bedrock Agent provides an alternative analysis path where the AI autonomously decides '
    'which data to fetch based on the user\'s question. Unlike the Chat flow (which pre-gathers '
    'all data), the Agent flow is iterative — the agent calls tools one at a time until it has '
    'enough context to answer.'
)

add_heading(doc, '3.3.1 Agent Invocation Flow', 3)
doc.add_paragraph(
    'Step 1 — User Question\n'
    '  User submits a question \u2192 POST /members/agent/invoke\n'
    '  The request includes the question text and selected account IDs.\n\n'
    'Step 2 — Agent Receives Question\n'
    '  Bedrock Agent (IDG5VJGUOZ5W) receives the question along with system instructions\n'
    '  that describe the available tools and how to use them.\n\n'
    'Step 3 — Tool Selection\n'
    '  The agent analyzes the question and decides which tool(s) to call. For example:\n'
    '    \u2022 "What are my top costs?" \u2192 getCostData\n'
    '    \u2022 "Compare last 3 months" \u2192 getMonthlyComparison\n'
    '    \u2022 "Which EC2 instances are idle?" \u2192 getEC2Instances\n'
    '    \u2022 "How can I save money?" \u2192 getCostData + getOptimizationTips + getEC2Instances\n\n'
    'Step 4 — Action Lambda Execution\n'
    '  The agent calls the action group Lambda (SlashMyBill-AgentAction) with the selected\n'
    '  tool name and parameters. The Lambda:\n'
    '    a. Assumes the cross-account role for the target account\n'
    '    b. Executes the specific AWS API calls for that tool\n'
    '    c. Returns structured data to the agent\n\n'
    'Step 5 — Iterative Reasoning\n'
    '  The agent may call multiple tools iteratively. After each tool response, it decides\n'
    '  whether it has enough data to answer or needs additional context. For example:\n'
    '    \u2192 getCostData reveals high EC2 spend\n'
    '    \u2192 Agent decides to also call getEC2Instances for details\n'
    '    \u2192 Agent sees idle instances, calls getOptimizationTips for recommendations\n\n'
    'Step 6 — Response Generation\n'
    '  Once the agent has sufficient data, it generates a final response.\n\n'
    'Step 7 — Streaming\n'
    '  The response is streamed back to the frontend in real-time.'
)

add_heading(doc, '3.3.2 Agent Action Tools (12)', 3)
add_table(doc,
    ['Tool Name', 'Data Retrieved', 'AWS APIs Used'],
    [
        ['getCostData', 'Cost by service (30 days), daily trend', 'ce:GetCostAndUsage'],
        ['getMonthlyComparison', '6-month cost trend by service', 'ce:GetCostAndUsage (MONTHLY)'],
        ['getEC2Instances', 'Instance list + CPU metrics', 'ec2:DescribeInstances, CloudWatch'],
        ['getRDSInstances', 'RDS instances + CPU + connections', 'rds:DescribeDBInstances, CloudWatch'],
        ['getLambdaFunctions', 'Functions + invocations + errors', 'lambda:ListFunctions, CloudWatch'],
        ['getS3Buckets', 'Buckets + lifecycle status', 's3:ListBuckets, s3:GetBucketLifecycle'],
        ['getEBSVolumes', 'Volumes + state + type + size', 'ec2:DescribeVolumes'],
        ['getNetworkResources', 'NAT GWs, VPC Endpoints, EIPs', 'ec2:Describe* (network)'],
        ['getBudgets', 'AWS Budgets + spend vs limit', 'budgets:DescribeBudgets'],
        ['getFinOpsSettings', 'Healthcheck results (cached)', 'DynamoDB lookup'],
        ['getOptimizationTips', 'Tips matching question context', 'DynamoDB query + filtering'],
        ['getAWSPricing', 'Savings Plans / RI pricing', 'pricing:GetProducts, ce:GetSavingsPlan*'],
    ],
    [2.0, 2.5, 3.5]
)

note_box(doc, 'The Bedrock Agent flow is complementary to the Chat flow. The Chat flow pre-gathers all data '
    '(faster for broad questions), while the Agent flow selectively fetches data (more efficient for targeted questions).')


# ── 3.4 FinOps Settings Healthcheck Flow ─────────────────────────────────────

add_heading(doc, '3.4 FinOps Settings Healthcheck Flow', 2)
doc.add_paragraph(
    'The FinOps Settings healthcheck audits a member\'s AWS account configuration against '
    'FinOps best practices. It detects the account type and runs a tailored checklist of '
    'scored and informational checks.'
)

add_heading(doc, '3.4.1 Healthcheck Execution Flow', 3)
doc.add_paragraph(
    'Step 1 — Trigger\n'
    '  User selects an account in Configure \u2192 FinOps Settings and clicks "Scan Settings".\n\n'
    'Step 2 — Assume Cross-Account Role\n'
    '  Backend assumes the cross-account role (SlashMyBill-{AccountID}) for the selected account.\n\n'
    'Step 3 — Detect Account Type\n'
    '  Call organizations:DescribeOrganization to retrieve the MasterAccountId.\n'
    '  Compare MasterAccountId with the connected AccountId:\n'
    '    \u2022 If they match \u2192 Management (payer) account\n'
    '    \u2022 If they differ \u2192 Linked (member) account\n'
    '    \u2022 If the call fails (no Organizations) \u2192 Standalone account (treated as management)\n\n'
    'Step 4 — Run Tailored Checklist\n'
    '  Based on the detected account type, run the appropriate set of checks.'
)

add_heading(doc, 'Management Account Checks', 3)
add_table(doc,
    ['Check', 'API Call', 'Pass Condition', 'Scored'],
    [
        ['Cost Allocation Tags', 'ce:ListCostAllocationTags', 'All user-defined tags activated', 'Yes'],
        ['AWS-Generated Tags', 'ce:ListCostAllocationTags (type=AWS)', 'aws:createdBy tag is active', 'Yes'],
        ['Cost Anomaly Detection', 'ce:GetAnomalyMonitors', 'At least one monitor exists', 'Yes'],
        ['Tag Backfill', 'ce:GetCostAllocationTagBackfillStatus', 'Backfill completed or in progress', 'Yes'],
        ['Budgets', 'budgets:DescribeBudgets', 'At least one budget configured', 'Yes'],
        ['Hourly Granularity', 'ce:GetCostAndUsage (HOURLY)', 'Hourly data available', 'Informational'],
        ['CE Preferences', 'ce:GetPreferences', 'Rightsizing recommendations enabled', 'Informational'],
        ['CUR Reports', 'cur:DescribeReportDefinitions', 'At least one CUR report configured', 'Informational'],
        ['Linked Billing Access', 'organizations:DescribeOrganization', 'Billing access enabled for linked accounts', 'Informational'],
    ],
    [2.0, 2.5, 2.0, 1.0]
)

add_heading(doc, 'Linked Account Checks', 3)
add_table(doc,
    ['Check', 'API Call', 'Pass Condition', 'Scored'],
    [
        ['Cost Anomaly Detection', 'ce:GetAnomalyMonitors', 'At least one monitor exists', 'Yes'],
        ['Compute Optimizer', 'compute-optimizer:GetEnrollmentStatus', 'Status = Active', 'Yes'],
        ['Budgets', 'budgets:DescribeBudgets', 'At least one budget configured', 'Yes'],
        ['Tag Coverage', 'tag:GetResources + tag policy keys', 'Coverage > 80% of resources tagged', 'Informational'],
        ['Hourly Granularity', 'ce:GetCostAndUsage (HOURLY)', 'Hourly data available', 'Informational'],
        ['Tag Activation Status', 'ce:ListCostAllocationTags', 'Tags activated by management admin', 'Informational'],
    ],
    [2.0, 2.5, 2.0, 1.0]
)

add_heading(doc, '3.4.2 Scoring and Caching', 3)
doc.add_paragraph(
    'Score Calculation:\n'
    '  Score = count of PASS items in scored group / total scored items\n'
    '  Example: 4 out of 5 scored checks pass \u2192 Score = 4/5 (80%)\n'
    '  Color coding: Green (\u226580%) | Amber (50-79%) | Red (<50%)\n\n'
    'Caching:\n'
    '  Results are cached in DynamoDB (MemberPortal-Members table) as part of the member record.\n'
    '  The cached results are used by:\n'
    '    \u2022 Dashboard KPI bar \u2192 FinOps Score card\n'
    '    \u2022 Act tab \u2192 FinOps Settings card in waste scan results\n'
    '    \u2022 AI Chat \u2192 Included in prompt context for relevant questions\n\n'
    'Fix Actions:\n'
    '  For checks with a Fix/Enable button, clicking it calls the appropriate AWS API through\n'
    '  the cross-account role to remediate the issue (e.g., activate tags, create anomaly monitor,\n'
    '  enroll in Compute Optimizer). The individual check result updates immediately without\n'
    '  requiring a full rescan.'
)

# ── 3.5 Dashboard Data Assembly ──────────────────────────────────────────────

add_heading(doc, '3.5 Dashboard Data Assembly', 2)
doc.add_paragraph(
    'The Observe tab dashboard assembles data from multiple sources in parallel to render '
    'KPI cards and 7 interactive chart widgets.'
)

add_heading(doc, '3.5.1 Data Assembly Flow', 3)
doc.add_paragraph(
    'Step 1 — Assume Cross-Account Roles\n'
    '  For each connected account in the member\'s selection, assume the cross-account role.\n\n'
    'Step 2 — Parallel Data Fetching\n'
    '  Fetch the following data streams in parallel for each account:\n'
    '    \u2022 Cost by service (30 days) \u2192 Treemap widget\n'
    '    \u2022 Daily cost trend (30 days) \u2192 Cost Trend widget\n'
    '    \u2022 Monthly cost trend (6 months) \u2192 Monthly Trend widget\n'
    '    \u2022 Waste detection (idle resources) \u2192 Waste Detection widget\n'
    '    \u2022 Rightsizing recommendations \u2192 Rightsizing widget\n'
    '    \u2022 Cost by region \u2192 Region breakdown\n'
    '    \u2022 Cost by tag \u2192 Tag Distribution widget\n'
    '    \u2022 Healthcheck results (cached) \u2192 FinOps Score KPI\n'
    '    \u2022 Budget data \u2192 Budget KPI card\n\n'
    'Step 3 — Compute KPIs\n'
    '    \u2022 Month-over-Month Change: (current_month - previous_month) / previous_month \u00d7 100\n'
    '    \u2022 Efficiency Score: (total_spend - potential_savings) / total_spend \u00d7 100\n'
    '    \u2022 Potential Savings: Sum of all identified waste and optimization opportunities\n'
    '    \u2022 FinOps Score: From cached healthcheck results\n'
    '    \u2022 Budget Utilization: Current spend / budget limit \u00d7 100\n\n'
    'Step 4 — Cache Results\n'
    '  Dashboard data is cached for 5 minutes, keyed by the set of selected account IDs.\n'
    '  Subsequent requests within the cache window return instantly.\n\n'
    'Step 5 — Live Business Metrics (Separate Cache)\n'
    '  Live Business Metrics are fetched separately with a 30-minute cache:\n'
    '    \u2022 Cognito user counts\n'
    '    \u2022 CloudFront distribution metrics\n'
    '    \u2022 ELB load balancer counts\n'
    '    \u2022 Route 53 hosted zone counts\n'
    '    \u2022 DynamoDB item counts\n'
    '    \u2022 Lambda invocation counts\n'
    '    \u2022 S3 object counts\n'
    '  These are auto-discovered from the connected accounts and plotted alongside cost data\n'
    '  to compute unit economics (cost per user, cost per transaction, etc.).'
)

tip_box(doc, 'The 5-minute cache prevents redundant API calls when switching between tabs. '
    'Click the refresh button to force a fresh data fetch.')

# ── 3.6 Tag Policy and Coverage Analysis ─────────────────────────────────────

add_heading(doc, '3.6 Tag Policy and Coverage Analysis', 2)
doc.add_paragraph(
    'The tag policy system lets members define required tag keys and measure coverage across '
    'their connected accounts.'
)

add_heading(doc, '3.6.1 Tag Policy Configuration', 3)
doc.add_paragraph(
    'Step 1 — Define Required Tags\n'
    '  Member defines required tag keys in Configure \u2192 Tag Policy (e.g., Environment,\n'
    '  Owner, CostCenter, Application).\n\n'
    'Step 2 — Store Policy\n'
    '  The tag policy is stored in DynamoDB as the tagPolicy attribute on the member record\n'
    '  in the MemberPortal-Members table.\n\n'
    'Step 3 — Policy Propagation\n'
    '  The tag policy keys are used across multiple features:\n'
    '    \u2022 Plan \u2192 Tag Resources: Pre-populates required keys in the bulk-tag form\n'
    '    \u2022 FinOps Settings \u2192 Tag Coverage check: Evaluates coverage against policy keys\n'
    '    \u2022 AI Chat: Recommends the member\'s specific tag keys instead of generic ones\n'
    '    \u2022 Dashboard \u2192 Tag Distribution widget: Defaults to showing policy keys'
)

add_heading(doc, '3.6.2 Tag Coverage Scan Flow', 3)
doc.add_paragraph(
    'Step 1 — Trigger\n'
    '  User navigates to Plan \u2192 Tag Resources and initiates a scan.\n\n'
    'Step 2 — Resource Discovery\n'
    '  For each connected account, call tag:GetResources with ResourcesPerPage=100.\n'
    '  Paginate through all resources to build a complete inventory.\n\n'
    'Step 3 — Coverage Evaluation\n'
    '  For each discovered resource, check if ALL required tag keys (from the tag policy)\n'
    '  are present on the resource. A resource with any missing required key is classified\n'
    '  as "needs tagging".\n\n'
    'Step 4 — Calculate Coverage\n'
    '  Coverage = resources with all required keys / total resources \u00d7 100\n'
    '  This percentage is displayed in the Tag Resources view and used by the FinOps Settings\n'
    '  tag coverage check.\n\n'
    'Step 5 — Display Results\n'
    '  Resources are listed with green \u2713 badges (fully tagged) or red badges (missing keys).\n'
    '  Members can select untagged resources and bulk-apply tags in one action.'
)


# ── 3.7 Invoice Explorer Flow ────────────────────────────────────────────────

add_heading(doc, '3.7 Invoice Explorer Flow', 2)
doc.add_paragraph(
    'The Invoice Explorer provides a structured, tabular interface for browsing AWS invoices '
    'across all connected accounts. It uses a cache-first architecture with DynamoDB to avoid '
    'repeated cross-account API calls and supports filtering, sorting, pagination, and CSV export.'
)

add_heading(doc, '3.7.1 Data Sync and Caching', 3)
doc.add_paragraph(
    'Step 1 — Cache Check\n'
    '  When a member opens the Invoices tab and selects an account, the system queries\n'
    '  the MemberPortal-Invoices DynamoDB table for cached invoice records.\n'
    '  Key structure: pk = {memberEmail}#{accountId}, sk = {YYYY-MM}#{serviceName}\n\n'
    'Step 2 — Cache Miss: Cross-Account Sync\n'
    '  If no cached data exists for the requested month, the Invoice Sync Service:\n'
    '    a. Assumes the cross-account role (SlashMyBill-{AccountID})\n'
    '    b. Calls Cost Explorer GetCostAndUsage with SERVICE granularity\n'
    '    c. Calls GetCostAndUsage with DAILY granularity for daily breakdown\n'
    '    d. Calls GetCostAndUsage grouped by SERVICE + USAGE_TYPE for usage details\n'
    '    e. Normalizes responses into flat DynamoDB records\n'
    '    f. Writes records with BatchWriteItem (TTL = 90 days)\n\n'
    'Step 3 — Cache Hit: Direct Return\n'
    '  Cached records are returned directly from DynamoDB without any cross-account calls.\n'
    '  This provides single-digit millisecond latency for repeat queries.\n\n'
    'Step 4 — Manual Refresh\n'
    '  Members can force a data refresh (rate-limited to 1 per account per 5 minutes).\n'
    '  The refresh deletes old records for the specified months and re-syncs from Cost Explorer.'
)

add_heading(doc, '3.7.2 API Endpoints', 3)
add_table(doc,
    ['Endpoint', 'Method', 'Purpose'],
    [
        ['/members/invoices', 'GET', 'List invoices with filters, sorting, and pagination'],
        ['/members/invoices/summary', 'GET', 'Spending totals, month-over-month change, top services'],
        ['/members/invoices/services', 'GET', 'Distinct service names for filter dropdown'],
        ['/members/invoices/refresh', 'POST', 'Force re-sync from Cost Explorer (rate-limited)'],
    ],
    [2.5, 1.0, 4.5]
)

add_heading(doc, '3.7.3 Server-Side Processing', 3)
doc.add_paragraph(
    'Filtering:\n'
    '  \u2022 Service: case-insensitive exact match\n'
    '  \u2022 Month: exact match on YYYY-MM\n'
    '  \u2022 Cost range: inclusive min/max with 2 decimal precision\n'
    '  \u2022 Search: case-insensitive substring on service name or usage type\n'
    '  \u2022 All filters combined with AND logic\n\n'
    'Sorting:\n'
    '  \u2022 By cost (numeric), service (alphabetical), or date (chronological)\n'
    '  \u2022 Secondary sort by cost descending for equal values\n'
    '  \u2022 Default: cost descending\n\n'
    'Pagination:\n'
    '  \u2022 Server-side with configurable page size (1-200, default 50)\n'
    '  \u2022 Returns metadata: page, pageSize, totalItems, totalPages\n\n'
    'Summary:\n'
    '  \u2022 Total cost for current month (sum of all items)\n'
    '  \u2022 Month-over-month percentage change\n'
    '  \u2022 Top 5 services by spend with percentage of total'
)

add_heading(doc, '3.7.4 DynamoDB Table: MemberPortal-Invoices', 3)
add_table(doc,
    ['Attribute', 'Type', 'Key', 'Description'],
    [
        ['pk', 'String', 'Partition Key', '{memberEmail}#{accountId}'],
        ['sk', 'String', 'Sort Key', '{YYYY-MM}#{serviceName}'],
        ['month', 'String', 'GSI Sort Key', 'Billing month (YYYY-MM)'],
        ['service', 'String', '\u2014', 'AWS service name'],
        ['cost', 'Number', '\u2014', 'Total cost in USD (2 decimal places)'],
        ['dailyCosts', 'Map', '\u2014', 'Daily breakdown {day: cost}'],
        ['usageTypes', 'List', '\u2014', 'Usage type details [{type, cost, unit, quantity}]'],
        ['ttl', 'Number', '\u2014', 'TTL epoch (90 days from sync)'],
        ['lastSyncedAt', 'String', '\u2014', 'ISO 8601 timestamp of last refresh'],
    ],
    [1.5, 1.0, 1.5, 4.0]
)

note_box(doc, 'Invoice data is automatically cleaned up after 90 days via DynamoDB TTL. '
    'Rate limiting uses a special record (pk=REFRESH#{accountId}, sk=RATE_LIMIT) with a short TTL.')


# ══════════════════════════════════════════════════════════════════════════════
# 4. Security Model
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '4. Security Model', 1)
doc.add_paragraph(
    'SlashMyBill implements defense-in-depth security across authentication, authorization, '
    'cross-account access, and data handling.'
)

add_table(doc,
    ['Layer', 'Mechanism', 'Details'],
    [
        ['Authentication', 'JWT Tokens (HS256)', '24-hour expiry, issued on login, validated on every API call'],
        ['Registration', '3-Step OTP Flow', 'Email verification via SES, 5-minute TTL, rate-limited to 1 per 60 seconds'],
        ['Password Storage', 'bcrypt with salt', 'Passwords never stored in plaintext'],
        ['Cross-Account Access', 'STS AssumeRole', 'ExternalId = SHA-256(member_email) prevents confused deputy attacks'],
        ['Credential Lifecycle', 'Temporary STS tokens', 'No long-lived credentials stored — fresh tokens per operation'],
        ['Lateral Access Prevention', '_verify_account_ownership()', 'Every account operation validates the account belongs to the requesting member'],
        ['API Security', 'CORS Restriction', 'API Gateway CORS restricted to slashmycloudbill.com origin'],
        ['Data Encryption', 'AWS-managed encryption', 'DynamoDB SSE at rest, S3 default encryption, HTTPS in transit'],
        ['IAM Permissions', 'Least privilege', 'ReadOnlyAccess + minimal inline policy for billing and cleanup actions'],
        ['Deployment', 'GitHub OIDC', 'No stored AWS credentials in CI/CD — OIDC federation with IAM role'],
        ['Write Actions', 'JIT Safety Checks', 'Resource state re-verified immediately before every destructive action'],
        ['CloudFormation', 'Versioned templates', 'Write permissions require updated CF template deployment by the member'],
    ],
    [2.0, 2.0, 4.0]
)

note_box(doc, 'The CloudFormation template includes a self-management permission that allows the stack to '
    'delete its own IAM role during account disconnection — enabling clean teardown without manual intervention.')

# ══════════════════════════════════════════════════════════════════════════════
# 5. Data Flow
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '5. Data Flow', 1)

add_heading(doc, '5.1 Data That Is NOT Stored', 2)
doc.add_paragraph(
    'The following data is fetched live from AWS APIs on every request and never persisted:\n\n'
    '  \u2022 Cost data — All Cost Explorer queries (cost by service, daily/monthly trends,\n'
    '    usage type breakdowns) are fetched in real-time\n'
    '  \u2022 Resource inventory — EC2 instances, RDS instances, Lambda functions, S3 buckets,\n'
    '    EBS volumes, NAT Gateways, EIPs, VPC Endpoints, KMS keys\n'
    '  \u2022 CloudWatch metrics — CPU utilization, database connections, invocation counts\n'
    '  \u2022 Compute Optimizer recommendations\n'
    '  \u2022 AWS Pricing data\n'
    '  \u2022 Budget details\n\n'
    'This design ensures members always see current data and eliminates the risk of stale\n'
    'or leaked cost information.'
)

add_heading(doc, '5.2 Data That IS Stored', 2)
add_table(doc,
    ['Data', 'Storage', 'Purpose'],
    [
        ['Member profiles', 'MemberPortal-Members', 'Email, hashed password, preferences, tag policy, allocation rules'],
        ['Account connections', 'MemberPortal-Accounts', 'Account ID, name, priority, hourly status, connection state'],
        ['Healthcheck results', 'MemberPortal-Members (cached)', 'FinOps Settings scan results for dashboard KPI and AI context'],
        ['Last scan findings', 'MemberPortal-Members (cached)', 'Top waste findings for Chat tab welcome screen'],
        ['Tag policy', 'MemberPortal-Members', 'Required tag keys and coverage threshold'],
        ['Schedules', 'MemberPortal-Members', 'Scheduler configurations and execution history'],
        ['AI feedback', 'MemberPortal-AgentFeedback', 'Thumbs up/down ratings and correction text per interaction'],
        ['Business metrics', 'MemberPortal-BusinessMetrics', 'Monthly business volumes for unit economics'],
        ['Knowledge base', 'ViewMyBill-CostOptimizationTips', '92 optimization tips with check specifications'],
        ['OTP codes', 'ViewMyBill-OTP', 'Verification codes with 5-minute TTL auto-expiry'],
    ],
    [2.0, 2.5, 3.5]
)

add_heading(doc, '5.3 DynamoDB Design', 2)
doc.add_paragraph(
    'Member data uses a single-table design pattern in DynamoDB:\n\n'
    '  \u2022 MemberPortal-Members: PK = email, stores profile, preferences, cached scan results,\n'
    '    allocation rules, tag policy, and scheduler configurations as nested attributes\n'
    '  \u2022 MemberPortal-Accounts: PK = memberEmail, SK = accountId, stores per-account\n'
    '    connection details and hourly granularity status\n\n'
    'The knowledge base (ViewMyBill-CostOptimizationTips) uses a separate table with:\n'
    '  PK = service (e.g., "EC2", "S3", "RDS")\n'
    '  SK = tipId (e.g., "ec2-001", "s3-002")\n'
    '  This allows efficient queries for all tips related to a specific service.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. CI/CD Pipeline
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '6. CI/CD Pipeline', 1)
doc.add_paragraph(
    'SlashMyBill uses GitHub Actions with AWS OIDC authentication for fully automated deployment.'
)

add_heading(doc, '6.1 Pipeline Trigger', 2)
doc.add_paragraph(
    'The pipeline triggers on push to the main branch (or manual workflow dispatch).\n'
    'Authentication uses GitHub OIDC federation with an IAM role — no stored AWS credentials.'
)

add_heading(doc, '6.2 Deployment Steps', 2)
doc.add_paragraph(
    'The pipeline executes the following steps in order:\n\n'
    '  1. Package Lambda Functions\n'
    '     Package 8 Lambda functions with their Python dependencies into ZIP archives\n'
    '     and upload to S3 (includes member-handler, agent-action, bill-analyzer,\n'
    '     scheduler-executor, admin-handler, upload-handler, otp-handler, contact-form).\n\n'
    '  2. Deploy CloudFormation Stack\n'
    '     Deploy the infrastructure stack with CAPABILITY_NAMED_IAM. This creates or updates\n'
    '     all Lambda functions, API Gateway routes, DynamoDB tables, IAM roles, and EventBridge\n'
    '     Scheduler resources.\n\n'
    '  3. Update Lambda Function Code\n'
    '     Update each Lambda function\'s code from the S3 packages to ensure the latest\n'
    '     version is deployed.\n\n'
    '  4. Seed DynamoDB Knowledge Base\n'
    '     Run the seed script to upsert all 92 optimization tips into the\n'
    '     ViewMyBill-CostOptimizationTips table. Tips are versioned — new fields are added\n'
    '     without overwriting existing feedback scores.\n\n'
    '  5. Update API Gateway Routes\n'
    '     Ensure all 25+ routes are configured with correct Lambda integrations.\n\n'
    '  6. Update Bedrock Agent\n'
    '     Update the Bedrock Agent instructions and OpenAPI schema to reflect any changes\n'
    '     to the 12 action tools.\n\n'
    '  7. Generate Documentation\n'
    '     Run _gen_ug.py and _gen_hld.py to generate the User Guide and High-Level Design\n'
    '     documents as .docx files.\n\n'
    '  8. Deploy Frontend\n'
    '     Sync website files to the S3 bucket (slashmycloudbill.com), inject the API Gateway\n'
    '     URL into frontend JavaScript, and invalidate the CloudFront cache to serve the\n'
    '     latest version.'
)

tip_box(doc, 'The pipeline is idempotent — running it multiple times produces the same result. '
    'CloudFormation handles resource creation vs update automatically.')

# ══════════════════════════════════════════════════════════════════════════════
# 7. Knowledge Base
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '7. Knowledge Base', 1)
doc.add_paragraph(
    'The knowledge base is a curated collection of 92 AWS cost optimization tips stored in '
    'DynamoDB. It serves as the ground truth for the scan engine, AI prompt context, and '
    'the Act tab card rendering.'
)

add_heading(doc, '7.1 Tip Schema', 2)
add_table(doc,
    ['Field', 'Type', 'Description'],
    [
        ['service', 'String (PK)', 'AWS service name: EC2, RDS, S3, EBS, Lambda, NAT, CloudFront, ELB, KMS, DynamoDB, ElastiCache, ECS, EKS, EFS, Data Transfer, General'],
        ['tipId', 'String (SK)', 'Unique identifier (e.g., ec2-001, s3-002, rds-003)'],
        ['category', 'String', 'right-sizing | pricing-model | scheduling | cleanup | lifecycle | architecture | monitoring'],
        ['title', 'String', 'Human-readable tip title'],
        ['description', 'String', 'Detailed explanation of the optimization opportunity'],
        ['estimatedSavings', 'String', 'Savings guidance (e.g., "20% of instance cost", "$3.65/month per EIP")'],
        ['difficulty', 'String', 'easy | medium | hard'],
        ['automatedCheck', 'String', 'Machine-readable check specification — describes the API calls and data fields to evaluate'],
        ['checkImplemented', 'Boolean', 'Whether the scan engine has a check function for this tip'],
        ['implementedInAct', 'Boolean', 'Whether the tip has a one-click action in the Act tab'],
        ['actionType', 'String', 'deep-link | advisory | delete | modify — determines the Act tab button behavior'],
        ['actionTarget', 'String', 'Which SlashMyBill feature handles this tip (e.g., "Act \u2192 Waste Cleanup", "Plan \u2192 Budget")'],
        ['level', 'Number', '1 (hygiene) | 2 (optimization) | 3 (architecture)'],
        ['serviceKey', 'String', 'Cost Explorer service name for presence-gating (e.g., "Amazon Elastic Compute Cloud - Compute")'],
        ['positiveCount', 'Number', 'Feedback score from user interactions (incremented on thumbs-up)'],
    ],
    [1.8, 1.2, 5.0]
)

add_heading(doc, '7.2 Coverage Statistics', 2)
doc.add_paragraph(
    'The knowledge base covers 16 AWS service categories with the following implementation status:\n\n'
    '  \u2022 Total tips: 92\n'
    '  \u2022 Check coverage: 100% — every tip has an automatedCheck specification\n'
    '  \u2022 Act coverage: 98% — nearly every tip has an actionType and actionTarget\n'
    '  \u2022 Scan implementation: Tips with checkImplemented=true have automated check functions\n'
    '    in the _SCAN_REGISTRY that run during the waste scan\n\n'
    'Service distribution:\n'
    '  EC2, RDS, S3, EBS, Lambda, NAT Gateway, CloudFront, ELB, KMS, DynamoDB,\n'
    '  ElastiCache, ECS, EKS, EFS, Data Transfer, General\n\n'
    'How tips are used:\n'
    '  1. Scan Engine — Tips with checkImplemented=true drive the automated waste scan.\n'
    '     The scan registry maps each tipId to a check function.\n'
    '  2. AI Prompts — Tips matching the user\'s question context are loaded into the\n'
    '     Bedrock prompt to provide domain-specific recommendations.\n'
    '  3. Act Tab — Tips with implementedInAct=true render as actionable cards with\n'
    '     one-click cleanup buttons.\n'
    '  4. Chat Widget — Top findings from the last scan are displayed on the Chat\n'
    '     welcome screen with "Ask \u25b6" buttons for follow-up questions.'
)

note_box(doc, 'Tips are seeded on every deployment. New tips or schema changes are applied automatically '
    'without overwriting existing feedback scores (positiveCount is preserved via conditional updates).')

# ── Save Document ─────────────────────────────────────────────────────────────

doc.save('SlashMyBill-HLD-v13.docx')
print('HLD saved: SlashMyBill-HLD-v13.docx')
