from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255,255,255)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F3864')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shd)
    for ri, row in enumerate(rows):
        tr = table.rows[ri+1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
        if ri % 2 == 0:
            for ci in range(len(row)):
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E8EDF5')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'), 'clear')
                tr.cells[ci]._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def tip_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('💡 Tip: ' + text)
    run.font.color.rgb = RGBColor(22, 101, 52)
    run.font.italic = True

def note_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run('⚠ Note: ' + text)
    run.font.color.rgb = RGBColor(146, 64, 14)
    run.font.italic = True

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SlashMyBill')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run('User Guide')
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(99, 102, 241)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Version 4.0  |  {datetime.date.today().strftime("%B %Y")}  |  AWS FinOps Platform').font.size = Pt(10)
doc.add_paragraph()

# ── Introduction ──────────────────────────────────────────────────────────────
add_heading(doc, '1. Introduction', 1)
doc.add_paragraph(
    'SlashMyBill is an AI-powered AWS cost optimization platform. It connects to your AWS accounts '
    'using read-only (and optionally write) access, analyzes your spending, and helps you identify '
    'and act on savings opportunities — all through a conversational interface and automated cleanup tools.\n\n'
    'Platform URL: https://www.eshkolai.com/members/'
)

add_heading(doc, '1.1 What You Can Do', 2)
add_table(doc,
    ['Feature', 'Description'],
    [
        ['Observe', 'FinOps dashboard with 7 interactive charts — cost by service, trends, waste, rightsizing'],
        ['Chat', 'Ask natural language questions about your AWS costs and get AI-powered answers'],
        ['Act', 'Scan for idle resources, clean up with one click, and create automated stop/start schedules'],
        ['Plan', 'Create AWS Budgets with alerts and manage resource tags for cost allocation'],
        ['Configure', 'Connect and manage your AWS accounts, audit FinOps settings with one-click fixes'],
    ],
    [2.0, 6.0]
)

# ── Getting Started ───────────────────────────────────────────────────────────
add_heading(doc, '2. Getting Started', 1)

add_heading(doc, '2.1 Create Your Account', 2)
doc.add_paragraph(
    '1. Go to https://www.eshkolai.com/members/\n'
    '2. Click "Register"\n'
    '3. Enter your email address and click "Send Code"\n'
    '4. Check your email for a 6-digit verification code (valid for 5 minutes)\n'
    '5. Enter the code and click "Verify"\n'
    '6. Set a password (minimum 8 characters) and click "Create Account"\n'
    '7. Log in with your email and password'
)
tip_box(doc, 'If you do not receive the verification email, check your spam folder. You can request a new code after 60 seconds.')

add_heading(doc, '2.2 Connect Your First AWS Account', 2)
doc.add_paragraph(
    'SlashMyBill uses a secure cross-account IAM role to read your AWS cost data. '
    'No credentials are stored — access is granted via a CloudFormation template you deploy.\n\n'
    'Steps:\n'
    '1. Go to the Configure tab\n'
    '2. Click "+ Add Account"\n'
    '3. Enter your 12-digit AWS Account ID and an optional name\n'
    '4. Click "Add" — the Setup Wizard will open automatically\n\n'
    'Setup Wizard (4 steps):\n'
    '  Step 1 — Download Template: Click "Download CF Template" to get the CloudFormation YAML file\n'
    '  Step 2 — Deploy in AWS: Click "Open in CloudFormation Console" or upload the template manually. '
    'Review the IAM permissions and click "Create Stack"\n'
    '  Step 3 — Wait: The stack takes 1-2 minutes to deploy\n'
    '  Step 4 — Test & Configure: Click "Test Connection" to verify access. '
    'The system will also check if hourly cost granularity is enabled'
)
note_box(doc, 'The CloudFormation stack creates an IAM role named SlashMyBill-{AccountID} in your AWS account. '
    'This role has ReadOnlyAccess plus billing/cost permissions. It does NOT have access to your application data.')

add_heading(doc, '2.3 Enable Hourly Cost Tracking (Optional)', 2)
doc.add_paragraph(
    'For real-time hourly cost charts, enable hourly granularity in AWS Cost Explorer:\n\n'
    '1. Sign in to the management (payer) account\n'
    '2. Go to AWS Cost Management → Cost Explorer → Settings\n'
    '3. Check "Hourly and Resource Level Data"\n'
    '4. Click Save\n'
    '5. Wait 24-48 hours for data to appear\n\n'
    'Cost: ~$0.01 per 1,000 usage records/month (typically < $1/month)\n\n'
    'To check status: click the ⏱ button next to any account in the Configure tab, '
    'or run "Test Connection" — the result will show ⏱✓ or ⏱✗'
)
note_box(doc, 'This setting must be enabled from the management (payer) account. '
    'It cannot be enabled via API — it requires a manual action in the AWS Console.')

# ── Observe Tab ───────────────────────────────────────────────────────────────
add_heading(doc, '3. Observe Tab — FinOps Dashboard', 1)
doc.add_paragraph(
    'The Observe tab is your FinOps command center. It loads automatically when you log in '
    'and shows data from all your connected accounts.'
)

add_heading(doc, '3.1 KPI Bar', 2)
doc.add_paragraph(
    'At the top of the dashboard, four key metrics are shown:\n'
    '  • Month-over-Month: Cost change vs previous month (green = decrease, red = increase)\n'
    '  • Efficiency Score: 0-100% score based on identified waste vs total spend\n'
    '  • Potential Savings: Total monthly savings identified — click to open Chat with a savings question\n'
    '  • Accounts: Number of connected accounts included in the view\n'
    '  • FinOps Score: Settings health score (e.g., 7/8) — click to navigate to Configure → FinOps Settings\n'
    '  • Budget: Current spend vs budget limit — click to navigate to Plan → Budget'
)

add_heading(doc, '3.2 Dashboard Widgets', 2)
add_table(doc,
    ['Widget', 'What It Shows', 'How to Use'],
    [
        ['Cost by Service', 'Treemap of spending by AWS service', 'Click a tile to drill into usage types. Click again to go back.'],
        ['Cost Trend', 'Daily or hourly cost over time', 'Toggle Daily/Hourly. Red markers = anomalies (> 2× average).'],
        ['Cost Allocation', 'Costs by business unit (virtual tags)', 'Click "Manage Rules" to define business units.'],
        ['Waste Detection', 'Idle resources with dollar amounts', 'Click "Chat ▶" to ask the AI about specific waste items.'],
        ['Rightsizing', 'Over-provisioned instances', 'Shows Compute Optimizer recommendations with savings.'],
        ['Monthly Trend', 'Cost by service month-over-month', 'Stacked bar chart — each color is a service.'],
        ['Unit Cost Trend', 'Cost per business unit (e.g., per user)', 'Requires business metrics to be configured.'],
    ],
    [1.8, 2.5, 3.7]
)

add_heading(doc, '3.3 Account Selection', 2)
doc.add_paragraph(
    'Use the account selector dropdown (top right of the dashboard) to choose which accounts '
    'to include. All accounts are selected by default. Your selection is preserved when you '
    'switch between Observe, Chat, and Act tabs.'
)

add_heading(doc, '3.4 Cost by Service Drill-Down', 2)
doc.add_paragraph(
    'The Cost by Service treemap supports 2-phase drill-down:\n\n'
    '  Level 1 (Services): Shows all AWS services as colored tiles sized by cost\n'
    '  Level 2 (Usage Types): Click any service tile to see its usage type breakdown '
    '(e.g., EC2-Other breaks into VolumeUsage.gp3, NatGateway-Hours, DataTransfer)\n\n'
    'Navigation:\n'
    '  • Click a tile → drill into usage types\n'
    '  • Click "← All Services" breadcrumb → return to service view\n'
    '  • Click "Details ↗" → open side panel with bar chart + AI chat button'
)

add_heading(doc, '3.5 Live Business Metrics Widget', 2)
doc.add_paragraph(
    'The Live Business Metrics widget auto-discovers real operational metrics from your connected '
    'AWS accounts and plots them alongside cost data to compute unit economics.\n\n'
    'Auto-discovered metrics include:\n'
    '  • Cognito user counts\n'
    '  • DynamoDB item counts\n'
    '  • API Gateway request counts\n'
    '  • Route 53 DNS queries\n'
    '  • CloudWatch custom metrics\n'
    '  • Lambda invocations\n'
    '  • S3 object counts\n\n'
    'The widget displays a dual-axis chart:\n'
    '  • Purple bars (left axis): Volume / business metric count\n'
    '  • Amber line (right axis): Cost per unit\n\n'
    'Use the metric selector dropdown to switch between "Auto-Discovered" and "Manual" metric groups. '
    'The cost dimension selector lets you choose Total Account Cost, per-service cost, or tag-based cost '
    'for the unit economics calculation. Each auto-discovered metric is auto-mapped to a default cost dimension '
    '(e.g., Cognito users → Cognito service cost).\n\n'
    'The chart shows a 6-month time-series by default. This widget replaces the old manual "Unit Cost Trend" widget.'
)
tip_box(doc, 'Auto-discovered metrics require no manual configuration — just connect your AWS account and the widget populates automatically.')

add_heading(doc, '3.6 Tag Distribution Widget', 2)
doc.add_paragraph(
    'The Tag Distribution widget shows resource count distribution by tag value as a donut chart. '
    'It uses the AWS Resource Groups Tagging API to scan your connected accounts.\n\n'
    'Click "Manage Tags ▶" to navigate to Plan → Tag Resources for bulk tag management.'
)

add_heading(doc, '3.7 Budget KPI Card', 2)
doc.add_paragraph(
    'The Budget KPI card appears in the KPI bar and shows your current spend vs budget limit '
    'as a progress bar. If you have active AWS Budgets configured via the Plan tab, the card '
    'displays the budget with the highest utilization percentage.\n\n'
    'Click the Budget KPI card to navigate directly to Plan → Budget for detailed budget management.'
)
# ── Chat Tab ──────────────────────────────────────────────────────────────────
add_heading(doc, '4. Chat Tab — AI Agent', 1)
doc.add_paragraph(
    'The Chat tab lets you ask natural language questions about your AWS costs. '
    'The AI analyzes your real account data and provides specific, actionable answers.'
)

add_heading(doc, '4.1 Top Findings Widget', 2)
doc.add_paragraph(
    'When you open the Chat tab, the welcome screen shows your top savings findings:\n\n'
    '  • 🔴 Red = high savings (> $20/month)\n'
    '  • 🟡 Yellow = medium savings ($5-20/month)\n'
    '  • 🟢 Green = low savings (< $5/month)\n\n'
    'Click any finding row or its "Ask ▶" button to populate the Ask box with a suggested question. '
    'Then press Enter or click Ask to submit.\n\n'
    'Click "🔍 Scan for Savings Opportunities" to run a fresh scan.\n'
    'Click "↻ Refresh Findings" (in the header) to rescan at any time, even mid-conversation.'
)

add_heading(doc, '4.2 Asking Questions', 2)
doc.add_paragraph(
    'Type your question in the Ask box and press Enter or click Ask.\n\n'
    'Example questions:\n'
    '  • "How efficient is my account?"\n'
    '  • "Where can I save money?"\n'
    '  • "Compare my costs over the last 3 months"\n'
    '  • "Which S3 buckets need lifecycle policies?"\n'
    '  • "List Lambda transactions for Jan, Feb, March"\n'
    '  • "Which EC2 instances are over-provisioned?"\n'
    '  • "How do I set up AWS Budgets with cost alerts?"\n'
    '  • "Which of my instances can use Spot pricing?"\n\n'
    'The AI will show the AWS API commands it executed, then provide a detailed answer '
    'with specific resource IDs, dollar amounts, and actionable recommendations.\n\n'
    'The AI always recommends SlashMyBill in-app features when relevant — for example, '
    '"Go to Plan → Budget" to create cost alerts, or "Go to Act → Scheduler" to automate '
    'stop/start schedules — instead of directing you to the AWS Console.'
)
tip_box(doc, 'For multi-account questions, select multiple accounts using the account dropdown. '
    'The AI will analyze all selected accounts and provide per-account breakdowns.')

add_heading(doc, '4.3 Answer Features', 2)
doc.add_paragraph(
    'Each AI answer includes:\n'
    '  • Commands log: Shows exactly which AWS APIs were called\n'
    '  • Drill-down buttons: Follow-up questions based on the answer\n'
    '  • Show as table: Convert data to a sortable table (only shown when relevant)\n'
    '  • 👍/👎 Feedback: Rate the answer to improve future responses\n\n'
    'Click "📋 Copy" to copy the answer text to your clipboard.'
)

add_heading(doc, '4.4 Font Size', 2)
doc.add_paragraph(
    'Use the A- and A+ buttons in the header to adjust the chat font size. '
    'The setting applies to both the chat messages and the findings widget.'
)

# ── Act Tab ───────────────────────────────────────────────────────────────────
add_heading(doc, '5. Act Tab — Resource Cleanup', 1)
doc.add_paragraph(
    'The Act tab scans your AWS accounts for idle and wasted resources, '
    'then lets you clean them up with one click.'
)

add_heading(doc, '5.1 Running a Scan', 2)
doc.add_paragraph(
    '1. Select accounts using the dropdown (same style as Chat/Observe)\n'
    '2. Click "🔍 Scan for Waste"\n'
    '3. Wait 10-20 seconds for the scan to complete\n'
    '4. Review the 7 category cards\n\n'
    'Each category always shows — either a ✓ clean card or a ⚠ findings card with savings amount.'
)

add_heading(doc, '5.2 Understanding the Cards', 2)
add_table(doc,
    ['Card', 'What It Finds', 'Action'],
    [
        ['🌐 Elastic IPs', 'Unassociated IPs ($3.65/mo each)', 'Release Address'],
        ['💾 EBS Volumes', 'Unattached volumes ($0.10/GB/mo)', 'Delete Volume'],
        ['⚖️ Load Balancers', 'LBs with 0 healthy targets ($16/mo)', 'Delete Load Balancer'],
        ['🪣 S3 Buckets', 'No lifecycle policy or inactive 90+ days', 'Apply Lifecycle / Browse'],
        ['🖥️ EC2 Instances', 'Avg CPU < 5% over 14 days', 'Stop Instance'],
        ['🗄️ RDS Instances', 'Avg CPU < 5%, < 2 connections over 14 days', 'Delete (with snapshot)'],
        ['📸 EBS Snapshots', 'Older than 180 days', 'Delete Snapshot'],
    ],
    [1.8, 3.0, 2.2]
)

add_heading(doc, '5.3 S3 Bucket Details', 2)
doc.add_paragraph(
    'The S3 card shows each flagged bucket with:\n'
    '  • Size and estimated monthly cost\n'
    '  • Last activity (days since last object change)\n'
    '  • Reason badges: "No lifecycle", "Inactive 95d", "Empty"\n\n'
    'Click "Browse" on any bucket to see its contents:\n'
    '  • Object list with size, last modified date, and age\n'
    '  • Objects 90+ days old highlighted in red\n'
    '  • Sort by: Oldest first | Largest first | Newest first\n'
    '  • "Apply Lifecycle Policy" — adds Intelligent-Tiering after 90 days\n'
    '  • "Delete All Objects" — permanently removes all objects (bucket remains)'
)
note_box(doc, 'Delete All Objects is irreversible. A confirmation dialog will show the object count and size before proceeding.')

add_heading(doc, '5.4 Safety Guardrails', 2)
doc.add_paragraph(
    'Before every cleanup action, the system performs a Just-In-Time (JIT) check:\n'
    '  • EIP: Verifies still unassociated (skips if now attached)\n'
    '  • EBS Volume: Verifies still in "available" state (skips if reattached)\n'
    '  • Load Balancer: Verifies still has 0 healthy targets (skips if traffic resumed)\n'
    '  • EC2: Checks for Auto Scaling Group membership (detaches from ASG before stopping)\n'
    '  • RDS: Always creates a final snapshot before deletion\n'
    '  • Snapshots: Verifies still older than 180 days'
)

add_heading(doc, '5.5 IAM Permissions for Write Actions', 2)
doc.add_paragraph(
    'Write actions (delete, stop, apply lifecycle) require an updated CloudFormation template. '
    'If you see a blue "⚠ Requires updated IAM role" banner:\n\n'
    '1. Go to Configure tab\n'
    '2. Click "↓ Download CF Template" for the affected account\n'
    '3. In AWS CloudFormation, find the stack SlashMyBill-Access-{AccountID}\n'
    '4. Click Update → Replace current template → upload the new file\n'
    '5. Review the new IAM permissions and confirm\n\n'
    'Click "How to update →" in the banner for step-by-step guidance.'
)

add_heading(doc, '5.6 Scheduler — Automated Stop/Start', 2)
doc.add_paragraph(
    'The Scheduler lets you create automated stop/start schedules for your AWS resources. '
    'Instead of manually stopping dev instances every evening, SlashMyBill does it for you.\n\n'
    'How to create a schedule:\n'
    '  1. Go to Act → Scheduler sub-tab\n'
    '  2. Click "+ New Schedule"\n'
    '  3. Select the target account\n'
    '  4. Choose a schedule type:\n'
    '     Stop/Start types: EC2 Stop/Start, RDS Stop/Start, ASG Scale to 0, EKS Scale to 0, '
    'SageMaker Stop, Redshift Pause, WorkSpaces Auto-Stop, ELB Teardown\n'
    '     Review types: Waste Scan, Snapshot Cleanup, gp2→gp3 Migration, SP/RI Review\n'
    '  5. Set a name, frequency (weekdays/daily/weekly/monthly/quarterly), days, stop/start times, and timezone\n'
    '  6. Select specific resources from the list or use a tag filter (e.g., Environment=dev)\n'
    '  7. Click "Create Schedule"\n\n'
    'What happens behind the scenes:\n'
    '  • SlashMyBill creates real EventBridge Scheduler rules in the platform account\n'
    '  • At the scheduled time, the executor Lambda assumes your cross-account role and performs the action\n'
    '  • For stop/start types, two schedules are created — one for stop and one for start\n\n'
    'Schedule cards show:\n'
    '  • Active/Paused status badge\n'
    '  • Next execution time in your configured timezone\n'
    '  • Execution history: ✅ success, ⚠️ partial (some resources failed), ❌ failure\n'
    '  • Per-resource details for partial/failed runs\n\n'
    'Controls on each card:\n'
    '  • Pause — disables the EventBridge schedule without deleting it\n'
    '  • Resume — re-enables a paused schedule\n'
    '  • Delete — removes the schedule and its EventBridge rules permanently'
)
tip_box(doc, 'For stop/start types, two schedules are created — one for stop and one for start. Both use the same days and timezone.')
note_box(doc, 'The cross-account IAM role must include write permissions for the scheduled action type. Redeploy the latest CloudFormation template if needed.')

add_heading(doc, '5.7 Plan Tab — Budget Management', 2)
doc.add_paragraph(
    'The Plan tab lets you create and manage AWS Budgets directly from SlashMyBill. '
    'Budgets are created in your actual AWS account via the AWS Budgets API.\n\n'
    'Creating a budget:\n'
    '  1. Go to Plan → Budget\n'
    '  2. Click "+ Create Budget"\n'
    '  3. Enter a monthly budget amount (e.g., $500)\n'
    '  4. Configure alert thresholds — default thresholds are 50%, 75%, 100%, and 120% of the budget\n'
    '  5. Add email addresses for notifications\n'
    '  6. Optionally add tag-based budget filtering (e.g., Environment=production) to scope the budget to specific resources\n'
    '  7. Click "Create"\n\n'
    'Managing budgets:\n'
    '  • Edit: Change the budget amount or alert thresholds\n'
    '  • Delete: Remove the budget from your AWS account\n'
    '  • View: See existing budgets with spend vs limit progress bars showing current utilization'
)
tip_box(doc, 'Budget alerts come directly from AWS, not SlashMyBill — so they work even if you are not logged in.')
note_box(doc, 'For tag-based budgets, use the TagKeyValue format (e.g., user:Environment$production). The Plan tab handles this formatting automatically.')

add_heading(doc, '5.8 Plan Tab — Tag Resources', 2)
doc.add_paragraph(
    'The Tag Resources sub-tab scans all resources across your connected accounts for tag coverage '
    'and lets you bulk-apply tags.\n\n'
    'Features:\n'
    '  • Scans resources using the AWS Resource Groups Tagging API\n'
    '  • Shows tagged resources with green ✓ badges and untagged resources with red badges\n'
    '  • Pre-populates required tag keys: Environment, Owner, CostCenter, Application\n'
    '  • Select multiple resources and bulk-apply tags in one action\n'
    '  • Sticky table headers for easy scrolling through large resource lists\n\n'
    'To tag resources:\n'
    '  1. Go to Plan → Tag Resources\n'
    '  2. Review the resource list — untagged resources are highlighted\n'
    '  3. Select resources using the checkboxes\n'
    '  4. Enter tag key-value pairs\n'
    '  5. Click "Apply Tags" to tag all selected resources'
)

add_heading(doc, '5.9 Paddle Payment Integration', 2)
doc.add_paragraph(
    'SlashMyBill uses Paddle as its Merchant of Record for all payments.\n\n'
    '  • Upgrade plans: Click the "Upgrade" button in the header to switch to Growth or Scale\n'
    '  • Buy token top-ups: Click the 🪙 coin icon in the header to purchase additional tokens\n'
    '  • All payments are processed securely by Paddle, which handles tax, invoicing, and compliance'
)

# ── Configure Tab ─────────────────────────────────────────────────────────────
add_heading(doc, '6. Configure Tab — Account Management', 1)

add_heading(doc, '6.1 Account Table', 2)
doc.add_paragraph(
    'The Configure tab shows all your connected AWS accounts with:\n'
    '  • Status badge: pending | connected | failed | partial\n'
    '  • Hourly badge: ⏱✓ (enabled) or ⏱✗ (not enabled)\n'
    '  • Last tested date\n\n'
    'Action buttons per account:\n'
    '  ▲▼ — Reorder priority (affects dashboard and AI query order)\n'
    '  ↓ — Download CloudFormation template\n'
    '  ⚡ — Test connection (also checks hourly status)\n'
    '  ⏱ — Enable hourly granularity guide\n'
    '  ✏ — Edit account name or ID\n'
    '  🗑 — Delete account connection'
)

add_heading(doc, '6.2 Deleting an Account', 2)
doc.add_paragraph(
    'When you delete an account connection:\n'
    '  1. SlashMyBill assumes the cross-account role\n'
    '  2. Detaches all managed policies from the IAM role\n'
    '  3. Deletes all inline policies\n'
    '  4. Deletes the IAM role\n'
    '  5. Deletes the CloudFormation stack\n'
    '  6. Removes the account from SlashMyBill\n\n'
    'If the stack deletion fails (e.g., older template without iam:DetachRolePolicy), '
    'you will see a warning with instructions to clean up manually.'
)

add_heading(doc, '6.3 FinOps Settings Healthcheck', 2)
doc.add_paragraph(
    'The FinOps Settings section audits your AWS account\'s billing and cost management configuration '
    'against FinOps best practices. It detects whether your account is a management (payer) or linked '
    'account and runs a tailored checklist.\n\n'
    'To access:\n'
    '  1. Go to Configure tab\n'
    '  2. Click "FinOps Settings" in the left navigation\n'
    '  3. Select an account from the dropdown\n'
    '  4. Click "Scan Settings" to run the audit\n\n'
    'The scan checks different settings based on your account type:'
)

add_heading(doc, 'Management Account Checks (8 scored + 1 informational)', 3)
add_table(doc,
    ['Check', 'What It Verifies', 'Fix Available'],
    [
        ['Cost Allocation Tags (User-Defined)', 'All user-defined tags are activated for cost reporting', 'Yes — Activate all tags'],
        ['AWS-Generated Tags', 'aws:createdBy tag is active', 'Yes — Activate tag'],
        ['Cost Anomaly Detection', 'At least one anomaly monitor exists', 'Yes — Create monitor + email subscription'],
        ['Hourly Granularity', 'Hourly cost data is available', 'No — must enable in AWS Cost Explorer Settings'],
        ['CE Preferences (Right-Sizing)', 'Rightsizing recommendations are enabled', 'Yes — Enable preferences'],
        ['Cost and Usage Report (CUR)', 'At least one CUR report is configured', 'No — requires S3 bucket setup in AWS'],
        ['Tag Backfill', 'Historical billing data reflects current tags', 'Yes — Start backfill (up to 12 months)'],
        ['Linked Account Billing Access', 'Linked accounts can view billing data', 'Informational — not scored'],
        ['Budgets', 'At least one AWS Budget is configured', 'Link to Plan → Budget'],
    ],
    [2.5, 3.0, 2.5]
)

add_heading(doc, 'Linked Account Checks (6 scored)', 3)
add_table(doc,
    ['Check', 'What It Verifies', 'Fix Available'],
    [
        ['Resource Tag Coverage', 'Percentage of resources with tags (>80% = pass)', 'Link to Plan → Tag Resources'],
        ['Budgets', 'At least one AWS Budget is configured', 'Link to Plan → Budget'],
        ['Cost Anomaly Detection', 'At least one anomaly monitor exists', 'Yes — Create monitor + email subscription'],
        ['Compute Optimizer', 'AWS Compute Optimizer is enrolled', 'Yes — Enroll in Compute Optimizer'],
        ['Hourly Granularity', 'Hourly cost data is available', 'No — must be enabled from management account'],
        ['Tag Activation Status', 'Cost allocation tags are activated by management admin', 'Read-only — contact management admin'],
    ],
    [2.5, 3.0, 2.5]
)

doc.add_paragraph(
    'FinOps Score:\n'
    '  The score shows X/Y where X is the number of passing checks and Y is the total scored checks.\n'
    '  Informational items (like Linked Account Billing Access) are excluded from the score.\n'
    '  Color coding: Green (≥80%) | Amber (50-79%) | Red (<50%)\n\n'
    'Fix Actions:\n'
    '  For checks with a Fix/Enable/Setup button, click it to apply the fix directly through your\n'
    '  cross-account role. The item updates immediately without requiring a full rescan.\n\n'
    'Integration with other tabs:\n'
    '  • Dashboard: FinOps Score KPI card shows your score at a glance\n'
    '  • Act tab: FinOps Settings card appears in waste scan results when issues exist\n'
    '  • AI Chat: The AI recommends FinOps Settings fixes when relevant to your questions'
)
tip_box(doc, 'If fix actions fail with "Permission denied", redeploy the latest CloudFormation template for that account. '
    'The updated template includes the write permissions needed for FinOps Settings fixes.')
note_box(doc, 'Accounts connected before the FinOps Settings feature was released need their CloudFormation stack updated '
    'to include the new IAM permissions. Download the latest template from the Configure tab.')

# ── Virtual Tagging ───────────────────────────────────────────────────────────
add_heading(doc, '7. Virtual Tagging & Cost Allocation', 1)
doc.add_paragraph(
    'Virtual tagging lets you group AWS costs into business units without changing your '
    'actual AWS resource tags.\n\n'
    'To set up:\n'
    '1. In the Observe tab, click "Manage Rules" on the Cost Allocation widget\n'
    '2. Click "+ Add Business Unit"\n'
    '3. Name the business unit (e.g., "Data Science Team")\n'
    '4. Add rules:\n'
    '   • Dimension: Account | Service | Tag\n'
    '   • Operator: equals | contains | startsWith\n'
    '   • Value: the matching value\n'
    '5. Set rule logic: AND (all rules must match) or OR (any rule matches)\n'
    '6. For shared costs, choose: Even split | Proportional | Custom %\n'
    '7. Click Save\n\n'
    'The Cost Allocation treemap will update to show costs by business unit.'
)

# ── Unit Economics ────────────────────────────────────────────────────────────
add_heading(doc, '8. Unit Economics', 1)
doc.add_paragraph(
    'Unit Economics tracks your cost per business output (e.g., cost per user, cost per transaction).\n\n'
    'To add business metrics:\n'
    '1. In the Observe tab, find the Unit Cost Trend widget\n'
    '2. Click "Add Metric"\n'
    '3. Enter: Metric Name (e.g., "ActiveUsers"), Month (YYYY-MM), Volume (e.g., 50000)\n'
    '4. Click Save\n\n'
    'Auto-discovered IT metrics (no manual entry needed):\n'
    '  • DynamoDB items count\n'
    '  • API Gateway request count\n'
    '  • Lambda invocation count\n'
    '  • S3 bucket count\n\n'
    'The Unit Cost Trend widget shows a dual-axis chart:\n'
    '  • Bar (left axis): Business volume\n'
    '  • Line (right axis): Cost per unit\n\n'
    'If costs increase but cost-per-unit decreases, the AI will frame this as "efficient scaling".'
)

# ── Troubleshooting ───────────────────────────────────────────────────────────
add_heading(doc, '9. Troubleshooting', 1)
add_table(doc,
    ['Issue', 'Cause', 'Solution'],
    [
        ['Connection test fails', 'CloudFormation stack not deployed or role name mismatch', 'Re-run the Setup Wizard and deploy the template'],
        ['Hourly chart shows no data', 'Hourly granularity not enabled in Cost Explorer', 'Enable in AWS Cost Explorer Settings (management account)'],
        ['Scan fails with "Service Unavailable"', 'Lambda timeout (scan too slow)', 'Try scanning fewer accounts at once'],
        ['Write action fails with Access Denied', 'Old CF template without write permissions', 'Redeploy the latest CF template (see Act tab warning)'],
        ['Stack deletion fails', 'Old template missing iam:DetachRolePolicy', 'Redeploy latest template first, then delete'],
        ['AI gives generic answers', 'Question too broad or data not fetched', 'Be specific: include service name, time period, account'],
        ['No findings in Act tab', 'Accounts are clean or scan not run yet', 'Click "Scan for Waste" to run a fresh scan'],
        ['OTP email not received', 'Email in spam or rate limit hit', 'Check spam; wait 60 seconds before requesting a new code'],
        ['Schedule creation fails', 'Account not connected or missing permissions', 'Verify account is connected and redeploy latest CF template'],
        ['Schedule shows ❌ failure', 'Cross-account role missing write permissions', 'Redeploy the latest CloudFormation template for the target account'],
        ['Budget creation fails with InvalidParameterException', 'Tag filter uses unsupported dimension', 'Use TagKeyValue format for tag-based budgets'],
        ['Live metrics shows no data', 'No connected accounts or permissions missing', 'Connect an account and redeploy latest CF template'],
        ['FinOps Settings fix fails with Permission denied', 'Old CF template without healthcheck write permissions', 'Download and redeploy the latest CF template from Configure tab'],
        ['FinOps Settings scan shows all errors', 'Cross-account role cannot be assumed', 'Verify account is connected and test the connection first'],
        ['Tag Backfill fix fails', 'BackfillFrom date format issue or backfill already running', 'Wait for any in-progress backfill to complete, then retry'],
        ['FinOps Score not showing on dashboard', 'No scan has been run yet', 'Go to Configure → FinOps Settings and run a scan'],
    ],
    [2.0, 2.5, 3.5]
)

# ── FAQ ───────────────────────────────────────────────────────────────────────
add_heading(doc, '10. Frequently Asked Questions', 1)

faqs = [
    ('Is my AWS data safe?',
     'SlashMyBill uses read-only access by default. The IAM role has ReadOnlyAccess plus billing permissions. '
     'No application data (S3 objects, database contents, etc.) is ever read. Write permissions are optional '
     'and only used when you explicitly click a cleanup action.'),
    ('Can SlashMyBill modify my AWS resources without my permission?',
     'No. Every cleanup action requires you to click a button and confirm a dialog. '
     'A JIT safety check runs immediately before each action to verify the resource state has not changed.'),
    ('How much does SlashMyBill cost?',
     'Contact the SlashMyBill team at https://www.eshkolai.com for pricing information.'),
    ('How often is the dashboard data refreshed?',
     'Dashboard data is cached for 5 minutes. Click the refresh button to force a reload. '
     'Cost Explorer data itself is updated by AWS once every 24 hours.'),
    ('Can I connect multiple AWS accounts?',
     'Yes. You can connect up to 5 accounts and analyze them together or individually. '
     'Use the account selector in each tab to choose which accounts to include.'),
    ('What happens when I delete an account connection?',
     'SlashMyBill removes the IAM role and CloudFormation stack from your AWS account, '
     'then removes the account from SlashMyBill. Your AWS resources are not affected.'),
    ('Does SlashMyBill support AWS Organizations?',
     'Yes. Connect your management (payer) account for organization-wide cost visibility. '
     'For linked accounts, hourly granularity must be enabled from the management account.'),
    ('Can SlashMyBill automatically stop my EC2 instances at night?',
     'Yes. Go to Act → Scheduler and create a stop/start schedule. SlashMyBill uses EventBridge '
     'Scheduler to automatically stop your instances at the configured time and start them again '
     'in the morning. You can pause or delete schedules at any time.'),
    ('Can I get expert help implementing the recommendations?',
     'Yes! SlashMyBill offers AI-driven consulting services. After analyzing your bill, click '
     '"Book a Free Consultation" to schedule a call with our FinOps experts. We implement all '
     'optimizations for you — you only pay 25% of the yearly savings we deliver.'),
    ('What is the Plan tab for?',
     'The Plan tab lets you create AWS Budgets with cost alerts and manage resource tags for '
     'cost allocation. Budgets are created directly in your AWS account so alerts come from AWS, '
     'not SlashMyBill.'),
    ('What is the FinOps Settings Healthcheck?',
     'FinOps Settings is a configuration audit in the Configure tab that checks your AWS account\'s '
     'billing best practices (cost allocation tags, anomaly detection, rightsizing, CUR reports, etc.) '
     'and lets you fix issues with one click. It detects whether your account is a management or linked '
     'account and shows the appropriate checklist. Your FinOps Score appears on the dashboard and in '
     'waste scan results.'),
    ('Why can\'t I reach 100% on the FinOps Score?',
     'All scored checks are achievable. The "Linked Account Billing Access" check is informational only '
     'and excluded from the score since it cannot be verified programmatically. If you see failing checks '
     'for Hourly Granularity or CUR Reports, these require manual setup in the AWS Console — they are '
     'scored but cannot be fixed via the SlashMyBill fix button.'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run('Q: ' + q)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 56, 100)
    doc.add_paragraph('A: ' + a)
    doc.add_paragraph()

doc.save('SlashMyBill-UserGuide-v4.docx')
print('User Guide saved: SlashMyBill-UserGuide-v4.docx')
