#!/usr/bin/env python3
"""
Made4Net Fortress & Factory - HLD Document Generator
Generates professional Word document for Sagi Van presentation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def create_hld_document():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ========================================
    # COVER PAGE
    # ========================================
    title = doc.add_heading('Made4Net Systems', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Fortress & Factory Architecture')
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(30, 60, 114)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    
    doc.add_paragraph()
    tagline = doc.add_paragraph()
    tagline_run = tagline.add_run('24/7 Operational Excellence & Security Posture')
    tagline_run.font.size = Pt(16)
    tagline_run.font.italic = True
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n')
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    info_data = [
        ('Document Type:', 'High-Level Design (HLD)'),
        ('Target Role:', 'Global Hosting Team Manager'),
        ('Prepared For:', 'Sagi Van - Made4Net Leadership'),
        ('Prepared By:', 'AWS Certified Security Specialist'),
        ('Date:', datetime.now().strftime('%B %d, %Y'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================
    # EXECUTIVE SUMMARY
    # ========================================
    doc.add_heading('Executive Summary', 1)
    
    exec_summary = [
        "The Made4Net Fortress & Factory architecture is designed specifically for 24/7 operational "
        "excellence and security posture management. This architecture addresses the core responsibilities "
        "of the Global Hosting Team Manager role, focusing on maintaining secure, compliant, and "
        "cost-efficient infrastructure for 800+ customer warehouses globally.",
        
        "This design draws from proven experience with the Israel Securities Authority and banking sector "
        "compliance requirements, implementing a 'Zero Trust' security model combined with automated "
        "maintenance strategies that eliminate manual intervention and reduce operational costs by 30%.",
        
        "The architecture is built on four foundational layers: Perimeter Security (Zero Trust), "
        "Automated Compute Maintenance, Resilient Data Layer, and Comprehensive Monitoring. Each layer "
        "incorporates AWS best practices and addresses specific challenges mentioned in the Made4Net "
        "hosting requirements."
    ]
    
    for para in exec_summary:
        p = doc.add_paragraph(para)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

    
    # ========================================
    # ARCHITECTURE OVERVIEW
    # ========================================
    doc.add_heading('1. Architecture Overview', 1)
    
    doc.add_heading('1.1 Design Philosophy: "Fortress & Factory"', 2)
    
    philosophy = [
        "The Fortress & Factory architecture embodies two critical principles:",
        
        "• FORTRESS: Impenetrable security perimeter with Zero Trust access controls, protecting "
        "sensitive warehouse operations data for global clients including Ingka/IKEA.",
        
        "• FACTORY: Automated, self-healing infrastructure that operates 24/7 without manual "
        "intervention, reducing operational overhead and human error."
    ]
    
    for para in philosophy:
        doc.add_paragraph(para)
    
    doc.add_heading('1.2 Key Design Objectives', 2)
    
    objectives_table = doc.add_table(rows=6, cols=2)
    objectives_table.style = 'Light List Accent 1'
    
    objectives_data = [
        ('Objective', 'Implementation'),
        ('Security', 'Zero Trust perimeter, encryption everywhere, continuous compliance monitoring'),
        ('Availability', '99.99% uptime with multi-region DR, automated failover'),
        ('Cost Efficiency', '30% cost reduction through automation and intelligent scheduling'),
        ('Compliance', 'Audit-ready with AWS Config, suitable for banking and retail audits'),
        ('Maintainability', 'Automated patching for mixed Linux/Windows fleets, zero-touch operations')
    ]
    
    for i, (obj, impl) in enumerate(objectives_data):
        objectives_table.rows[i].cells[0].text = obj
        objectives_table.rows[i].cells[1].text = impl
        if i == 0:
            objectives_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
            objectives_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # THE FOUR LAYERS
    # ========================================
    doc.add_heading('2. The Four Architecture Layers', 1)
    
    doc.add_paragraph(
        "The Fortress & Factory architecture is organized into four distinct layers, "
        "each addressing specific operational and security requirements."
    )
    
    # LAYER 1: PERIMETER
    doc.add_heading('2.1 Layer 1: The Perimeter - Zero Trust Access', 2)
    
    doc.add_paragraph(
        "Since Made4Net serves global clients across 800+ warehouses, the perimeter must be "
        "impenetrable yet maintain low-latency access for real-time warehouse operations."
    )
    
    perimeter_table = doc.add_table(rows=4, cols=3)
    perimeter_table.style = 'Medium Grid 3 Accent 1'
    
    perimeter_data = [
        ('Component', 'Role', 'Best Practice'),
        ('AWS WAF', 'Web Application Firewall on ALB', 'Geo-blocking rules, SQL injection & XSS protection'),
        ('Transit Gateway', 'Connects 800+ warehouses to VPCs', 'Network segmentation via route tables'),
        ('AWS Shield', 'DDoS Protection', 'Prevents volumetric attacks on platform')
    ]
    
    for i, row_data in enumerate(perimeter_data):
        for j, cell_text in enumerate(row_data):
            cell = perimeter_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        "Key Security Feature: Customer A's warehouse can never route packets to Customer B's "
        "hosted environment through Transit Gateway route table isolation."
    ).italic = True
    
    # LAYER 2: COMPUTE
    doc.add_heading('2.2 Layer 2: Compute - Automated No-Touch Maintenance', 2)
    
    doc.add_paragraph(
        "Managing Windows updates manually in a 24/7 environment is operationally risky. "
        "This layer eliminates manual intervention through AWS Systems Manager."
    )
    
    compute_table = doc.add_table(rows=4, cols=3)
    compute_table.style = 'Medium Grid 3 Accent 1'
    
    compute_data = [
        ('Component', 'Role', 'Best Practice'),
        ('AWS Systems Manager', 'Replaces SSH/RDP access', 'No port 22 or 3389 exposed to internet'),
        ('SSM Patch Manager', 'Automated patching', 'Critical updates during off-peak hours per region'),
        ('Golden AMIs', 'Immutable infrastructure', 'Bake patches into images, redeploy Auto Scaling Groups')
    ]
    
    for i, row_data in enumerate(compute_data):
        for j, cell_text in enumerate(row_data):
            cell = compute_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        "Interview Point: 'I utilize AWS Systems Manager to orchestrate rolling patches. "
        "Production is only patched after non-production passes health checks.'"
    ).italic = True

    
    doc.add_page_break()
    
    # LAYER 3: DATA
    doc.add_heading('2.3 Layer 3: Data - Resilience & Isolation', 2)
    
    doc.add_paragraph(
        "This layer addresses high availability requirements with encryption, automated backups, "
        "and disaster recovery capabilities."
    )
    
    data_table = doc.add_table(rows=4, cols=3)
    data_table.style = 'Medium Grid 3 Accent 1'
    
    data_data = [
        ('Component', 'Role', 'Best Practice'),
        ('AWS KMS', 'Key Management Service', 'Encrypt EBS volumes and RDS at rest'),
        ('TLS 1.2+', 'Data in transit encryption', 'All warehouse scanner to cloud communication'),
        ('AWS Backup + Cross-Region', 'Disaster Recovery', 'Pilot Light in secondary region (us-west-2)')
    ]
    
    for i, row_data in enumerate(data_data):
        for j, cell_text in enumerate(row_data):
            cell = data_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        "DR Strategy: If primary region (us-east-1) fails, Pilot Light copy in us-west-2 "
        "can be activated within 15 minutes."
    ).italic = True
    
    # LAYER 4: MONITORING
    doc.add_heading('2.4 Layer 4: Monitoring - Eyes on Glass', 2)
    
    doc.add_paragraph(
        "24/7 operations require continuous visibility. This layer provides intelligent threat "
        "detection and proactive health monitoring."
    )
    
    monitoring_table = doc.add_table(rows=4, cols=3)
    monitoring_table.style = 'Medium Grid 3 Accent 1'
    
    monitoring_data = [
        ('Component', 'Role', 'Best Practice'),
        ('Amazon GuardDuty', 'Intelligent threat detection', 'Spots bitcoin mining, malicious IPs'),
        ('CloudWatch Canaries', 'Synthetic monitoring', 'Simulate user transactions every minute'),
        ('AWS X-Ray', 'Distributed tracing', 'Trace latency: Wi-Fi, VPN, or SQL query?')
    ]
    
    for i, row_data in enumerate(monitoring_data):
        for j, cell_text in enumerate(row_data):
            cell = monitoring_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        "Proactive Alerting: Canaries page on-call team via PagerDuty/OpsGenie before "
        "warehouse customers call helpdesk."
    ).italic = True
    
    doc.add_page_break()

    
    # ========================================
    # INTERVIEW TALKING POINTS
    # ========================================
    doc.add_heading('3. Interview Talking Points for Sagi Van', 1)
    
    doc.add_paragraph(
        "These talking points directly address anticipated concerns from Made4Net leadership, "
        "bridging technical architecture to business value."
    )
    
    # Talking Point 1
    doc.add_heading('3.1 Patching Mixed OS Fleets (Linux/Windows)', 2)
    
    tp1_table = doc.add_table(rows=3, cols=2)
    tp1_table.style = 'Light Shading Accent 1'
    
    tp1_data = [
        ('Sagi\'s Concern', 'Your Response'),
        ('"How do we patch Windows servers without downtime?"',
         '"In my experience, manual patching is a risk. I utilize AWS Systems Manager to orchestrate '
         'rolling patches. For complex environments, I group instances by Patch Baselines—ensuring '
         'production is only patched after non-production passes health checks. This achieves 95%+ '
         'compliance with zero downtime."'),
        ('Business Impact', '• Zero downtime patching\n• 95%+ compliance rate\n• Reduced security vulnerabilities\n'
         '• Automated rollback on failure')
    ]
    
    for i, (label, value) in enumerate(tp1_data):
        tp1_table.rows[i].cells[0].text = label
        tp1_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in tp1_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Talking Point 2
    doc.add_heading('3.2 Cost Efficiency vs. Availability', 2)
    
    tp2_table = doc.add_table(rows=3, cols=2)
    tp2_table.style = 'Light Shading Accent 1'
    
    tp2_data = [
        ('Sagi\'s Concern', 'Your Response'),
        ('"Our AWS bill is too high."',
         '"I managed this at the Israel Securities Authority where I reduced costs by 30%. My strategy '
         'is Trusted Advisor + Automation. I use AWS Instance Scheduler to automatically stop '
         'non-production environments on nights and weekends. We don\'t pay for what we don\'t use."'),
        ('Business Impact', '• 30% cost reduction achieved\n• $15k-$20k monthly savings typical\n'
         '• Zero impact on production availability\n• Automated resource optimization')
    ]
    
    for i, (label, value) in enumerate(tp2_data):
        tp2_table.rows[i].cells[0].text = label
        tp2_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in tp2_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()

    
    # Talking Point 3
    doc.add_heading('3.3 Incident Resolution & Observability', 2)
    
    tp3_table = doc.add_table(rows=3, cols=2)
    tp3_table.style = 'Light Shading Accent 1'
    
    tp3_data = [
        ('Sagi\'s Concern', 'Your Response'),
        ('"What happens when the system is slow?"',
         '"I focus on Observability. At Shekel Brainweigh, I monitored 1,000+ IoT devices. I implement '
         'X-Ray or APM tools to trace latency. Is it the Warehouse Wi-Fi? The VPN? Or a slow SQL query? '
         'We need data to prove innocence or find the root cause quickly. Target MTTR: 8-15 minutes."'),
        ('Business Impact', '• Mean Time to Resolution: 8-15 minutes\n• Root cause identification, not guessing\n'
         '• Reduced customer complaints\n• Data-driven incident response')
    ]
    
    for i, (label, value) in enumerate(tp3_data):
        tp3_table.rows[i].cells[0].text = label
        tp3_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in tp3_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Talking Point 4
    doc.add_heading('3.4 Security Compliance for Banking & Retail', 2)
    
    tp4_table = doc.add_table(rows=3, cols=2)
    tp4_table.style = 'Light Shading Accent 1'
    
    tp4_data = [
        ('Sagi\'s Concern', 'Your Response'),
        ('"Our customers are banks and large retailers; they audit us."',
         '"I have architected compliant solutions for the banking sector (Bank Leumi). I use AWS Config '
         'to continuously record configuration changes. If an auditor asks, \'Who changed this Security '
         'Group last Tuesday?\', we have the log instantly. Compliance score: 95-100/100."'),
        ('Business Impact', '• Audit-ready at all times\n• Complete change history\n'
         '• 95-100 compliance score\n• Reduced audit preparation time')
    ]
    
    for i, (label, value) in enumerate(tp4_data):
        tp4_table.rows[i].cells[0].text = label
        tp4_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in tp4_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # TECHNICAL SPECIFICATIONS
    # ========================================
    doc.add_heading('4. Technical Specifications', 1)
    
    doc.add_heading('4.1 AWS Services Utilized', 2)
    
    services_table = doc.add_table(rows=13, cols=3)
    services_table.style = 'Medium List 1 Accent 1'
    
    services_data = [
        ('Service', 'Purpose', 'Configuration'),
        ('AWS WAF', 'Web Application Firewall', 'SQL injection, XSS, geo-blocking rules'),
        ('Transit Gateway', 'Network hub', 'Route tables for customer isolation'),
        ('AWS Shield', 'DDoS protection', 'Standard (automatic)'),
        ('Systems Manager', 'Server management', 'Patch Manager, State Manager, Session Manager'),
        ('EC2 Auto Scaling', 'Compute elasticity', 'Golden AMIs, health checks'),
        ('AWS KMS', 'Encryption key management', 'Customer-managed keys for compliance'),
        ('RDS', 'Managed database', 'Multi-AZ, encrypted, automated backups'),
        ('AWS Backup', 'Centralized backup', 'Cross-region replication'),
        ('GuardDuty', 'Threat detection', 'Continuous monitoring, ML-based'),
        ('CloudWatch', 'Monitoring & logging', 'Canaries, alarms, dashboards'),
        ('AWS X-Ray', 'Distributed tracing', 'Latency analysis, bottleneck identification'),
        ('AWS Config', 'Configuration tracking', 'Compliance rules, change history')
    ]
    
    for i, row_data in enumerate(services_data):
        for j, cell_text in enumerate(row_data):
            cell = services_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================
    # PERFORMANCE METRICS
    # ========================================
    doc.add_heading('4.2 Target Performance Metrics', 2)
    
    metrics_table = doc.add_table(rows=9, cols=3)
    metrics_table.style = 'Light Grid Accent 1'
    
    metrics_data = [
        ('Metric', 'Target', 'Measurement Method'),
        ('System Availability', '99.99%', 'CloudWatch Canaries + uptime monitoring'),
        ('Patch Compliance', '95%+', 'SSM Patch Manager compliance reports'),
        ('Mean Time to Resolution', '8-15 minutes', 'Incident tracking system'),
        ('Security Compliance Score', '95-100/100', 'AWS Config compliance dashboard'),
        ('Cost Reduction', '30%', 'AWS Cost Explorer + Trusted Advisor'),
        ('Cross-Region Replication Lag', '<150ms', 'CloudWatch metrics'),
        ('Failed Authentication Attempts', '<10/day', 'CloudWatch Logs Insights'),
        ('WAF Blocked Requests', '100-150/day', 'WAF metrics dashboard')
    ]
    
    for i, row_data in enumerate(metrics_data):
        for j, cell_text in enumerate(row_data):
            cell = metrics_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # COST ANALYSIS
    # ========================================
    doc.add_heading('5. Cost Analysis & Optimization', 1)
    
    doc.add_heading('5.1 Cost Optimization Strategies', 2)
    
    cost_strategies = [
        "1. AWS Instance Scheduler: Automatically stop/start non-production environments during "
        "off-hours (nights, weekends). Typical savings: $7,000-$10,000/month.",
        
        "2. Trusted Advisor Recommendations: Weekly review of idle resources, underutilized instances, "
        "and unattached EBS volumes. Typical savings: $3,000-$5,000/month.",
        
        "3. Reserved Instances: 1-year commitment for production workloads with predictable usage. "
        "Typical savings: $5,000-$8,000/month.",
        
        "4. S3 Lifecycle Policies: Automatic transition to Glacier for archival data. "
        "Typical savings: $1,000-$2,000/month."
    ]
    
    for strategy in cost_strategies:
        doc.add_paragraph(strategy)
    
    doc.add_paragraph()
    
    doc.add_heading('5.2 Monthly Cost Breakdown (Estimated)', 2)
    
    cost_table = doc.add_table(rows=8, cols=3)
    cost_table.style = 'Medium Shading 1 Accent 1'
    
    cost_data = [
        ('Service Category', 'Before Optimization', 'After Optimization'),
        ('Compute (EC2, Auto Scaling)', '$35,000', '$24,500'),
        ('Database (RDS)', '$12,000', '$9,000'),
        ('Storage (S3, EBS)', '$8,000', '$5,600'),
        ('Networking (Transit Gateway, VPN)', '$6,000', '$6,000'),
        ('Security & Monitoring', '$4,000', '$4,000'),
        ('Total Monthly Cost', '$65,000', '$49,100'),
        ('Monthly Savings', '-', '$15,900 (30% reduction)')
    ]
    
    for i, row_data in enumerate(cost_data):
        for j, cell_text in enumerate(row_data):
            cell = cost_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0 or i == 6 or i == 7:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()

    
    # ========================================
    # DISASTER RECOVERY
    # ========================================
    doc.add_heading('6. Disaster Recovery Strategy', 1)
    
    doc.add_heading('6.1 Multi-Region Architecture', 2)
    
    dr_overview = [
        "The Fortress & Factory architecture implements a Pilot Light disaster recovery strategy "
        "across two AWS regions:",
        
        "• Primary Region: us-east-1 (N. Virginia) - Active production workloads",
        "• DR Region: us-west-2 (Oregon) - Standby with minimal resources",
        
        "In the event of a regional failure, the DR region can be activated within 15 minutes, "
        "meeting the Recovery Time Objective (RTO) of <30 minutes."
    ]
    
    for para in dr_overview:
        doc.add_paragraph(para)
    
    doc.add_paragraph()
    
    doc.add_heading('6.2 DR Metrics', 2)
    
    dr_table = doc.add_table(rows=5, cols=2)
    dr_table.style = 'Light List Accent 1'
    
    dr_data = [
        ('Metric', 'Target'),
        ('Recovery Time Objective (RTO)', '<30 minutes'),
        ('Recovery Point Objective (RPO)', '<5 minutes'),
        ('Cross-Region Replication Lag', '<150ms'),
        ('Failover Test Frequency', 'Quarterly')
    ]
    
    for i, (label, value) in enumerate(dr_data):
        dr_table.rows[i].cells[0].text = label
        dr_table.rows[i].cells[1].text = value
        if i == 0:
            for cell in dr_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_page_break()
    
    # ========================================
    # DIAGRAM PLACEHOLDER
    # ========================================
    doc.add_heading('7. Architecture Diagrams', 1)
    
    doc.add_paragraph(
        "The following diagrams illustrate the Fortress & Factory architecture. "
        "These diagrams will be generated separately using draw.io format."
    )
    
    doc.add_paragraph()
    
    diagram_list = [
        "1. Overall Architecture Diagram - Shows all four layers and AWS services",
        "2. Security Flow Diagram - Zero Trust perimeter and WAF rules",
        "3. Patch Management Flow - SSM automation workflow",
        "4. Disaster Recovery Flow - Multi-region failover process",
        "5. Monitoring Dashboard - CloudWatch Canaries and GuardDuty integration"
    ]
    
    for diagram in diagram_list:
        p = doc.add_paragraph(diagram)
        p.style = 'List Bullet'
    
    doc.add_paragraph()
    
    placeholder = doc.add_paragraph()
    placeholder_run = placeholder.add_run('[DIAGRAM PLACEHOLDERS - Insert draw.io exports here]')
    placeholder_run.font.color.rgb = RGBColor(255, 0, 0)
    placeholder_run.font.size = Pt(14)
    placeholder_run.font.bold = True
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    
    # ========================================
    # IMPLEMENTATION ROADMAP
    # ========================================
    doc.add_heading('8. Implementation Roadmap', 1)
    
    doc.add_paragraph(
        "The following phased approach ensures minimal disruption to existing operations while "
        "implementing the Fortress & Factory architecture."
    )
    
    doc.add_paragraph()
    
    roadmap_table = doc.add_table(rows=5, cols=4)
    roadmap_table.style = 'Medium Grid 1 Accent 1'
    
    roadmap_data = [
        ('Phase', 'Duration', 'Key Activities', 'Success Criteria'),
        ('Phase 1: Foundation', '2-3 weeks',
         '• Deploy Transit Gateway\n• Configure WAF rules\n• Enable GuardDuty',
         '• Zero Trust perimeter active\n• Threat detection enabled'),
        ('Phase 2: Automation', '3-4 weeks',
         '• Implement SSM Patch Manager\n• Create Golden AMIs\n• Configure Auto Scaling',
         '• 95%+ patch compliance\n• Zero-touch patching'),
        ('Phase 3: DR & Backup', '2-3 weeks',
         '• Configure cross-region replication\n• Set up AWS Backup\n• Test failover',
         '• RTO <30 min achieved\n• Successful DR test'),
        ('Phase 4: Optimization', '2-3 weeks',
         '• Implement Instance Scheduler\n• Configure Trusted Advisor\n• Cost analysis',
         '• 30% cost reduction\n• Automated optimization')
    ]
    
    for i, row_data in enumerate(roadmap_data):
        for j, cell_text in enumerate(row_data):
            cell = roadmap_table.rows[i].cells[j]
            cell.text = cell_text
            if i == 0:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    total_timeline = doc.add_paragraph()
    total_timeline_run = total_timeline.add_run('Total Implementation Timeline: 9-13 weeks')
    total_timeline_run.font.bold = True
    total_timeline_run.font.size = Pt(12)
    total_timeline_run.font.color.rgb = RGBColor(30, 60, 114)
    
    doc.add_page_break()
    
    # ========================================
    # CONCLUSION
    # ========================================
    doc.add_heading('9. Conclusion', 1)
    
    conclusion = [
        "The Fortress & Factory architecture provides Made4Net with a secure, automated, and "
        "cost-efficient infrastructure platform that meets the demanding requirements of 24/7 "
        "warehouse operations for global clients.",
        
        "Key benefits delivered:",
        
        "• Security: Zero Trust perimeter with continuous threat detection and compliance monitoring",
        "• Availability: 99.99% uptime with multi-region disaster recovery",
        "• Cost Efficiency: 30% cost reduction through intelligent automation",
        "• Maintainability: Automated patching for mixed OS fleets with zero downtime",
        "• Compliance: Audit-ready infrastructure suitable for banking and retail customers",
        
        "This architecture draws from proven experience with the Israel Securities Authority and "
        "banking sector compliance, ensuring Made4Net can confidently serve its most demanding "
        "customers while maintaining operational excellence."
    ]
    
    for para in conclusion:
        if para.startswith('•'):
            p = doc.add_paragraph(para)
            p.style = 'List Bullet'
        else:
            p = doc.add_paragraph(para)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    doc.add_page_break()

    
    # ========================================
    # APPENDIX
    # ========================================
    doc.add_heading('Appendix A: Glossary', 1)
    
    glossary_table = doc.add_table(rows=11, cols=2)
    glossary_table.style = 'Light List Accent 1'
    
    glossary_data = [
        ('Term', 'Definition'),
        ('Zero Trust', 'Security model that requires verification for every access request'),
        ('Golden AMI', 'Pre-configured machine image with security patches and configurations'),
        ('Pilot Light', 'DR strategy with minimal resources running, ready to scale up'),
        ('MTTR', 'Mean Time to Resolution - average time to resolve incidents'),
        ('RTO', 'Recovery Time Objective - maximum acceptable downtime'),
        ('RPO', 'Recovery Point Objective - maximum acceptable data loss'),
        ('SSM', 'AWS Systems Manager - unified interface for managing AWS resources'),
        ('WAF', 'Web Application Firewall - protects against common web exploits'),
        ('GuardDuty', 'AWS threat detection service using machine learning'),
        ('Canary', 'Synthetic monitoring script that simulates user behavior')
    ]
    
    for i, (term, definition) in enumerate(glossary_data):
        glossary_table.rows[i].cells[0].text = term
        glossary_table.rows[i].cells[1].text = definition
        if i == 0:
            for cell in glossary_table.rows[i].cells:
                cell.paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading('Appendix B: References', 1)
    
    references = [
        "1. AWS Well-Architected Framework - Security Pillar",
        "2. AWS Well-Architected Framework - Operational Excellence Pillar",
        "3. AWS Systems Manager Best Practices",
        "4. AWS Disaster Recovery Whitepaper",
        "5. AWS Cost Optimization Best Practices",
        "6. Made4Net Global Hosting Team Manager Job Description"
    ]
    
    for ref in references:
        p = doc.add_paragraph(ref)
        p.style = 'List Number'
    
    # Save document
    doc.save('Made4Net-Fortress-Factory-HLD.docx')
    print("✓ HLD document generated: Made4Net-Fortress-Factory-HLD.docx")
    print(f"  Document size: {len(doc.element.xml)} bytes")
    print("  Sections: 9 main sections + 2 appendices")
    print("  Ready for Sagi Van presentation")

if __name__ == '__main__':
    create_hld_document()
