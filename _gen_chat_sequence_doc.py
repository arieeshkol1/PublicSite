#!/usr/bin/env python3
"""
Generate SlashMyBill-AI-Chat-Sequence-Flow.docx
Describes the end-to-end question processing sequence when a user asks a question in the Chat tab.
Uses python-docx for proper Word formatting with tables, bullet lists, heading hierarchy.
"""

import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_table_row(table, cells, header=False):
    """Add a row to a table with optional header styling."""
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = str(text)
        if header:
            set_cell_shading(row.cells[i], "2E4057")
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.size = Pt(9)
        else:
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return row

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_numbered(doc, text):
    """Add a numbered list item."""
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_body(doc, text):
    """Add body text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


# ── Main document builder ────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Style tweaks ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # ── Title page ──────────────────────────────────────────────────────────
    title = doc.add_heading('SlashMyCloudBill', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_heading('AI Chat — Question Processing Sequence Flow', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Version 1.0  |  June 2026\n').font.size = Pt(12)
    meta.add_run('Platform: slashmycloudbill.com\n').font.size = Pt(11)
    meta.add_run('AWS Region: us-east-1').font.size = Pt(11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 1. OVERVIEW
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('1. Overview', level=1)

    add_body(doc,
        'When a user types a question in the SlashMyBill Chat tab and clicks Send, the system '
        'executes a multi-stage pipeline that classifies the question, gathers relevant cloud data, '
        'searches the knowledge base for tips, and generates an AI-powered response using Amazon '
        'Bedrock Nova.')

    add_body(doc,
        'This document describes the complete end-to-end sequence from the moment the user '
        'submits a question until the response is rendered in the browser.')

    # ════════════════════════════════════════════════════════════════════════
    # 2. END-TO-END SEQUENCE DIAGRAM
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('2. End-to-End Sequence Diagram', level=1)

    add_body(doc, 'High-level data flow:')
    add_body(doc, 'User Browser → API Gateway → Member-Handler Lambda → [Pipeline] → Bedrock Nova → Response')

    doc.add_heading('2.1 Actors', level=2)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Actor', 'Role']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    actors = [
        ('User Browser', 'Frontend (members.js) sends POST request'),
        ('API Gateway', 'HTTP API (l2fd4h481h.execute-api.us-east-1.amazonaws.com)'),
        ('Member-Handler Lambda', 'Main processing Lambda (aws-bill-analyzer-member-api)'),
        ('DynamoDB', 'Members table, Accounts table, Cost Cache table, Tips table'),
        ('AWS STS', 'Cross-account role assumption'),
        ('AWS Cost Explorer', 'Cost and usage data'),
        ('AWS CloudWatch', 'Resource metrics (CPU, invocations, etc.)'),
        ('AWS EC2/RDS/Lambda/S3', 'Resource inventory APIs'),
        ('Bedrock Nova', 'AI model (us.amazon.nova-2-lite-v1:0)'),
        ('Transaction Logger', 'Audit trail to Audit_Transaction_Log table'),
    ]
    for row_data in actors:
        add_table_row(tbl, row_data)

    # ════════════════════════════════════════════════════════════════════════
    # 3. DETAILED STEP-BY-STEP FLOW
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('3. Detailed Step-by-Step Flow', level=1)

    # Step 1
    doc.add_heading('Step 1: Frontend Sends Request', level=2)
    add_bullet(doc, 'Endpoint: POST /members/accounts/ai-query')
    add_bullet(doc, 'Payload: { accountIds: ["714045115933"], question: "..." }')
    add_bullet(doc, 'Headers: Bearer token (Cognito access token)')
    add_bullet(doc, 'Timeout: Frontend waits up to 30 seconds')

    # Step 2
    doc.add_heading('Step 2: Authentication (validate_token)', level=2)
    add_numbered(doc, 'Extract Bearer token from Authorization header')
    add_numbered(doc, 'Call Cognito GetUser with the access token')
    add_numbered(doc, 'Extract member email from user attributes')
    add_numbered(doc, 'If Cognito fails, fall back to legacy JWT validation')

    # Step 3
    doc.add_heading('Step 3: Credit Check (_check_and_consume_credits)', level=2)
    add_numbered(doc, 'Read member record from DynamoDB (MemberPortal-Members)')
    add_numbered(doc, 'Check aiCreditsUsed vs tier limit (Free: 100, Growth: 300, Scale: 1500)')
    add_numbered(doc, 'If insufficient credits → return 403 with token info')
    add_numbered(doc, 'Consume 2 credits (AI_QUERY_CREDIT_COST) atomically')

    # Step 4
    doc.add_heading('Step 4: Account Ownership Verification', level=2)
    add_numbered(doc, 'Query MemberPortal-Accounts by memberEmail')
    add_numbered(doc, 'Verify all requested accountIds belong to the authenticated user')
    add_numbered(doc, 'If lateral access detected → return 403')

    # Step 5
    doc.add_heading('Step 5: Route Selection', level=2)
    add_bullet(doc, 'Multi-account (>1 accountId): → _invoke_multi_account()')
    add_bullet(doc, 'Single account: → _invoke_direct_model()')

    # Step 6
    doc.add_heading('Step 6: Intent Classification (_classify_intent)', level=2)
    add_bullet(doc, 'Engine: Keyword-based pattern matching (no LLM, <50ms)')
    add_bullet(doc, 'Categories: ec2, rds, s3, lambda, commitments, cost-general, network, storage, compute')
    add_bullet(doc, 'Output: Set of intent categories that control which APIs to call')
    add_bullet(doc, 'Example: "what about RIs and SPs?" → {commitments, cost-general}')

    # Step 7
    doc.add_heading('Step 7: Provider Detection (_route_to_connector)', level=2)
    add_numbered(doc, 'Read account record from DynamoDB')
    add_numbered(doc, 'Detect cloud provider (aws/azure/gcp) from cloudProvider field')
    add_numbered(doc, 'Return provider type and credential config')

    # Step 8
    doc.add_heading('Step 8: Tips Search (_search_tips)', level=2)
    add_numbered(doc, 'Query ViewMyBill-CostOptimizationTips DynamoDB table')
    add_numbered(doc, 'Match question keywords against tip titles/descriptions')
    add_numbered(doc, 'Filter by provider (aws/azure/gcp)')
    add_numbered(doc, 'Return top 3 relevant tips for prompt context')

    # Step 9
    doc.add_heading('Step 9: Data Gathering (_gather_account_data)', level=2)
    add_bullet(doc, 'Time Budget: 14 seconds maximum (leaves 12s for Bedrock)')
    add_bullet(doc, 'Time Guard: _time_left() function checked before each slow API')

    doc.add_heading('9a. STS AssumeRole', level=3)
    add_numbered(doc, 'Compute role ARN: arn:aws:iam::{accountId}:role/SlashMyBill-{accountId}')
    add_numbered(doc, 'Compute ExternalId: SHA-256 of member email')
    add_numbered(doc, 'Call sts:AssumeRole → get temporary credentials')

    doc.add_heading('9b. Cost Explorer (always runs first — fast with cache)', level=3)
    add_numbered(doc, 'Cache Check: Query Cost_Cache_Table for {email}#{accountId} with DAILY# prefix')
    add_numbered(doc, 'Cache HIT (25+ days): Use cached service breakdown, skip live CE')
    add_numbered(doc, 'Cache MISS: Call ce:GetCostAndUsage (monthly by service, last 30 days)')
    add_numbered(doc, 'Daily Trend: Call ce:GetCostAndUsage (daily, last 7 days)')
    add_numbered(doc, 'Usage Breakdown: For top services matching the question, fetch per-usage-type breakdown')

    doc.add_heading('9c. Monthly Cost Forecast (computed from daily data)', level=3)
    add_bullet(doc, 'Formula: (7-day avg daily cost × days in current month) + recurring fees')
    add_bullet(doc, 'Per-service forecast included in cost_forecast.by_service')
    add_bullet(doc, 'Recurring fees identified: Tax, Support, Registrar')

    doc.add_heading('9d. Conditional API Calls (based on intent + time budget)', level=3)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['API', 'Condition', 'Time Gate']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    api_calls = [
        ('EC2 DescribeInstances', 'Intent includes ec2_describe_instances', '_time_left() > 5'),
        ('RDS DescribeDBInstances', 'Intent includes rds_describe_instances', '_time_left() > 3'),
        ('Lambda ListFunctions', 'Intent includes lambda_list_functions', 'Always (fast)'),
        ('S3 ListBuckets', 'Intent includes s3_list_buckets', '_time_left() > 2'),
        ('NAT Gateways', 'Intent includes nat_gateways', '_time_left() > 3'),
        ('SP/RI Coverage', 'Intent includes sp_ri_coverage', '_time_left() > 3'),
        ('Pricing API', 'Always (last)', '_time_left() > 4'),
    ]
    for row_data in api_calls:
        add_table_row(tbl, row_data)

    doc.add_heading('9e. CloudWatch Metrics (per discovered resources)', level=3)
    add_bullet(doc, 'Lambda: invocations, duration, errors per function (30 days)')
    add_bullet(doc, 'EC2: CPU utilization per instance (30 days)')
    add_bullet(doc, 'RDS: CPU, connections per DB (30 days)')

    doc.add_heading('9f. Cost Efficiency Calculation', level=3)
    add_bullet(doc, 'Compute potential savings from: EBS waste, idle EIPs, over-provisioned instances')
    add_bullet(doc, 'Calculate efficiency score: (1 - savings/total) × 100')

    # Step 10
    doc.add_heading('Step 10: Prompt Assembly', level=2)
    add_numbered(doc, 'Serialize gathered data as JSON (trimmed to ~10KB max)')
    add_numbered(doc, 'Build prompt with system rules (anti-hallucination, pricing, platform features)')
    add_numbered(doc, 'Include tip citations from Step 8')
    add_numbered(doc, 'Include user question')
    add_numbered(doc, 'Include account data from Step 9')

    # Step 11
    doc.add_heading('Step 11: Bedrock Model Invocation', level=2)
    add_bullet(doc, 'Model: us.amazon.nova-2-lite-v1:0')
    add_bullet(doc, 'Max Tokens: 3000')
    add_bullet(doc, 'Temperature: 0.3')
    add_bullet(doc, 'Timeout: 27 seconds (enforced by ThreadPoolExecutor)')
    add_bullet(doc, 'If timeout: Return generic "analysis taking longer" message')

    # Step 12
    doc.add_heading('Step 12: Response Assembly', level=2)
    add_numbered(doc, 'Parse model output text')
    add_numbered(doc, 'Build chart data (service costs bar, daily trend line, efficiency doughnut)')
    add_numbered(doc, 'Include: answer, interactionId, commands list, tips found, chart data, top services')
    add_numbered(doc, 'Inject AI credits remaining')

    # Step 13
    doc.add_heading('Step 13: Transaction Logging (@transaction_log decorator)', level=2)
    add_numbered(doc, 'Capture request/response payloads (sanitized, truncated to 10KB)')
    add_numbered(doc, 'Record duration, status, user email, function name')
    add_numbered(doc, 'Write to Audit_Transaction_Log DynamoDB table')
    add_numbered(doc, 'Conditional write (attribute_not_exists) prevents duplicates')

    # Step 14
    doc.add_heading('Step 14: Response to Client', level=2)
    add_bullet(doc, 'Status: 200 (always, even for errors — error in body)')
    add_bullet(doc, 'Headers: CORS enabled')
    add_bullet(doc, 'Body: JSON with answer, charts, commands, credits')

    # ════════════════════════════════════════════════════════════════════════
    # 4. MULTI-ACCOUNT FLOW DIFFERENCES
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('4. Multi-Account Flow Differences', level=1)

    add_body(doc, 'When accountIds.length > 1:')
    add_numbered(doc, 'Detect provider per account via _route_to_connector')
    add_numbered(doc, 'Process all accounts concurrently (ThreadPoolExecutor, max 3 workers)')
    add_numbered(doc, 'Each account runs Steps 9a-9f independently')
    add_numbered(doc, 'Merge results: aggregate cost_by_service, daily trends, monthly trends')
    add_numbered(doc, 'Use multi-account prompt with per-account breakdown')
    add_numbered(doc, 'Build aggregate charts')

    # ════════════════════════════════════════════════════════════════════════
    # 5. PERFORMANCE CHARACTERISTICS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('5. Performance Characteristics', level=1)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Metric', 'Target', 'Actual (typical)']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    perf_data = [
        ('Cache HIT + simple question', '<8s', '5-12s'),
        ('Cache MISS (full CE fetch)', '<15s', '15-25s'),
        ('Multi-account (2 accounts)', '<15s', '10-20s'),
        ('Timeout threshold', '27s', 'Lambda returns fallback'),
        ('Bedrock model latency', '3-8s', '3-12s'),
    ]
    for row_data in perf_data:
        add_table_row(tbl, row_data)

    # ════════════════════════════════════════════════════════════════════════
    # 6. ERROR HANDLING
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('6. Error Handling', level=1)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Error', 'Handling']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    errors = [
        ('Auth failure', '401 returned immediately'),
        ('Account not owned', '403 returned'),
        ('STS AssumeRole fails', '403 with guidance to redeploy CFN'),
        ('API timeout (27s)', 'Generic "taking longer" message'),
        ('Bedrock model error', '"AI analysis error" with error detail'),
        ('Individual API fails', 'Silently skipped, data gathered from other sources'),
        ('Transaction logger fails', 'Swallowed (never affects response)'),
    ]
    for row_data in errors:
        add_table_row(tbl, row_data)

    # ════════════════════════════════════════════════════════════════════════
    # 7. KEY CONFIGURATION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('7. Key Configuration', level=1)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Setting', 'Value']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    config = [
        ('Lambda function', 'aws-bill-analyzer-member-api'),
        ('API Gateway', 'l2fd4h481h (HTTP API, $default stage)'),
        ('Bedrock Model', 'us.amazon.nova-2-lite-v1:0'),
        ('Region', 'us-east-1'),
        ('Lambda timeout', '30s'),
        ('API Gateway timeout', '29s'),
        ('Code timeout', '27s (ThreadPoolExecutor)'),
        ('Max gather time', '14s'),
        ('Credit cost', '2 per query'),
        ('Max response tokens', '3000'),
    ]
    for row_data in config:
        add_table_row(tbl, row_data)

    # ════════════════════════════════════════════════════════════════════════
    # 8. FILES INVOLVED
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('8. Files Involved', level=1)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['File', 'Role']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    files = [
        ('members/members.js', 'Frontend: sends question, renders response'),
        ('member-handler/lambda_function.py', 'Main handler: routing, data gathering, prompt'),
        ('member-handler/intent_classifier.py', 'Question → intent category mapping'),
        ('member-handler/sts_assume_role.py', 'Cross-account credential handling'),
        ('member-handler/provider_registry.py', 'Multi-cloud provider config'),
        ('member-handler/cache_service.py', 'DynamoDB cost cache read/write'),
        ('member-handler/tip_citation.py', 'Tips formatting for prompt'),
        ('transaction_logger.py', 'Audit logging decorator'),
        ('knowledge-base/aws-cost-optimization-tips.json', 'Tips knowledge base'),
    ]
    for row_data in files:
        add_table_row(tbl, row_data)

    # ── Save ────────────────────────────────────────────────────────────────
    doc.save('SlashMyBill-AI-Chat-Sequence-Flow.docx')
    print('SlashMyBill-AI-Chat-Sequence-Flow.docx generated successfully!')


# ── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    build_document()
