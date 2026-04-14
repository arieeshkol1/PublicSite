#!/usr/bin/env python3
"""
Generate SlashMyBill-HLD.docx - High-Level Design document for SlashMyCloudBill platform.
Uses python-docx for proper Word formatting with tables, bullet lists, heading hierarchy.
"""

import sys
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear',
    })
    shading.append(shd)

def add_table_row(table, cells, header=False):
    """Add a row to a table with optional header styling."""
    row = table.add_row()
    for i, text in enumerate(cells):
        row.cells[i].text = str(text)
        if header:
            set_cell_shading(row.cells[i], "2E4057")
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
                    run.font.size = Pt(9)
        else:
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return row

def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 0.63)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_numbered(doc, text):
    """Add a numbered list item."""
    p = doc.add_paragraph(text, style='List Number')
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)
    return p

def add_body(doc, text):
    """Add body text."""
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


# ── Main document builder ────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Style tweaks ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    # ── Title page ──────────────────────────────────────────────────────────
    title = doc.add_heading('SlashMyCloudBill', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x2E, 0x40, 0x57)

    subtitle = doc.add_heading('High-Level Design Document', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run('Version 2.0  |  April 2026\n').font.size = Pt(12)
    meta.add_run('Platform: slashmycloudbill.com\n').font.size = Pt(11)
    meta.add_run('AWS Region: us-east-1').font.size = Pt(11)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('1. Executive Summary', level=1)

    add_body(doc,
        'SlashMyCloudBill is a SaaS FinOps platform that helps organizations analyze, '
        'optimize, and reduce their Amazon Web Services cloud spending. The platform '
        'combines AI-powered bill analysis with automated cost optimization actions to '
        'deliver measurable savings.')

    add_body(doc,
        'Target Market: Small-to-medium businesses and startups running workloads on AWS '
        'who lack dedicated FinOps teams. The platform serves both self-service users '
        '(via tiered subscriptions) and managed-service clients (25% of yearly savings model).')

    doc.add_heading('Value Proposition', level=2)
    add_bullet(doc, 'Instant bill analysis — upload a PDF and get AI-generated optimization report in under 60 seconds')
    add_bullet(doc, 'Multi-account dashboard — unified view across all linked AWS accounts with 8 customizable widgets')
    add_bullet(doc, 'One-click cleanup actions — safely remove waste (unused EIPs, detached EBS, idle instances) with JIT safety checks')
    add_bullet(doc, 'AI-powered chat — ask natural-language questions about your cloud spend, powered by Amazon Bedrock Nova 2 Lite')
    add_bullet(doc, 'Zero credential storage — cross-account IAM roles with read-only access; no keys stored')

    # ════════════════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('2. System Architecture', level=1)

    add_body(doc,
        'The platform is built entirely on AWS serverless services in the us-east-1 region. '
        'The architecture follows an event-driven, API-first design with no persistent servers.')

    doc.add_heading('2.1 Component Inventory', level=2)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Component', 'Service', 'Identifier / ARN', 'Region']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    components = [
        ('DNS', 'Route 53', 'slashmycloudbill.com, eshkolai.com hosted zones', 'us-east-1'),
        ('CDN (SMB)', 'CloudFront', 'E2B3GXE4TJTH4Q', 'Global'),
        ('CDN (eshkolai)', 'CloudFront', 'E12JIHGHK40OLE', 'Global'),
        ('CF Function', 'CloudFront Functions', 'SlashMyCloudBill-Router', 'Global'),
        ('Website Bucket', 'S3', 'slashmycloudbill.com', 'us-east-1'),
        ('Storage Bucket', 'S3', 'aws-bill-analyzer-storage-991105135552', 'us-east-1'),
        ('API', 'API Gateway HTTP', 'ViewMyBill-API', 'us-east-1'),
        ('Bill Analyzer', 'Lambda', 'ViewMyBill-BillAnalyzer', 'us-east-1'),
        ('Upload Handler', 'Lambda', 'ViewMyBill-UploadHandler', 'us-east-1'),
        ('OTP Handler', 'Lambda', 'ViewMyBill-OTPHandler', 'us-east-1'),
        ('Member Handler', 'Lambda', 'ViewMyBill-MemberHandler', 'us-east-1'),
        ('Admin Handler', 'Lambda', 'ViewMyBill-AdminHandler', 'us-east-1'),
        ('Agent Action', 'Lambda', 'ViewMyBill-AgentAction', 'us-east-1'),
        ('AI Model', 'Amazon Bedrock', 'us.amazon.nova-2-lite-v1:0 (cross-region)', 'us-east-1'),
        ('User Pool', 'Cognito', 'SlashMyBill-Members', 'us-east-1'),
        ('Email', 'SES', 'noreply@slashmycloudbill.com', 'us-east-1'),
        ('TLS Certs', 'ACM', '*.slashmycloudbill.com, *.eshkolai.com', 'us-east-1'),
        ('CI/CD', 'GitHub Actions', 'OIDC → IAM Role', 'us-east-1'),
    ]
    for row_data in components:
        add_table_row(tbl, row_data)

    add_body(doc, 'See PublicSite.drawio for the full architecture diagram with data flow edges.')


    # ════════════════════════════════════════════════════════════════════════
    # 3. FRONTEND ARCHITECTURE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('3. Frontend Architecture', level=1)

    doc.add_heading('3.1 Landing Page (slashmycloudbill.com)', level=2)
    add_body(doc,
        'The public landing page is a single-page HTML/CSS/JS application hosted on S3 '
        'and served via CloudFront. The design is inspired by slashmybill.ai with a modern, '
        'conversion-focused layout.')
    add_bullet(doc, 'Hero section with bill upload CTA and animated cost-savings counter')
    add_bullet(doc, 'How-it-works 3-step flow (Upload → Analyze → Save)')
    add_bullet(doc, 'Pricing cards (Free / Growth / Scale tiers)')
    add_bullet(doc, 'Contact form (→ contact-form-handler Lambda → SES)')
    add_bullet(doc, 'Member login / register modal with OTP verification')
    add_bullet(doc, 'Responsive design with mobile-first approach')

    doc.add_heading('3.2 Member Portal', level=2)
    add_body(doc,
        'The member portal is a SPA at /members/ with 4 main tabs:')

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Tab', 'Purpose', 'Key Features']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    tabs = [
        ('Observe', 'Dashboard & Monitoring',
         '8 customizable widgets, multi-account selector, date range picker, auto-refresh'),
        ('Chat', 'AI-Powered Console',
         'Natural language queries, direct model + agent mode, conversation history, feedback (👍/👎), token tracking'),
        ('Act', 'Cleanup Actions',
         'Tips-driven scan engine, 7 cleanup categories, JIT safety checks, one-click execution, action history'),
        ('Configure', 'Account Management',
         'Add/remove AWS accounts, IAM role setup wizard, CloudFormation template download, tier & billing info'),
    ]
    for row_data in tabs:
        add_table_row(tbl, row_data)

    doc.add_heading('3.3 Admin Panel', level=2)
    add_body(doc, 'The admin panel at /admin/ provides back-office management:')
    add_bullet(doc, 'Leads Management — view/search/export bill analysis leads with status tracking')
    add_bullet(doc, 'Tips Management — CRUD for cost optimization tips (30+ tips in knowledge base)')
    add_bullet(doc, 'Feedback Viewer — review AI chat feedback with thumbs up/down and comments')
    add_bullet(doc, 'Authentication via legacy JWT (admin username + bcrypt password hash)')

    doc.add_heading('3.4 CloudFront Function Routing', level=2)
    add_body(doc,
        'The SlashMyCloudBill-Router CloudFront Function handles URL rewriting:')
    add_bullet(doc, '/members → /members/index.html')
    add_bullet(doc, '/members/* → pass through (SPA assets)')
    add_bullet(doc, '/admin → /admin/index.html')
    add_bullet(doc, '/admin/* → pass through')
    add_bullet(doc, '/slashMyBill → /slashMyBill/index.html (bill analyzer tool)')
    add_bullet(doc, 'Default: serve from S3 origin with index.html fallback')

    # ════════════════════════════════════════════════════════════════════════
    # 4. API LAYER
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('4. API Layer', level=1)

    doc.add_heading('4.1 API Gateway Routes', level=2)
    add_body(doc, 'ViewMyBill-API is an HTTP API (API Gateway v2) with the following routes:')

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Method', 'Route', 'Lambda Target', 'Auth']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    routes = [
        ('POST', '/analyze', 'Bill Analyzer', 'None'),
        ('POST', '/upload', 'Upload Handler', 'None'),
        ('POST', '/otp/send', 'OTP Handler', 'None'),
        ('POST', '/otp/verify', 'OTP Handler', 'None'),
        ('POST', '/members/register', 'Member Handler', 'None'),
        ('POST', '/members/login', 'Member Handler', 'None'),
        ('GET',  '/members/accounts', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/accounts', 'Member Handler', 'Cognito/JWT'),
        ('PUT',  '/members/accounts', 'Member Handler', 'Cognito/JWT'),
        ('DELETE','/members/accounts', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/accounts/template', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/accounts/test', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/query', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/query-agent', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/feedback', 'Member Handler', 'Cognito/JWT'),
        ('GET',  '/members/dashboard', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/scan', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/members/actions', 'Member Handler', 'Cognito/JWT'),
        ('GET',  '/members/profile', 'Member Handler', 'Cognito/JWT'),
        ('PUT',  '/members/profile', 'Member Handler', 'Cognito/JWT'),
        ('POST', '/admin/login', 'Admin Handler', 'None'),
        ('GET',  '/admin/leads', 'Admin Handler', 'JWT'),
        ('GET',  '/admin/tips', 'Admin Handler', 'JWT'),
        ('POST', '/admin/tips', 'Admin Handler', 'JWT'),
        ('PUT',  '/admin/tips', 'Admin Handler', 'JWT'),
        ('DELETE','/admin/tips', 'Admin Handler', 'JWT'),
        ('GET',  '/admin/feedback', 'Admin Handler', 'JWT'),
    ]
    for row_data in routes:
        add_table_row(tbl, row_data)

    doc.add_heading('4.2 CORS Configuration', level=2)
    add_body(doc, 'CORS is configured at the API Gateway level:')
    add_bullet(doc, 'Allowed Origins: https://slashmycloudbill.com, https://www.slashmycloudbill.com, https://www.eshkolai.com')
    add_bullet(doc, 'Allowed Methods: GET, POST, PUT, DELETE, OPTIONS')
    add_bullet(doc, 'Allowed Headers: Content-Type, Authorization, X-Requested-With')
    add_bullet(doc, 'Max Age: 86400 seconds')

    doc.add_heading('4.3 Authentication Flow', level=2)
    add_body(doc,
        'The API supports dual authentication: Cognito (primary for members) and legacy JWT '
        '(for admin panel). Member endpoints validate the Authorization header first against '
        'Cognito User Pool tokens, falling back to JWT verification if Cognito validation fails. '
        'This ensures backward compatibility during the migration period.')


    # ════════════════════════════════════════════════════════════════════════
    # 5. LAMBDA FUNCTIONS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('5. Lambda Functions', level=1)

    # 5.1 Bill Analyzer
    doc.add_heading('5.1 Bill Analyzer', level=2)
    add_body(doc,
        'Parses uploaded AWS bills (PDF), sends extracted data to Amazon Bedrock for AI analysis, '
        'and generates a branded PDF optimization report.')
    add_bullet(doc, 'PDF Parsing: PyPDF2 for text extraction, pdfplumber as fallback for table-heavy bills')
    add_bullet(doc, 'AI Analysis: Sends structured bill data to Bedrock Nova 2 Lite with a FinOps-tuned prompt')
    add_bullet(doc, 'Report Generation: ReportLab-based PDF with branded header (SlashMyBill logo), charts, and actionable recommendations')
    add_bullet(doc, 'Async Self-Invoke: For large bills, the Lambda invokes itself asynchronously to avoid API Gateway 30s timeout')
    add_bullet(doc, 'Lead Capture: Saves email + bill metadata to ViewMyBill-Leads DynamoDB table')
    add_bullet(doc, 'Storage: Uploads report PDF to S3 with 24-hour lifecycle expiration')
    add_body(doc, 'Runtime: Python 3.12 | Memory: 512 MB | Timeout: 300s')

    # 5.2 Upload Handler
    doc.add_heading('5.2 Upload Handler', level=2)
    add_body(doc,
        'Handles multipart form data uploads and generates S3 presigned URLs.')
    add_bullet(doc, 'Parses multipart/form-data boundary from the request body')
    add_bullet(doc, 'Validates file type (PDF only) and size (max 10 MB)')
    add_bullet(doc, 'Generates S3 presigned upload URL with 5-minute expiration')
    add_bullet(doc, 'Returns presigned URL + object key for client-side direct upload')
    add_body(doc, 'Runtime: Python 3.12 | Memory: 256 MB | Timeout: 30s')

    # 5.3 OTP Handler
    doc.add_heading('5.3 OTP Handler', level=2)
    add_body(doc,
        'Manages one-time password generation and verification for email-based authentication.')
    add_bullet(doc, '6-digit cryptographically random code generation (secrets module)')
    add_bullet(doc, 'DynamoDB storage with TTL of 5 minutes (300 seconds)')
    add_bullet(doc, 'Rate limiting: max 1 OTP per email per 60 seconds')
    add_bullet(doc, 'SES email delivery with branded HTML template')
    add_bullet(doc, 'Verification endpoint consumes the OTP (single-use)')
    add_bullet(doc, 'Max 5 verification attempts before OTP invalidation')
    add_body(doc, 'Runtime: Python 3.12 | Memory: 128 MB | Timeout: 10s')

    # 5.4 Member Handler
    doc.add_heading('5.4 Member Handler', level=2)
    add_body(doc,
        'The largest Lambda function — handles all member portal operations:')

    doc.add_heading('Registration & Auth', level=3)
    add_bullet(doc, 'Cognito-based registration with pre-verified flow (from bill analysis) or OTP flow')
    add_bullet(doc, 'Login via Cognito InitiateAuth (USER_PASSWORD_AUTH)')
    add_bullet(doc, 'Token refresh via Cognito refresh tokens')
    add_bullet(doc, 'Profile management (display name, company, preferences)')

    doc.add_heading('Account Management', level=3)
    add_bullet(doc, 'CRUD operations for linked AWS accounts')
    add_bullet(doc, 'IAM role setup: generates CloudFormation template for SlashMyBill-{AccountID} role')
    add_bullet(doc, 'Connection test: STS AssumeRole validation')
    add_bullet(doc, 'Tier enforcement: Free=1 account, Growth/Scale=20 accounts')

    doc.add_heading('AI Query (Direct Model)', level=3)
    add_bullet(doc, 'Gathers account data via _gather_account_data() — Cost Explorer, EC2, RDS, S3, EBS, ELB')
    add_bullet(doc, 'Builds context-rich prompt with real cost data')
    add_bullet(doc, 'Sends to Bedrock Nova 2 Lite (us.amazon.nova-2-lite-v1:0)')
    add_bullet(doc, 'Multi-account support: queries across all linked accounts')
    add_bullet(doc, 'Token cost: 2 tokens per AI question')

    doc.add_heading('AI Query (Agent Mode)', level=3)
    add_bullet(doc, 'Routes to Bedrock Agent with action groups')
    add_bullet(doc, 'Agent can autonomously gather data and execute analysis')
    add_bullet(doc, 'Feedback loop: thumbs up/down stored in MemberPortal-AgentFeedback')

    doc.add_heading('Dashboard Data Aggregation', level=3)
    add_bullet(doc, 'Fetches Cost Explorer data for all linked accounts')
    add_bullet(doc, 'Computes widget data: cost by service, trends, waste detection, unit costs')
    add_bullet(doc, 'Virtual tagging for business unit cost allocation')
    add_bullet(doc, 'Caches results in MemberPortal-BusinessMetrics (15-minute TTL)')

    doc.add_heading('Actions / Scan Engine', level=3)
    add_bullet(doc, 'Tips-driven scan: loads optimization tips from DynamoDB')
    add_bullet(doc, 'Scans linked accounts for waste across 7 categories')
    add_bullet(doc, 'Executes cleanup actions with JIT safety checks')
    add_bullet(doc, 'Token cost: scan=10 tokens, cleanup action=50 tokens')

    add_body(doc, 'Runtime: Python 3.12 | Memory: 1024 MB | Timeout: 300s')

    # 5.5 Admin Handler
    doc.add_heading('5.5 Admin Handler', level=2)
    add_body(doc, 'Back-office API for the admin panel:')
    add_bullet(doc, 'Leads CRUD: list, search, update status, export')
    add_bullet(doc, 'Tips Management: create, read, update, delete optimization tips')
    add_bullet(doc, 'Feedback Viewer: paginated list of AI chat feedback with filters')
    add_bullet(doc, 'Billing Sync: aggregate member usage for invoicing')
    add_bullet(doc, 'Auth: JWT-based (admin username + bcrypt password hash)')
    add_body(doc, 'Runtime: Python 3.12 | Memory: 256 MB | Timeout: 30s')

    # 5.6 Agent Action
    doc.add_heading('5.6 Agent Action', level=2)
    add_body(doc,
        'Bedrock Agent action group Lambda — invoked by the Bedrock Agent to perform '
        'automated FinOps tasks:')
    add_bullet(doc, 'get_cost_summary: Fetches Cost Explorer data for a given account and date range')
    add_bullet(doc, 'get_resource_inventory: Lists EC2, RDS, S3, EBS resources')
    add_bullet(doc, 'get_optimization_tips: Retrieves relevant tips from the knowledge base')
    add_bullet(doc, 'execute_action: Performs a cleanup action (with safety validation)')
    add_body(doc, 'Runtime: Python 3.12 | Memory: 256 MB | Timeout: 60s')


    # ════════════════════════════════════════════════════════════════════════
    # 6. DATA MODEL
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('6. Data Model', level=1)
    add_body(doc, 'All data is stored in Amazon DynamoDB with on-demand capacity mode.')

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Table Name', 'Partition Key', 'Sort Key', 'Key Attributes', 'TTL', 'GSIs']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(8)

    tables_data = [
        ('ViewMyBill-Leads', 'email', 'timestamp',
         'name, company, bill_s3_key, report_s3_key, status, source',
         'None', 'status-index (status, timestamp)'),
        ('ViewMyBill-OTP', 'email', 'None',
         'otp_code, created_at, attempts, verified',
         'ttl (5 min)', 'None'),
        ('ViewMyBill-CostOptimizationTips', 'tip_id', 'None',
         'category, title, description, severity, service, check_type, remediation',
         'None', 'category-index (category)'),
        ('MemberPortal-Members', 'email', 'None',
         'name, company, password_hash, tier, tokens_remaining, tokens_reset_date, '
         'bonus_tokens, created_at, cognito_sub',
         'None', 'cognito-sub-index (cognito_sub)'),
        ('MemberPortal-Accounts', 'member_email', 'account_id',
         'account_name, role_arn, external_id, status, linked_at, last_scan',
         'None', 'None'),
        ('MemberPortal-AgentFeedback', 'feedback_id', 'None',
         'member_email, query, response, rating (up/down), comment, model, timestamp',
         'None', 'member-index (member_email, timestamp)'),
        ('MemberPortal-BusinessMetrics', 'member_email#account_id', 'metric_date',
         'widget_type, data_json, computed_at',
         'ttl (15 min)', 'None'),
    ]
    for row_data in tables_data:
        add_table_row(tbl, row_data)

    # ════════════════════════════════════════════════════════════════════════
    # 7. AUTHENTICATION & AUTHORIZATION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('7. Authentication & Authorization', level=1)

    doc.add_heading('7.1 Cognito User Pool', level=2)
    add_body(doc, 'Amazon Cognito User Pool "SlashMyBill-Members" configuration:')
    add_bullet(doc, 'Sign-in: email (username alias)')
    add_bullet(doc, 'Password policy: min 8 chars, requires uppercase, lowercase, number, special char')
    add_bullet(doc, 'MFA: optional (TOTP)')
    add_bullet(doc, 'App client: SlashMyBill-WebApp (USER_PASSWORD_AUTH, ALLOW_REFRESH_TOKEN_AUTH)')
    add_bullet(doc, 'Token validity: Access=1 hour, ID=1 hour, Refresh=30 days')

    doc.add_heading('7.2 Registration Flow (3-Step)', level=2)
    add_numbered(doc, 'send-otp: User enters email → OTP Handler generates 6-digit code → SES sends email')
    add_numbered(doc, 'verify-otp: User enters code → OTP Handler validates against DynamoDB → returns verification token')
    add_numbered(doc, 'create-account: User submits name + password + verification token → Member Handler creates Cognito user + DynamoDB record')

    doc.add_heading('7.3 Pre-Verified Flow', level=2)
    add_body(doc,
        'Users who have already analyzed a bill are pre-verified. The bill analysis lead record '
        'serves as proof of email ownership, allowing them to skip the OTP step during registration.')

    doc.add_heading('7.4 Token Validation', level=2)
    add_bullet(doc, 'Primary: Cognito ID token validation (JWKS verification, issuer check, expiry check)')
    add_bullet(doc, 'Fallback: Legacy JWT validation (HS256, shared secret) for backward compatibility')
    add_bullet(doc, 'Admin endpoints: JWT-only (separate admin credentials)')

    doc.add_heading('7.5 Cross-Account IAM Role', level=2)
    add_body(doc,
        'Each linked customer account creates an IAM role named SlashMyBill-{AccountID}:')
    add_bullet(doc, 'Trust policy: allows sts:AssumeRole from platform account (991105135552)')
    add_bullet(doc, 'Permissions: ReadOnlyAccess + ce:GetCostAndUsage + limited write for cleanup actions')
    add_bullet(doc, 'External ID: unique per member for confused deputy protection')
    add_bullet(doc, 'CloudFormation template provided for one-click role creation')


    # ════════════════════════════════════════════════════════════════════════
    # 8. AI/ML INTEGRATION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('8. AI/ML Integration', level=1)

    doc.add_heading('8.1 Amazon Bedrock — Nova 2 Lite', level=2)
    add_body(doc,
        'The platform uses Amazon Bedrock with the Nova 2 Lite model via cross-region '
        'inference profile (us.amazon.nova-2-lite-v1:0). This provides low-latency, '
        'cost-effective AI inference for FinOps analysis.')
    add_bullet(doc, 'Model ID: us.amazon.nova-2-lite-v1:0')
    add_bullet(doc, 'Access: Cross-region inference profile (automatic region routing)')
    add_bullet(doc, 'Max tokens: 4096 output tokens per request')
    add_bullet(doc, 'Temperature: 0.3 (low creativity, high accuracy for financial data)')
    add_bullet(doc, 'Use cases: bill analysis, cost optimization Q&A, resource recommendations')

    doc.add_heading('8.2 Knowledge Base', level=2)
    add_body(doc,
        'The knowledge base is stored in the ViewMyBill-CostOptimizationTips DynamoDB table '
        'with 30+ curated optimization tips across categories:')
    add_bullet(doc, 'Compute: rightsizing, Savings Plans, Reserved Instances, Spot usage')
    add_bullet(doc, 'Storage: S3 lifecycle policies, EBS optimization, snapshot cleanup')
    add_bullet(doc, 'Network: NAT Gateway optimization, data transfer reduction')
    add_bullet(doc, 'Database: RDS rightsizing, Aurora Serverless migration, DynamoDB capacity')
    add_bullet(doc, 'General: tag governance, account consolidation, unused resource cleanup')

    doc.add_heading('8.3 Data Gathering Pipeline', level=2)
    add_body(doc,
        'The _gather_account_data() function collects real-time data from customer accounts:')
    add_numbered(doc, 'STS AssumeRole into customer account')
    add_numbered(doc, 'Cost Explorer: GetCostAndUsage (last 30 days, grouped by service)')
    add_numbered(doc, 'EC2: DescribeInstances (running, stopped, instance types, utilization)')
    add_numbered(doc, 'RDS: DescribeDBInstances (engine, size, multi-AZ, utilization)')
    add_numbered(doc, 'S3: ListBuckets + GetBucketMetrics (size, request counts)')
    add_numbered(doc, 'EBS: DescribeVolumes (unattached, type, size, IOPS)')
    add_numbered(doc, 'ELB: DescribeLoadBalancers (idle detection)')
    add_numbered(doc, 'EIP: DescribeAddresses (unassociated)')

    doc.add_heading('8.4 Multi-Account Query Support', level=2)
    add_body(doc,
        'When a member has multiple linked accounts, AI queries aggregate data across all '
        'accounts. The prompt includes a per-account breakdown so the model can provide '
        'cross-account optimization recommendations (e.g., consolidating Reserved Instances).')

    doc.add_heading('8.5 Feedback Loop', level=2)
    add_body(doc,
        'Every AI response includes thumbs up/down buttons. Feedback is stored in '
        'MemberPortal-AgentFeedback with the full query, response, rating, and optional '
        'comment. This data is used to improve prompts and identify model weaknesses.')

    # ════════════════════════════════════════════════════════════════════════
    # 9. PRICING & TOKEN SYSTEM
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('9. Pricing & Token System', level=1)

    doc.add_heading('9.1 Subscription Tiers', level=2)

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Feature', 'Free', 'Growth ($50/mo)', 'Scale ($200/mo)', 'Managed']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    tiers = [
        ('Monthly Tokens', '100', '300', '1,500', 'Unlimited'),
        ('Linked Accounts', '1', '20', '20', 'Unlimited'),
        ('AI Questions', '✓', '✓', '✓', '✓'),
        ('Dashboard Widgets', '4 basic', '8 full', '8 full + custom', '8 full + custom'),
        ('Cleanup Actions', '—', '✓', '✓', '✓ (managed)'),
        ('Scan Engine', '—', '✓', '✓', '✓'),
        ('Priority Support', '—', '—', '✓', '✓'),
        ('Dedicated FinOps', '—', '—', '—', '✓'),
    ]
    for row_data in tiers:
        add_table_row(tbl, row_data)

    doc.add_heading('9.2 Token Costs', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Action', 'Token Cost', 'Notes']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    costs = [
        ('AI Question (direct)', '2', 'Per query to Bedrock model'),
        ('AI Question (agent)', '3', 'Per agent invocation'),
        ('Account Scan', '10', 'Full scan across all 7 categories'),
        ('Cleanup Action', '50', 'Per resource cleanup execution'),
        ('Dashboard Refresh', '0', 'Free — cached for 15 minutes'),
        ('Bill Analysis', '0', 'Free — lead generation funnel'),
    ]
    for row_data in costs:
        add_table_row(tbl, row_data)

    doc.add_heading('9.3 Token Reset & Bonus', level=2)
    add_bullet(doc, 'Monthly Reset: tokens_remaining resets to tier allocation on billing date')
    add_bullet(doc, 'Unused tokens do NOT roll over')
    add_bullet(doc, 'Bonus Tokens: one-time top-up purchases stored in bonus_tokens field')
    add_bullet(doc, 'Bonus tokens are consumed AFTER monthly allocation is exhausted')
    add_bullet(doc, 'Bonus tokens do NOT expire')

    doc.add_heading('9.4 Managed Service', level=2)
    add_body(doc,
        'For enterprise clients, SlashMyCloudBill offers a managed service option: '
        'a dedicated FinOps engineer reviews the account monthly and implements optimizations. '
        'Pricing: 25% of verified yearly savings (minimum $500/month).')


    # ════════════════════════════════════════════════════════════════════════
    # 10. DASHBOARD WIDGETS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('10. Dashboard Widgets', level=1)
    add_body(doc,
        'The Observe tab features 8 customizable widgets rendered with ECharts. '
        'Members can reorder, resize, and toggle widget visibility.')

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Widget', 'Chart Type', 'Data Source', 'Features']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    widgets = [
        ('Cost by Service', 'Treemap', 'Cost Explorer (GroupBy SERVICE)',
         '2-level drill-down (service → usage type), color-coded by spend magnitude'),
        ('Cost Trend', 'Line chart', 'Cost Explorer (daily/hourly)',
         'Daily/hourly toggle, 7/14/30/90 day range, anomaly highlighting'),
        ('Cost Allocation by BU', 'Stacked bar', 'Virtual tagging engine',
         'Business unit mapping via tag rules, untagged resource detection'),
        ('Waste Detection', 'Gauge + list', 'Scan engine results',
         'Waste score (0-100), top 10 waste items with estimated monthly savings'),
        ('Monthly Cost by Service', 'Stacked bar', 'Cost Explorer (monthly)',
         'Month-over-month comparison, top 5 services + "Other" grouping'),
        ('Unit Cost Trend', 'Line chart', 'Cost Explorer + custom metrics',
         'Cost per user/transaction/request, configurable denominator'),
        ('Cost by Region', 'Pie chart', 'Cost Explorer (GroupBy REGION)',
         'Interactive legend, click to filter, data transfer cost overlay'),
        ('Savings Plans & RI', 'Donut + table', 'Savings Plans + RI utilization APIs',
         'Coverage %, utilization %, expiring commitments, recommendation engine'),
    ]
    for row_data in widgets:
        add_table_row(tbl, row_data)

    # ════════════════════════════════════════════════════════════════════════
    # 11. ACTIONS / SCAN ENGINE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('11. Actions / Scan Engine', level=1)

    doc.add_heading('11.1 Tips-Driven Scan', level=2)
    add_body(doc,
        'The scan engine loads optimization tips from ViewMyBill-CostOptimizationTips and '
        'evaluates each tip against the customer\'s linked accounts. Each tip defines a '
        'check_type (e.g., "unused_eip", "detached_ebs") and the engine runs the corresponding '
        'AWS API calls to detect matching resources.')
    add_bullet(doc, '30+ built-in checks across 7 categories')
    add_bullet(doc, 'Severity levels: critical, high, medium, low, info')
    add_bullet(doc, 'Estimated monthly savings calculated per finding')
    add_bullet(doc, 'Results cached in MemberPortal-BusinessMetrics (15-minute TTL)')

    doc.add_heading('11.2 Cleanup Categories', level=2)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Category', 'Check Examples', 'Action', 'Safety']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    categories = [
        ('Elastic IPs', 'Unassociated EIPs', 'Release EIP', 'Verify not in DNS records'),
        ('EBS Volumes', 'Detached volumes, gp2→gp3 migration', 'Delete / Modify type',
         'Check for recent snapshots, verify no pending attachments'),
        ('Load Balancers', 'Idle ALB/NLB (0 healthy targets)', 'Delete LB + target group',
         'Verify no DNS CNAME pointing to LB'),
        ('S3 Buckets', 'No lifecycle policy, public access', 'Add lifecycle / Block public',
         'Check for static website hosting, verify no CloudFront origin'),
        ('EC2 Instances', 'Stopped >7 days, oversized', 'Stop / Terminate / Rightsize',
         'ASG detach before terminate, create AMI backup'),
        ('RDS Instances', 'Idle (0 connections >7 days), oversized', 'Stop / Delete / Rightsize',
         'Final snapshot before delete, check for read replicas'),
        ('Snapshots', 'Orphaned snapshots (no parent volume)', 'Delete snapshot',
         'Verify not used by any AMI, check age >30 days'),
    ]
    for row_data in categories:
        add_table_row(tbl, row_data)

    doc.add_heading('11.3 JIT Safety Checks', level=2)
    add_body(doc,
        'Before every cleanup action, the engine performs just-in-time safety validation:')
    add_numbered(doc, 'Re-fetch resource state (confirm still exists and matches expected state)')
    add_numbered(doc, 'Check for dependencies (DNS records, ASG membership, AMI references)')
    add_numbered(doc, 'Verify no recent activity (CloudTrail events in last 24 hours)')
    add_numbered(doc, 'Create safety backup if applicable (snapshot, AMI, final DB snapshot)')
    add_numbered(doc, 'Execute action with error handling and rollback capability')
    add_numbered(doc, 'Log action to MemberPortal-BusinessMetrics for audit trail')


    # ════════════════════════════════════════════════════════════════════════
    # 12. SECURITY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('12. Security', level=1)

    doc.add_heading('12.1 Encryption', level=2)
    add_bullet(doc, 'In Transit: TLS 1.2+ enforced on all endpoints (CloudFront, API Gateway, S3)')
    add_bullet(doc, 'At Rest: S3 SSE-S3 (AES-256) for all buckets, DynamoDB encryption enabled')
    add_bullet(doc, 'Cognito tokens: signed JWTs with RS256 (asymmetric)')

    doc.add_heading('12.2 Data Lifecycle', level=2)
    add_bullet(doc, 'Uploaded bills: S3 lifecycle rule deletes after 24 hours')
    add_bullet(doc, 'Generated reports: S3 lifecycle rule deletes after 24 hours')
    add_bullet(doc, 'OTP codes: DynamoDB TTL auto-deletes after 5 minutes')
    add_bullet(doc, 'Dashboard cache: DynamoDB TTL auto-deletes after 15 minutes')

    doc.add_heading('12.3 CORS Restrictions', level=2)
    add_bullet(doc, 'API Gateway CORS allows only: slashmycloudbill.com, www.slashmycloudbill.com, www.eshkolai.com')
    add_bullet(doc, 'S3 bucket CORS: restricted to same origins')
    add_bullet(doc, 'CloudFront: custom headers for origin verification')

    doc.add_heading('12.4 IAM Least Privilege', level=2)
    add_bullet(doc, 'Each Lambda has a dedicated IAM role with minimal permissions')
    add_bullet(doc, 'Bill Analyzer: S3 read/write (storage bucket only) + Bedrock InvokeModel + DynamoDB PutItem (leads)')
    add_bullet(doc, 'Member Handler: DynamoDB CRUD (member tables) + STS AssumeRole + Bedrock + SES')
    add_bullet(doc, 'No Lambda has admin or wildcard (*) permissions')

    doc.add_heading('12.5 Cross-Account Access', level=2)
    add_bullet(doc, 'ReadOnly base policy — no write access to customer resources by default')
    add_bullet(doc, 'Cleanup actions require explicit opt-in (additional policy attached to role)')
    add_bullet(doc, 'External ID required for confused deputy protection')
    add_bullet(doc, 'Role session name includes member email for CloudTrail attribution')

    doc.add_heading('12.6 No Credential Storage', level=2)
    add_body(doc,
        'The platform never stores AWS access keys or secret keys. All cross-account access '
        'uses IAM role assumption with temporary credentials (STS). Customer credentials are '
        'never transmitted to or stored by the platform.')

    # ════════════════════════════════════════════════════════════════════════
    # 13. CI/CD PIPELINE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('13. CI/CD Pipeline', level=1)

    doc.add_heading('13.1 GitHub Actions Workflow', level=2)
    add_body(doc,
        'The deployment pipeline is defined in .github/workflows/deploy.yml and triggers '
        'on push to main branch when relevant files change.')

    doc.add_heading('13.2 OIDC Authentication', level=2)
    add_body(doc,
        'No AWS credentials are stored in GitHub. The pipeline uses OpenID Connect (OIDC) '
        'to assume an IAM role (github-oidc-role) with deployment permissions. The trust '
        'policy restricts access to the specific GitHub repository and branch.')

    doc.add_heading('13.3 Pipeline Steps', level=2)
    add_numbered(doc, 'Checkout: Clone repository')
    add_numbered(doc, 'Configure AWS: OIDC → AssumeRole (github-oidc-role)')
    add_numbered(doc, 'Package Lambdas: pip install + zip for each Lambda (bill-analyzer, upload-handler, otp-handler, member-handler, admin-handler, agent-action)')
    add_numbered(doc, 'Upload packages: S3 → aws-bill-analyzer-storage/lambda-packages/')
    add_numbered(doc, 'Deploy CloudFormation: Update aws-bill-analyzer-viewmybill stack')
    add_numbered(doc, 'Update Lambda code: aws lambda update-function-code for each function')
    add_numbered(doc, 'Deploy Frontend (eshkolai): Sync to S3 www.eshkolai.com bucket')
    add_numbered(doc, 'Deploy Frontend (SMB): Sync to S3 slashmycloudbill.com bucket')
    add_numbered(doc, 'Invalidate CloudFront: Create invalidation for both distributions')
    add_numbered(doc, 'Seed Knowledge Base: Update DynamoDB tips table from knowledge-base/aws-cost-optimization-tips.json')

    doc.add_heading('13.4 Dual Deployment', level=2)
    add_body(doc,
        'The pipeline deploys to both domains:')
    add_bullet(doc, 'slashmycloudbill.com — primary domain, CloudFront E2B3GXE4TJTH4Q')
    add_bullet(doc, 'eshkolai.com — legacy domain, CloudFront E12JIHGHK40OLE (redirects to SMB for most paths)')

    # ════════════════════════════════════════════════════════════════════════
    # 14. INFRASTRUCTURE AS CODE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading('14. Infrastructure as Code', level=1)

    doc.add_heading('14.1 CloudFormation Stacks', level=2)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(['Stack Name', 'Template', 'Purpose', 'Key Resources']):
        hdr[i].text = h
        set_cell_shading(hdr[i], "2E4057")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(9)

    stacks = [
        ('aws-bill-analyzer-viewmybill', 'infrastructure/viewmybill-stack.yaml',
         'Core application stack',
         'API Gateway, 6 Lambdas, 7 DynamoDB tables, IAM roles, Cognito User Pool'),
        ('github-oidc-role', 'infrastructure/github-oidc-role.yaml',
         'CI/CD authentication',
         'IAM OIDC Provider, IAM Role with deployment permissions'),
        ('(CLI-managed)', 'infrastructure/deploy-cloudfront.ps1',
         'CDN + static hosting',
         'CloudFront distribution, S3 bucket, Route 53 records, ACM certificate'),
    ]
    for row_data in stacks:
        add_table_row(tbl, row_data)

    doc.add_heading('14.2 Direct CLI Resources', level=2)
    add_body(doc,
        'Some resources are managed via AWS CLI scripts rather than CloudFormation:')
    add_bullet(doc, 'slashmycloudbill.com S3 bucket + CloudFront distribution (deploy-cloudfront.ps1)')
    add_bullet(doc, 'Route 53 hosted zones and records (deploy-domain.ps1)')
    add_bullet(doc, 'ACM certificates (manual request + DNS validation)')
    add_bullet(doc, 'CloudFront Function (cf-function-slashmycloudbill.js)')
    add_bullet(doc, 'Bedrock Agent configuration (create-bedrock-agent.py)')
    add_bullet(doc, 'SES domain verification and sending configuration')

    # ── Save ────────────────────────────────────────────────────────────────
    doc.save('SlashMyBill-HLD.docx')
    print('✅ SlashMyBill-HLD.docx generated successfully!')
    print('   Open in Microsoft Word or Google Docs to view.')


# ── Entry point ──────────────────────────────────────────────────────────────

if __name__ == '__main__':
    build_document()
