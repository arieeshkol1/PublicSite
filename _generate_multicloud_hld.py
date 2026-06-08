"""Generate SlashMyBill Multi-Cloud Support HLD as a Word document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# --- TITLE PAGE ---
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('SlashMyBill')
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Multi-Cloud Support\nHigh-Level Design Document')
run.bold = True
run.font.size = Pt(24)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run(f'Version 1.0\n{datetime.date.today().strftime("%B %d, %Y")}\n\n')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run2 = meta.add_run('Platform: https://www.eshkolai.com/members/\nAWS Account: 991105135552 (us-east-1)')
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
    doc.add_paragraph()
    return table

# --- SECTION 1: EXECUTIVE SUMMARY ---
add_heading('1. Executive Summary')
doc.add_paragraph(
    'Multi-Cloud Support extends SlashMyBill from an AWS-only FinOps platform to a unified '
    'multi-cloud solution supporting AWS, Microsoft Azure, and Google Cloud Platform (GCP). '
    'Members connect accounts from all three providers within the same user account, view '
    'unified cost data in a single dashboard, and receive cloud-specific optimization tips \u2014 '
    'all powered by the same AI agent.'
)
doc.add_paragraph(
    'The extension uses a Provider Connector pattern where each cloud provider has a dedicated '
    'connector module implementing a common interface. A shared Cost Normalizer transforms '
    'provider-specific responses into a unified schema for display and AI analysis. The existing '
    'AWS flow remains completely unchanged.'
)
p = doc.add_paragraph()
p.add_run('Supported Providers: ').bold = True
p.add_run('AWS | Azure | GCP')
p = doc.add_paragraph()
p.add_run('Platform URL: ').bold = True
p.add_run('https://www.eshkolai.com/members/')
p = doc.add_paragraph()
p.add_run('AWS Account: ').bold = True
p.add_run('991105135552 (us-east-1)')

doc.add_page_break()

# --- SECTION 2: ARCHITECTURE OVERVIEW ---
add_heading('2. Architecture Overview')
add_heading('2.1 Core Components (New/Extended)', level=2)
add_table(
    ['Component', 'Technology', 'Purpose'],
    [
        ['AWS Connector', 'Python 3.12 (existing)', 'STS AssumeRole + Cost Explorer (unchanged)'],
        ['Azure Connector', 'Python 3.12 (new)', 'OAuth2 Service Principal + Azure Cost Management API'],
        ['GCP Connector', 'Python 3.12 (new)', 'JWT Service Account + GCP Cloud Billing API'],
        ['Cost Normalizer', 'Python 3.12 (new)', 'Transforms provider-specific cost data to common schema'],
        ['Member Handler', 'Python 3.12 Lambda (extended)', 'Multi-cloud account CRUD, connection testing, unified dashboard'],
        ['Admin Handler', 'Python 3.12 Lambda (extended)', 'Multi-cloud tips management, per-provider filtering'],
        ['Tips Sync Lambda', 'Python 3.12 Lambda (extended)', 'Daily sync of Azure + GCP tips from curated JSON files'],
        ['KMS', 'AWS KMS (existing)', 'Encrypts Azure Client Secrets and GCP private keys'],
        ['Accounts Table', 'DynamoDB (extended)', 'New cloudProvider + credentials attributes'],
        ['Tips Table', 'DynamoDB (extended)', 'New cloudProvider attribute per tip'],
        ['Cost Cache Table', 'DynamoDB (extended)', 'Cache key includes cloudProvider'],
        ['Member Portal', 'HTML/CSS/JS (extended)', 'Provider selection, unified dashboard, provider icons'],
        ['Admin Panel', 'HTML/CSS/JS (extended)', 'Provider tabs, sync status display'],
    ]
)

add_heading('2.2 Multi-Cloud Access Model', level=2)
add_table(
    ['Provider', 'Authentication Method', 'Required Permissions'],
    [
        ['AWS', 'STS AssumeRole with ExternalId (SHA-256 of email)', 'ReadOnlyAccess + Cost Explorer inline policy (unchanged)'],
        ['Azure', 'OAuth2 Client Credentials (Service Principal)', '"Cost Management Reader" role on subscription'],
        ['GCP', 'Self-signed JWT (Service Account key)', '"Billing Account Viewer" + "BigQuery User" roles'],
    ]
)

add_heading('2.3 System Context Diagram', level=2)
doc.add_paragraph(
    'The system follows a hub-and-spoke architecture where the Member Handler Lambda acts as '
    'the central orchestrator, dispatching to provider-specific connectors based on the '
    'cloudProvider attribute stored with each account.'
)
doc.add_paragraph()
p = doc.add_paragraph(style='List Bullet')
p.add_run('Member Browser').bold = True
p.add_run(' \u2192 CloudFront + API Gateway \u2192 Member Handler Lambda')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Member Handler').bold = True
p.add_run(' \u2192 AWS Connector \u2192 Customer AWS Account (STS AssumeRole)')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Member Handler').bold = True
p.add_run(' \u2192 Azure Connector \u2192 Azure Cost Management API (OAuth2)')
p = doc.add_paragraph(style='List Bullet')
p.add_run('Member Handler').bold = True
p.add_run(' \u2192 GCP Connector \u2192 GCP Cloud Billing API (JWT)')
p = doc.add_paragraph(style='List Bullet')
p.add_run('All Connectors').bold = True
p.add_run(' \u2192 Cost Normalizer \u2192 Unified Dashboard + Cost Cache')
p = doc.add_paragraph(style='List Bullet')
p.add_run('EventBridge (Daily)').bold = True
p.add_run(' \u2192 Tips Sync Lambda \u2192 S3 Knowledge Base \u2192 DynamoDB Tips Table')

doc.add_page_break()

# --- SECTION 3: CONNECTION FLOWS ---
add_heading('3. Connection Flows')

add_heading('3.1 Azure Connection Flow', level=2)
doc.add_paragraph('The Azure connection follows a multi-step process:')
steps = [
    'Member selects "Azure" in the Add Account modal',
    'Member enters Subscription ID, Tenant ID, and Client ID',
    'API stores account with connectionStatus: "pending"',
    'Member creates Service Principal in Azure AD (assigns "Cost Management Reader" role)',
    'Member provides Client Secret \u2192 KMS encrypts \u2192 stored in DynamoDB credentials map',
    'Member clicks "Test Connection"',
    'System decrypts Client Secret via KMS',
    'System authenticates: POST /oauth2/token (Tenant ID, Client ID, Secret)',
    'System queries: POST /costManagement/query (Subscription ID)',
    'On success: connectionStatus updated to "connected", lastTestedAt set',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

add_heading('3.2 GCP Connection Flow', level=2)
doc.add_paragraph('The GCP connection follows a similar multi-step process:')
steps = [
    'Member selects "GCP" in the Add Account modal',
    'Member enters Project ID and uploads Service Account JSON key file',
    'API validates key file structure (type, project_id, private_key_id, private_key, client_email)',
    'KMS encrypts private_key \u2192 stored in DynamoDB credentials map (status: "pending")',
    'Member clicks "Test Connection"',
    'System decrypts private key via KMS',
    'System creates self-signed JWT (iss=client_email, signed with RS256)',
    'System exchanges JWT for access token: POST /oauth2/token',
    'System queries: GET /billingAccounts/.../projects (Project ID)',
    'On success: connectionStatus updated to "connected", lastTestedAt set',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

add_heading('3.3 AWS Connection Flow (Unchanged)', level=2)
doc.add_paragraph(
    'The existing AWS connection flow remains identical: Member enters 12-digit Account ID \u2192 '
    'Deploys CloudFormation template \u2192 STS AssumeRole \u2192 ce:GetCostAndUsage \u2192 status: "connected".'
)

doc.add_page_break()

# --- SECTION 4: UNIFIED COST DATA PIPELINE ---
add_heading('4. Unified Cost Data Pipeline')

add_heading('4.1 Data Retrieval per Provider', level=2)
add_table(
    ['Provider', 'API', 'Authentication', 'Response Format'],
    [
        ['AWS', 'ce:GetCostAndUsage', 'STS AssumeRole', 'ResultsByTime \u2192 Groups \u2192 Metrics'],
        ['Azure', 'Microsoft.CostManagement/query', 'OAuth2 Bearer token', 'Rows: [cost, date, serviceName, currency]'],
        ['GCP', 'Cloud Billing API / BigQuery', 'JWT Bearer token', 'BillingAccount \u2192 Project \u2192 Services'],
    ]
)

add_heading('4.2 Common Normalized Schema', level=2)
doc.add_paragraph('All provider-specific cost data is transformed into a common schema:')
add_table(
    ['Field', 'Type', 'Example'],
    [
        ['date', 'String (ISO date)', '2026-01-15'],
        ['service_name', 'String', 'Virtual Machines / Compute Engine / Amazon EC2'],
        ['cost_amount', 'Float', '45.23'],
        ['currency', 'String', 'USD'],
        ['cloud_provider', 'String', 'aws | azure | gcp'],
        ['account_id', 'String', 'Provider-specific identifier'],
    ]
)

add_heading('4.3 Dashboard Aggregation', level=2)
doc.add_paragraph('The unified dashboard aggregation pipeline:')
steps = [
    'Retrieve cost data from each provider in parallel (AWS, Azure, GCP)',
    'Cost Normalizer transforms each response to common schema',
    'Aggregate totals per provider and calculate overall total',
    'Compute provider breakdown percentages',
    'Cache results in Cost_Cache_Table (key includes cloudProvider)',
    'Render: Combined summary + Provider pie chart + Filter toggles',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}', style='List Number')

add_heading('4.4 Filtering & Display Rules', level=2)
rules = [
    'Only accounts with connectionStatus: "connected" contribute to cost calculations',
    'Accounts with "pending" or "failed" status show a warning indicator (excluded from totals)',
    'If one provider\'s API fails, others still return data (graceful degradation)',
    'Provider color coding: AWS (orange #FF9900), Azure (blue #0078D4), GCP (red #EA4335)',
]
for rule in rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_page_break()

# --- SECTION 5: MULTI-CLOUD OPTIMIZATION TIPS ---
add_heading('5. Multi-Cloud Optimization Tips')

add_heading('5.1 Knowledge Base Structure', level=2)
doc.add_paragraph('The knowledge base is extended with provider-specific tip files:')
add_table(
    ['File', 'Status', 'Description'],
    [
        ['aws-cost-optimization-tips.json', 'Existing (32+ tips)', 'AWS-specific optimization recommendations'],
        ['azure-cost-optimization-tips.json', 'New', 'Azure-specific optimization recommendations'],
        ['gcp-cost-optimization-tips.json', 'New', 'GCP-specific optimization recommendations'],
    ]
)

add_heading('5.2 Daily Tips Sync (EventBridge + Lambda)', level=2)
add_table(
    ['Aspect', 'Detail'],
    [
        ['Schedule', 'Daily at 02:00 UTC (configurable)'],
        ['Trigger', 'EventBridge scheduled rule'],
        ['Source', 'S3 knowledge base JSON files'],
        ['Upsert Logic', 'Compare by (service + tipId), update only changed tips'],
        ['Removed Tips', 'Marked deprecated: true (never deleted)'],
        ['Failure Handling', 'If one provider fails, continue with others'],
        ['Logging', 'Summary: tips added, updated, deprecated, errors'],
    ]
)

add_heading('5.3 Member Tips Display', level=2)
rules = [
    'Members see only tips for providers they have connected',
    'Tips grouped by Cloud_Provider with provider-specific tabs',
    'Admin can filter and manage tips per provider',
    'Admin panel shows last sync timestamp and status per provider',
]
for rule in rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_page_break()

# --- SECTION 6: DATA MODELS ---
add_heading('6. Data Models')

add_heading('6.1 Accounts Table (MemberPortal-Accounts)', level=2)
doc.add_paragraph('Partition Key: memberEmail (String) | Sort Key: accountId (String)')
add_table(
    ['Attribute', 'Type', 'Required', 'Description'],
    [
        ['memberEmail', 'String (PK)', 'Yes', "Member's email"],
        ['accountId', 'String (SK)', 'Yes', 'AWS: 12-digit, Azure: Subscription UUID, GCP: Project ID'],
        ['cloudProvider', 'String', 'Yes*', '"aws" / "azure" / "gcp" (*defaults to "aws" for legacy)'],
        ['connectionStatus', 'String', 'Yes', '"pending" / "connected" / "failed"'],
        ['credentials', 'Map', 'No', 'Encrypted provider-specific credentials'],
        ['accountName', 'String', 'Yes', 'Display name'],
        ['addedAt', 'String', 'Yes', 'ISO 8601 timestamp'],
        ['lastTestedAt', 'String', 'No', 'Last successful connection test'],
        ['roleName', 'String', 'No', 'AWS only: IAM role name'],
    ]
)

add_heading('6.2 Credentials Map Structure', level=2)
doc.add_paragraph('Provider-specific credentials stored in the credentials attribute:')
add_table(
    ['Provider', 'Fields Stored'],
    [
        ['AWS', '(none \u2014 uses STS AssumeRole with roleName)'],
        ['Azure', 'tenantId, clientId, encryptedClientSecret (KMS ciphertext)'],
        ['GCP', 'clientEmail, projectId, privateKeyId, encryptedPrivateKey (KMS ciphertext)'],
    ]
)

add_heading('6.3 Tips Table (ViewMyBill-CostOptimizationTips)', level=2)
doc.add_paragraph('Partition Key: service (String) | Sort Key: tipId (String)')
add_table(
    ['Attribute', 'Type', 'Required', 'Description'],
    [
        ['service', 'String (PK)', 'Yes', 'Service name'],
        ['tipId', 'String (SK)', 'Yes', 'Unique tip ID (e.g., "azure-vm-001")'],
        ['cloudProvider', 'String', 'Yes*', '"aws" / "azure" / "gcp" (*backfill "aws" for legacy)'],
        ['category', 'String', 'Yes', 'Tip category'],
        ['title', 'String', 'Yes', 'Short title'],
        ['description', 'String', 'Yes', 'Detailed recommendation'],
        ['estimatedSavings', 'String', 'No', 'Expected savings range'],
        ['difficulty', 'String', 'No', '"easy" / "medium" / "hard"'],
        ['deprecated', 'Boolean', 'No', 'True if removed from source file'],
        ['lastSyncedAt', 'String', 'No', 'Last sync timestamp'],
    ]
)

add_heading('6.4 Cost Cache Table', level=2)
doc.add_paragraph(
    'Cache key format extended to include provider:\n'
    'Old: {memberEmail}#{accountId}#{dateRange}\n'
    'New: {memberEmail}#{cloudProvider}#{accountId}#{dateRange}\n\n'
    'This ensures the same accountId across different providers produces distinct cache entries.'
)

doc.add_page_break()

# --- SECTION 7: API CHANGES ---
add_heading('7. API Changes')

add_heading('7.1 Member Handler (Extended Routes)', level=2)
add_table(
    ['Method', 'Path', 'Change', 'Description'],
    [
        ['POST', '/members/accounts', 'Modified', 'Accept cloudProvider, route to provider-specific validation'],
        ['PUT', '/members/accounts', 'New', 'Update credentials (e.g., add Azure Client Secret)'],
        ['POST', '/members/accounts/test', 'Modified', 'Dispatch to correct connector based on cloudProvider'],
        ['GET', '/members/accounts', 'Modified', 'Return cloudProvider, backfill "aws" for legacy'],
        ['GET', '/members/dashboard-data', 'Modified', 'Aggregate cost data across all providers'],
        ['POST', '/members/accounts/ai-query', 'Modified', 'Include multi-cloud context in AI prompt'],
    ]
)

add_heading('7.2 Admin Handler (Extended Routes)', level=2)
add_table(
    ['Method', 'Path', 'Change', 'Description'],
    [
        ['GET', '/admin/tips', 'Modified', 'Accept ?cloudProvider= filter'],
        ['POST', '/admin/tips', 'Modified', 'Require cloudProvider field'],
        ['GET', '/admin/tips-sync/status', 'New', 'Per-provider sync status + last timestamp'],
    ]
)

add_heading('7.3 Input Validation Rules', level=2)
add_table(
    ['Provider', 'Field', 'Format'],
    [
        ['AWS', 'accountId', 'Exactly 12 digits (^\\d{12}$)'],
        ['Azure', 'subscriptionId', 'UUID format (^[0-9a-f]{8}-...-[0-9a-f]{12}$)'],
        ['Azure', 'tenantId', 'UUID format (same as above)'],
        ['GCP', 'projectId', '6-30 chars, lowercase + hyphens (^[a-z][a-z0-9-]{4,28}[a-z0-9]$)'],
    ]
)

doc.add_page_break()

# --- SECTION 8: SECURITY MODEL ---
add_heading('8. Security Model (Multi-Cloud Extension)')
add_table(
    ['Layer', 'Mechanism'],
    [
        ['Credential Encryption', 'AWS KMS with dedicated key (Azure secrets + GCP private keys)'],
        ['Decryption Scope', 'Only at connection-test or cost-retrieval time; plaintext discarded immediately'],
        ['API Response Security', 'Sensitive credentials never returned in API responses'],
        ['KMS Key Policy', 'Restricted to Member Portal Lambda execution role only'],
        ['KMS Failure Handling', 'Returns generic 500 error; detailed error logged internally'],
        ['Input Validation', 'Provider-specific format validation on all write operations'],
        ['Backward Compatibility', 'Legacy records without cloudProvider default to "aws"'],
        ['Authentication', 'Existing JWT token required for all multi-cloud operations'],
        ['Credential Storage', 'Envelope encryption: KMS encrypts data key, data key encrypts credentials'],
    ]
)

doc.add_page_break()

# --- SECTION 9: AI AGENT MULTI-CLOUD CONTEXT ---
add_heading('9. AI Agent Multi-Cloud Context')

add_heading('9.1 Enhanced AI Prompt Context', level=2)
doc.add_paragraph(
    'When a member sends an AI query, the system now includes in the prompt context:'
)
items = [
    'List of connected cloud providers and their account identifiers',
    'Provider-specific cost data scoped to the relevant provider',
    'Instructions for the model to provide cloud-specific recommendations',
    'Knowledge base tips filtered to connected providers only',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

add_heading('9.2 AI Behavior Rules', level=2)
add_table(
    ['Scenario', 'AI Response'],
    [
        ['Member asks about a connected provider', 'Scope context to that provider\'s accounts'],
        ['Member asks about an unconnected provider', 'Inform them no accounts are connected, suggest connecting'],
        ['General cost question', 'Include data from all connected providers'],
        ['Optimization question', 'Provide provider-specific recommendations'],
    ]
)

doc.add_page_break()

# --- SECTION 10: FRONTEND CHANGES ---
add_heading('10. Frontend Changes')

add_heading('10.1 Member Portal', level=2)
add_table(
    ['Element', 'Description'],
    [
        ['Add Account Modal', 'Provider selection step (3 cards: AWS, Azure, GCP with logos)'],
        ['Provider-Specific Forms', 'AWS: 12-digit ID | Azure: Subscription/Tenant/Client IDs | GCP: Project ID + file upload'],
        ['Accounts List', 'Provider icon + color coding next to each account'],
        ['Identifier Labels', '"Account ID" (AWS), "Subscription ID" (Azure), "Project ID" (GCP)'],
        ['Dashboard Header', 'Summary count of connected accounts per provider'],
        ['Dashboard Chart', 'Provider breakdown pie chart (% spend per provider)'],
        ['Filter Toggle', 'Filter cost data by Cloud_Provider'],
        ['Connection Instructions', 'Provider-specific setup guides (downloadable scripts)'],
    ]
)

add_heading('10.2 Admin Panel', level=2)
add_table(
    ['Element', 'Description'],
    [
        ['Tip Creation Form', '"Cloud Provider" dropdown (AWS, Azure, GCP)'],
        ['Tips List', 'Provider tabs/filter controls'],
        ['Sync Status', 'Per-provider last sync timestamp + status indicator'],
    ]
)

doc.add_page_break()

# --- SECTION 11: BACKWARD COMPATIBILITY ---
add_heading('11. Backward Compatibility Guarantees')
add_table(
    ['Aspect', 'Guarantee'],
    [
        ['Existing AWS accounts', 'Continue working without any changes'],
        ['Legacy records (no cloudProvider)', 'Default to "aws" in all read operations'],
        ['CloudFormation template generation', 'Unchanged for AWS accounts'],
        ['AWS connection test', 'Same behavior and response format'],
        ['AWS cost retrieval', 'Same STS AssumeRole + Cost Explorer flow'],
        ['AWS-only members', 'Same dashboard experience + provider icons added'],
        ['Existing API endpoints', 'No breaking changes'],
        ['Error messages', 'Existing AWS errors unchanged; new providers use same format'],
    ]
)

doc.add_page_break()

# --- SECTION 12: ERROR HANDLING ---
add_heading('12. Error Handling')

add_heading('12.1 Provider Connection Errors', level=2)
add_table(
    ['Scenario', 'Provider', 'HTTP', 'User Message'],
    [
        ['Invalid Service Principal', 'Azure', '400', 'Verify credentials and "Cost Management Reader" role'],
        ['Expired Client Secret', 'Azure', '400', 'Generate a new secret in Azure AD'],
        ['Insufficient permissions', 'Azure', '400', 'Verify "Cost Management Reader" role on subscription'],
        ['Invalid service account key', 'GCP', '400', 'Verify key file and "Billing Account Viewer" role'],
        ['Billing not enabled', 'GCP', '400', 'Ensure billing is enabled for this project'],
        ['KMS decryption failure', 'Any', '500', 'Unable to access credentials. Contact support.'],
        ['Provider API timeout', 'Any', '504', 'Connection timed out. Try again.'],
        ['Provider API rate limit', 'Any', '429', 'Rate limit reached. Wait and retry.'],
        ['Invalid cloudProvider', 'Any', '400', 'Supported values: aws, azure, gcp'],
        ['Duplicate account', 'Any', '409', 'Account already connected'],
        ['Malformed GCP key file', 'GCP', '400', 'Required: type, project_id, private_key_id, private_key, client_email'],
    ]
)

add_heading('12.2 Graceful Degradation', level=2)
rules = [
    'If one provider\'s API fails during dashboard load, other providers still return data',
    'Tips sync continues processing remaining providers if one fails',
    'Failed accounts excluded from cost calculations (warning shown to user)',
    'DynamoDB write throttling handled with exponential backoff (3 retries)',
]
for rule in rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_page_break()

# --- SECTION 13: INFRASTRUCTURE CHANGES ---
add_heading('13. Infrastructure Changes')

add_heading('13.1 Lambda Functions (Modified)', level=2)
add_table(
    ['Function', 'Change', 'Memory', 'Timeout'],
    [
        ['aws-bill-analyzer-member-api', 'Extended (multi-cloud connectors)', '256 MB', '120s'],
        ['aws-bill-analyzer-admin-api', 'Extended (provider filter)', '128 MB', '30s'],
        ['Tips Sync Lambda', 'Extended (Azure + GCP sources)', '128 MB', '60s'],
    ]
)

add_heading('13.2 New Dependencies (Member Handler)', level=2)
add_table(
    ['Package', 'Purpose'],
    [
        ['PyJWT', 'GCP JWT signing (already present for member auth)'],
        ['cryptography', 'RS256 key signing for GCP service account'],
        ['requests / urllib3', 'Azure OAuth2 + Cost Management API calls'],
    ]
)

add_heading('13.3 New AWS Resources', level=2)
add_table(
    ['Resource', 'Purpose'],
    [
        ['KMS Key (multi-cloud-credentials)', 'Encrypt Azure/GCP credentials at rest'],
        ['EventBridge Rule (tips-sync-daily)', 'Trigger daily tips sync at 02:00 UTC'],
        ['S3 objects (knowledge-base/)', 'azure-cost-optimization-tips.json, gcp-cost-optimization-tips.json'],
    ]
)

add_heading('13.4 DynamoDB Table Changes', level=2)
add_table(
    ['Table', 'Change'],
    [
        ['MemberPortal-Accounts', 'Add cloudProvider (String) + credentials (Map) attributes'],
        ['ViewMyBill-CostOptimizationTips', 'Add cloudProvider (String) + deprecated (Boolean) + lastSyncedAt (String)'],
        ['Cost_Cache_Table', 'Extended cache key format (includes cloudProvider)'],
    ]
)

doc.add_page_break()

# --- SECTION 14: MODULE STRUCTURE ---
add_heading('14. Module Structure')
doc.add_paragraph('New and modified files for the multi-cloud feature:')
add_table(
    ['Path', 'Status', 'Purpose'],
    [
        ['member-handler/lambda_function.py', 'Extended', 'Multi-cloud routing logic'],
        ['member-handler/cost_normalizer.py', 'New', 'Unified cost transformation'],
        ['member-handler/connectors/__init__.py', 'New', 'Connector package init'],
        ['member-handler/connectors/base_connector.py', 'New', 'ProviderConnector interface'],
        ['member-handler/connectors/aws_connector.py', 'New', 'Refactored from existing logic'],
        ['member-handler/connectors/azure_connector.py', 'New', 'OAuth2 + Cost Management'],
        ['member-handler/connectors/gcp_connector.py', 'New', 'JWT + Cloud Billing'],
        ['knowledge-base/azure-cost-optimization-tips.json', 'New', 'Azure tips source'],
        ['knowledge-base/gcp-cost-optimization-tips.json', 'New', 'GCP tips source'],
        ['admin-handler/lambda_function.py', 'Extended', 'Provider filter support'],
        ['tips-sync/lambda_function.py', 'Extended', 'Multi-provider sync'],
        ['members/members.js', 'Extended', 'Provider selection UI'],
        ['admin/admin.js', 'Extended', 'Provider tabs + sync status'],
    ]
)

doc.add_page_break()

# --- SECTION 15: KEY DESIGN DECISIONS ---
add_heading('15. Key Design Decisions')
add_table(
    ['#', 'Decision', 'Rationale'],
    [
        ['1', 'Provider Connector Pattern', 'Isolates provider logic; adding future providers requires only a new connector module'],
        ['2', 'Credentials in DynamoDB (KMS-encrypted)', 'Avoids separate secrets store; leverages existing table; KMS provides envelope encryption'],
        ['3', 'Backward-compatible schema', 'Zero disruption for existing members; no data migration needed'],
        ['4', 'Cost Normalizer as separate module', 'Single responsibility; testable independently; reusable across dashboard + AI'],
        ['5', 'Daily tips sync (not real-time)', 'Tips are curated content, not live data; daily is sufficient; reduces complexity'],
        ['6', 'Cloud Billing API over BigQuery for GCP', 'Simpler setup (no billing export required); BigQuery as optional enhancement'],
        ['7', 'Existing AWS flow unchanged', 'Risk mitigation; proven flow stays stable; multi-cloud is additive only'],
    ]
)

# --- SAVE ---
output_path = r'c:\Users\Michal\Desktop\PublicSite\SlashMyBill-MultiCloud-HLD.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
